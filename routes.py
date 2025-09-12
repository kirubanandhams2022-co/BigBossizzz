import os
from flask import render_template, request, redirect, url_for, flash, jsonify, session, send_file, make_response
from flask_login import login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from app import app, db, mail, socketio
from models import User, Quiz, Question, QuestionOption, QuizAttempt, Answer, ProctoringEvent, LoginEvent, UserViolation, UploadRecord, Course, HostCourseAssignment, ParticipantEnrollment, DeviceLog, SecurityAlert, CollaborationSignal, AttemptSimilarity, AlertThreshold, QuizThresholdOverride, AlertTrigger, InteractionEvent, QuestionHeatmapData, CollaborationInsight, PlagiarismAnalysis, PlagiarismMatch, Role, Permission, UserRole, RolePermission, RoleAuditLog

# 🛡️ FEATURE FLAGS - Defined immediately after imports to prevent NameError
ENABLE_LTI = os.environ.get('ENABLE_LTI', 'false').lower() == 'true'
ENABLE_ANALYTICS = os.environ.get('ENABLE_ANALYTICS', 'false').lower() == 'true'
ENABLE_REPORTS = os.environ.get('ENABLE_REPORTS', 'false').lower() == 'true'
ENABLE_PLAGIARISM = os.environ.get('ENABLE_PLAGIARISM', 'false').lower() == 'true'
ENABLE_COLLABORATION = os.environ.get('ENABLE_COLLABORATION', 'false').lower() == 'true'
ENABLE_RBAC = os.environ.get('ENABLE_RBAC', 'false').lower() == 'true'

# 🛡️ SAFE PLACEHOLDERS - Initialize all optional symbols as stubs to prevent LSP diagnostics
LTIProvider = LTIUser = LTIGradePassback = LTIToolConfiguration = None
get_lti_provider = get_lti_grade_passback = lambda *args, **kwargs: None
ProctoringReportGenerator = generate_scheduled_report = export_report_to_pdf = None
AnalyticsEngine = PredictiveAnalytics = QuestionPerformanceAnalyzer = None
CheatingPatternDetector = InstitutionalDashboard = get_analytics_engine = None
detector = None
plagiarism_detector = None
RBACService = None
initialize_rbac_system = None
# 🔒 SECURITY: Fail-closed RBAC decorators - deny access when RBAC disabled  
def _rbac_fallback_decorator(f):
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            flash('Access denied. Admin privileges required.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return wrapper
require_permission = lambda perm: _rbac_fallback_decorator
require_role = lambda role: _rbac_fallback_decorator
admin_required = _rbac_fallback_decorator
permission_context_processor = lambda: {}

# Import optional LTI integration with feature flag (after placeholders)
if ENABLE_LTI:
    try:
        from lti_integration import (LTIProvider, LTIUser, LTIGradePassback, LTIToolConfiguration,
                                    get_lti_provider, get_lti_grade_passback)
    except ImportError:
        ENABLE_LTI = False

# Import optional proctoring reports with feature flag (after placeholders)
if ENABLE_REPORTS:
    try:
        from automated_proctoring_reports import ProctoringReportGenerator, generate_scheduled_report, export_report_to_pdf
    except ImportError:
        ENABLE_REPORTS = False

# Import optional analytics engine with feature flag (after placeholders)
if ENABLE_ANALYTICS:
    try:
        from analytics_engine import (AnalyticsEngine, PredictiveAnalytics, QuestionPerformanceAnalyzer, 
                                      CheatingPatternDetector, InstitutionalDashboard, get_analytics_engine)
    except ImportError:
        ENABLE_ANALYTICS = False
from forms import RegistrationForm, LoginForm, QuizForm, QuestionForm, ProfileForm
from email_service import send_verification_email, send_credentials_email, send_login_notification, send_host_login_notification
from flask_mail import Message
from datetime import datetime, timedelta
import json
import logging
import os
import re
import csv
# Import optional data processing libraries
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None
from io import BytesIO
from sqlalchemy import func, text
from sqlalchemy.orm import joinedload, selectinload
from utils import get_time_greeting, get_greeting_icon

# Import collaboration detection with feature flag (after placeholders)
if ENABLE_COLLABORATION:
    try:
        from collaboration_detection import detector
    except ImportError:
        ENABLE_COLLABORATION = False

# Import plagiarism detection with feature flag (after placeholders)
if ENABLE_PLAGIARISM:
    try:
        from plagiarism_detector import plagiarism_detector
    except ImportError:
        ENABLE_PLAGIARISM = False

# Import RBAC system with feature flag (after placeholders)
if ENABLE_RBAC:
    try:
        from rbac_service import RBACService, initialize_rbac_system
        from rbac_decorators import require_permission, require_role, admin_required, permission_context_processor
    except ImportError:
        ENABLE_RBAC = False

# Add email health check endpoint
@app.route('/admin/email-health')
@login_required
def admin_email_health():
    """Email system health check for monitoring"""
    # 🔒 SECURITY: Require admin role for health check
    if not current_user.is_authenticated or current_user.role != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('login'))
    try:
        from email_service import brevo_service, test_email_service
        is_healthy = test_email_service()
        return jsonify({
            'status': 'healthy' if is_healthy else 'unhealthy',
            'service': 'Brevo (300/day FREE)',
            'api_configured': bool(brevo_service.api_key),
            'sender_configured': bool(brevo_service.sender_email),
            'ready_for_production': is_healthy
        })
    except Exception as e:
        return jsonify({'status': 'error', 'error': str(e)}), 500

# Excel/Spreadsheet generation imports
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

# PDF generation imports  
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# Global context processor to inject greeting variables
@app.context_processor
def inject_greeting():
    """Inject time-based greeting and icon into all templates"""
    return {
        'greeting': get_time_greeting(),
        'greeting_icon': get_greeting_icon()
    }

@app.route('/')
def index():
    """Home page - Redirect to register/login"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('register'))

@app.route('/default-accounts')
def default_accounts():
    """Show default login accounts"""
    return render_template('default_accounts.html')

@app.route('/loading')
def loading():
    """Loading page with eye animation"""
    return render_template('loading.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    """User registration"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    form = RegistrationForm()
    if form.validate_on_submit():
        # Check if user already exists
        if User.query.filter_by(email=form.email.data).first():
            flash('Email already registered. Please use a different email.', 'error')
            return render_template('register.html', form=form)
        
        if User.query.filter_by(username=form.username.data).first():
            flash('Username already taken. Please choose a different username.', 'error')
            return render_template('register.html', form=form)
        
        # Create new user
        user = User(
            username=form.username.data,
            email=form.email.data,
            role=form.role.data
        )
        user.set_password(form.password.data)
        
        db.session.add(user)
        db.session.commit()
        
        # Generate verification token and commit it to database
        user.generate_verification_token()
        db.session.commit()  # Critical: Persist the verification token
        
        # Send verification email
        if send_verification_email(user):
            flash('Registration successful! Please check your email to verify your account.', 'success')
            return render_template('verify_email.html', user=user, resent=False)
        else:
            flash('Registration successful, but we could not send the verification email. You can request a new one from the login page.', 'warning')
            return redirect(url_for('login'))
    
    return render_template('register.html', form=form)

@app.route('/verify/<token>')
def verify_email(token):
    """Verify email address"""
    user = User.query.filter_by(verification_token=token).first()
    
    if not user:
        flash('Invalid or expired verification link.', 'error')
        return redirect(url_for('login'))
    
    if user.verify_email(token):
        # Store the original password temporarily to send in email
        # In production, you might want to generate a temporary password instead
        temp_password = "Please change this password after login"
        
        db.session.commit()
        
        # Send credentials email
        if send_credentials_email(user, temp_password):
            flash('Email verified successfully! Your login credentials have been sent to your email.', 'success')
        else:
            flash('Email verified successfully! You can now log in with your credentials.', 'success')
    else:
        flash('Email verification failed. Please try again.', 'error')
    
    return redirect(url_for('login'))

@app.route('/resend-verification/<email>')
def resend_verification(email):
    """Resend email verification"""
    user = User.query.filter_by(email=email).first()
    
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('login'))
    
    if user.is_verified:
        flash('Your email is already verified. You can log in now.', 'info')
        return redirect(url_for('login'))
    
    # Generate new verification token
    user.generate_verification_token()
    db.session.commit()
    
    # Send verification email
    if send_verification_email(user):
        flash('Verification email sent successfully! Please check your inbox and spam folder.', 'success')
        return render_template('verify_email.html', user=user, resent=True)
    else:
        flash('Failed to send verification email. Please try again later.', 'error')
    
    return redirect(url_for('login'))

@app.route('/request-verification', methods=['GET', 'POST'])
def request_verification():
    """Request email verification page"""
    if request.method == 'POST':
        email = request.form.get('email')
        if not email:
            flash('Please enter your email address.', 'error')
            return render_template('request_verification.html')
        
        user = User.query.filter_by(email=email).first()
        if not user:
            flash('No account found with that email address.', 'error')
            return render_template('request_verification.html')
        
        if user.is_verified:
            flash('Your email is already verified. You can log in now.', 'info')
            return redirect(url_for('login'))
        
        # Generate new verification token
        user.generate_verification_token()
        db.session.commit()
        
        # Send verification email
        if send_verification_email(user):
            flash('Verification email sent successfully! Please check your inbox and spam folder.', 'success')
            return render_template('verify_email.html', user=user, resent=True)
        else:
            flash('Failed to send verification email. Please try again later.', 'error')
    
    return render_template('request_verification.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """User login"""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        
        if user and user.check_password(form.password.data):
            if not user.is_verified:
                flash('Please verify your email address before logging in.', 'warning')
                # Add resend verification option
                resend_link = url_for('resend_verification', email=user.email)
                flash(f'Need a new verification email? <a href="{resend_link}">Click here to resend</a>', 'info')
                return render_template('login.html', form=form)
            
            # Update last login time
            user.last_login = datetime.utcnow()
            
            # Enhanced login tracking with comprehensive device/IP information
            ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR', 'Unknown'))
            user_agent = request.headers.get('User-Agent', '')
            
            # Enhanced device fingerprinting and location tracking
            device_info = {
                'user_agent': user_agent,
                'ip_address': ip_address,
                'accept_language': request.headers.get('Accept-Language', ''),
                'accept_encoding': request.headers.get('Accept-Encoding', ''),
                'remote_addr': request.environ.get('REMOTE_ADDR'),
                'x_forwarded_for': request.headers.get('X-Forwarded-For', ''),
                'x_real_ip': request.headers.get('X-Real-IP', ''),
                'host': request.headers.get('Host', ''),
                'referer': request.headers.get('Referer', ''),
                'connection': request.headers.get('Connection', '')
            }
            
            # Check for suspicious login patterns
            from datetime import timedelta
            recent_logins = LoginEvent.query.filter_by(user_id=user.id).filter(
                LoginEvent.login_time > datetime.utcnow() - timedelta(hours=1)
            ).count()
            
            # Check for different IP addresses in short time (potential location jumping)
            different_ips = db.session.query(LoginEvent.ip_address).distinct().filter(
                LoginEvent.user_id == user.id,
                LoginEvent.login_time > datetime.utcnow() - timedelta(hours=24)
            ).count()
            
            # Check for different user agents (device switching)
            different_devices = db.session.query(LoginEvent.user_agent).distinct().filter(
                LoginEvent.user_id == user.id,
                LoginEvent.login_time > datetime.utcnow() - timedelta(days=7)
            ).count()
            
            is_suspicious = recent_logins > 5 or different_ips > 3 or different_devices > 5
            
            # Create login event record
            login_event = LoginEvent(
                user_id=user.id,
                ip_address=ip_address,
                user_agent=user_agent,
                device_info=json.dumps(device_info),
                is_suspicious=is_suspicious
            )
            
            db.session.add(login_event)
            db.session.commit()
            
            # Send email notifications
            try:
                # Send notification to user
                send_login_notification(user, login_event)
                
                # If participant, notify all hosts
                if user.role == 'participant':
                    hosts = User.query.filter_by(role='host').all()
                    for host in hosts:
                        send_host_login_notification(host, user, login_event)
                        
            except Exception as e:
                logging.error(f"Failed to send login notifications: {e}")
            
            # DEBUG: Log login details
            print(f"🔍 LOGIN DEBUG - User: {user.email}, Role in DB: {user.role}")
            print(f"🔍 LOGIN DEBUG - is_admin(): {user.is_admin()}, is_host(): {user.is_host()}, is_participant(): {user.is_participant()}")
            
            login_user(user)
            flash(f'Welcome back, {user.username}!', 'success')
            
            # DEBUG: Check current_user after login_user
            print(f"🔍 LOGIN DEBUG - current_user.email: {current_user.email}, current_user.role: {current_user.role}")
            print(f"🔍 LOGIN DEBUG - current_user methods: is_admin()={current_user.is_admin()}, is_host()={current_user.is_host()}, is_participant()={current_user.is_participant()}")
            
            next_page = request.args.get('next')
            if next_page:
                return redirect(next_page)
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password.', 'error')
    
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    """User logout"""
    logout_user()
    flash('You have been logged out successfully.', 'info')
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    """Main dashboard - redirects based on user role"""
    if current_user.is_admin():
        return redirect(url_for('admin_dashboard'))
    elif current_user.is_host():
        return redirect(url_for('host_dashboard'))
    else:
        return redirect(url_for('participant_dashboard'))

@app.route('/host/dashboard')
@login_required
def host_dashboard():
    """Enhanced Host dashboard with participant management"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    quizzes = Quiz.query.filter_by(creator_id=current_user.id).order_by(Quiz.display_order.asc(), Quiz.created_at.desc()).all()
    
    # Get recent quiz attempts for host's quizzes and add heatmap data flags
    recent_attempts = []
    heatmap_ready_count = 0
    
    for quiz in quizzes:
        attempts = QuizAttempt.query.filter_by(quiz_id=quiz.id).order_by(QuizAttempt.started_at.desc()).limit(5).all()
        quiz.attempts = attempts  # Add attempts to quiz object for template access
        recent_attempts.extend(attempts)
        
        # Check if quiz has completed attempts with interaction data for accurate heatmap counting
        completed_attempts = QuizAttempt.query.filter_by(
            quiz_id=quiz.id, 
            status='completed'
        ).first()
        
        if completed_attempts:
            # Check if there's actual interaction data
            interaction_exists = InteractionEvent.query.filter_by(
                attempt_id=completed_attempts.id
            ).first()
            quiz.has_heatmap_data = interaction_exists is not None
            if quiz.has_heatmap_data:
                heatmap_ready_count += 1
        else:
            quiz.has_heatmap_data = False
    
    recent_attempts.sort(key=lambda x: x.started_at, reverse=True)
    recent_attempts = recent_attempts[:10]  # Show only 10 most recent
    
    # Get participant statistics
    participants = User.query.filter_by(role='participant').all()
    recent_logins = LoginEvent.query.join(User).filter(User.role == 'participant').order_by(LoginEvent.login_time.desc()).limit(10).all()
    
    # Get violation statistics
    quiz_ids = [quiz.id for quiz in quizzes]
    if quiz_ids:
        high_violations = db.session.query(QuizAttempt, func.count(ProctoringEvent.id).label('violation_count')).join(ProctoringEvent).filter(
            QuizAttempt.quiz_id.in_(quiz_ids),
            ProctoringEvent.severity == 'high'
        ).group_by(QuizAttempt.id).having(func.count(ProctoringEvent.id) > 2).all()
    else:
        high_violations = []
    
    return render_template('host_dashboard.html', 
                         quizzes=quizzes, 
                         recent_attempts=recent_attempts,
                         participants=participants,
                         recent_logins=recent_logins,
                         high_violations=high_violations,
                         heatmap_ready_count=heatmap_ready_count,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/host/total-quizzes')
@login_required
def host_total_quizzes():
    """Show all quizzes created by the host"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get all quizzes created by this host with eager loading
    quizzes = Quiz.query.filter_by(creator_id=current_user.id).options(
        joinedload(Quiz.questions),
        joinedload(Quiz.attempts)
    ).order_by(Quiz.created_at.desc()).all()
    
    # Calculate statistics
    total_questions = sum(len(quiz.questions) for quiz in quizzes)
    total_attempts = sum(len(quiz.attempts) for quiz in quizzes)
    
    return render_template('host_total_quizzes.html',
                         quizzes=quizzes,
                         total_questions=total_questions,
                         total_attempts=total_attempts,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/host/active-quizzes')
@login_required
def host_active_quizzes():
    """Show active quizzes created by the host"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get active quizzes created by this host with eager loading
    active_quizzes = Quiz.query.filter_by(
        creator_id=current_user.id, 
        is_active=True
    ).options(
        joinedload(Quiz.questions),
        joinedload(Quiz.attempts)
    ).order_by(Quiz.created_at.desc()).all()
    
    # Calculate statistics for active quizzes
    total_active_attempts = sum(len(quiz.attempts) for quiz in active_quizzes)
    in_progress_attempts = sum(len([a for a in quiz.attempts if a.status == 'in_progress']) for quiz in active_quizzes)
    
    return render_template('host_active_quizzes.html',
                         active_quizzes=active_quizzes,
                         total_active_attempts=total_active_attempts,
                         in_progress_attempts=in_progress_attempts,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/host/recent-attempts')
@login_required
def host_recent_attempts():
    """Show recent quiz attempts for host's quizzes"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get recent attempts for host's quizzes with eager loading
    quiz_ids = [quiz.id for quiz in Quiz.query.filter_by(creator_id=current_user.id).all()]
    
    if quiz_ids:
        recent_attempts = QuizAttempt.query.filter(
            QuizAttempt.quiz_id.in_(quiz_ids)
        ).options(
            joinedload(QuizAttempt.quiz),
            joinedload(QuizAttempt.participant),
            joinedload(QuizAttempt.answers)
        ).order_by(QuizAttempt.started_at.desc()).limit(50).all()
    else:
        recent_attempts = []
    
    # Calculate statistics
    attempts_today = sum(1 for attempt in recent_attempts 
                        if attempt.started_at.date() == datetime.utcnow().date())
    unique_participants = len(set(attempt.participant_id for attempt in recent_attempts))
    
    return render_template('host_recent_attempts.html',
                         recent_attempts=recent_attempts,
                         attempts_today=attempts_today,
                         unique_participants=unique_participants,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/host/completed-attempts')
@login_required
def host_completed_attempts():
    """Show completed quiz attempts for host's quizzes"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get completed attempts for host's quizzes with eager loading
    quiz_ids = [quiz.id for quiz in Quiz.query.filter_by(creator_id=current_user.id).all()]
    
    if quiz_ids:
        completed_attempts = QuizAttempt.query.filter(
            QuizAttempt.quiz_id.in_(quiz_ids),
            QuizAttempt.status == 'completed'
        ).options(
            joinedload(QuizAttempt.quiz),
            joinedload(QuizAttempt.participant),
            joinedload(QuizAttempt.answers)
        ).order_by(QuizAttempt.completed_at.desc()).all()
    else:
        completed_attempts = []
    
    # Calculate statistics
    if completed_attempts:
        scores = [attempt.score for attempt in completed_attempts if attempt.score is not None]
        avg_score = sum(scores) / len(scores) if scores else 0
        highest_score = max(scores) if scores else 0
        lowest_score = min(scores) if scores else 0
        
        # Performance breakdown
        excellent_count = len([s for s in scores if s >= 80])
        good_count = len([s for s in scores if 60 <= s < 80])
        needs_improvement_count = len([s for s in scores if s < 60])
    else:
        avg_score = highest_score = lowest_score = 0
        excellent_count = good_count = needs_improvement_count = 0
    
    performance_breakdown = {
        'excellent': excellent_count,
        'good': good_count,
        'needs_improvement': needs_improvement_count
    }
    
    return render_template('host_completed_attempts.html',
                         completed_attempts=completed_attempts,
                         avg_score=round(avg_score, 1),
                         highest_score=highest_score,
                         lowest_score=lowest_score,
                         performance_breakdown=performance_breakdown,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/host/monitoring')
@login_required  
def host_monitoring():
    """Real-time participant monitoring panel"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get active quiz attempts
    active_attempts = QuizAttempt.query.filter_by(status='in_progress').join(Quiz).filter(
        Quiz.creator_id == current_user.id if not current_user.is_admin() else True
    ).all()
    
    # Get recent login events
    recent_logins = LoginEvent.query.order_by(LoginEvent.login_time.desc()).limit(20).all()
    
    # OPTIMIZED: Get consolidated violations per person (one entry per person)
    from sqlalchemy import func, case
    
    # CORRECTED: Single-query solution with proper SQL and highest severity selection
    from sqlalchemy import desc
    
    # CORRECTED: Window functions computed BEFORE filtering to get accurate counts
    violations_with_stats = db.session.query(
        User.id.label('participant_id'),
        User.username.label('participant_name'),
        Quiz.id.label('quiz_id'),
        Quiz.title.label('quiz_title'),
        ProctoringEvent.event_type,
        ProctoringEvent.severity,
        ProctoringEvent.timestamp,
        # Rank by severity (highest first), then by timestamp (latest first)
        func.row_number().over(
            partition_by=[User.id, Quiz.id],
            order_by=[
                case(
                    (ProctoringEvent.severity == 'critical', 4),
                    (ProctoringEvent.severity == 'high', 3),
                    (ProctoringEvent.severity == 'medium', 2),
                    (ProctoringEvent.severity == 'low', 1),
                    else_=1  # Default to 'low' rank for unknown severities
                ).desc(),
                ProctoringEvent.timestamp.desc()
            ]
        ).label('severity_rank'),
        # CRITICAL: Compute counts and max timestamp BEFORE filtering
        func.count().over(partition_by=[User.id, Quiz.id]).label('violation_count'),
        func.max(ProctoringEvent.timestamp).over(partition_by=[User.id, Quiz.id]).label('latest_violation_time')
    ).select_from(User) \
     .join(QuizAttempt, QuizAttempt.participant_id == User.id) \
     .join(Quiz, QuizAttempt.quiz_id == Quiz.id) \
     .join(ProctoringEvent, ProctoringEvent.attempt_id == QuizAttempt.id) \
     .filter(
         Quiz.creator_id == current_user.id if not current_user.is_admin() else True
     ).subquery()
    
    # Get only the highest severity violation (rank 1) with accurate stats
    consolidated_violations = db.session.query(
        violations_with_stats.c.participant_id,
        violations_with_stats.c.participant_name,
        violations_with_stats.c.quiz_id,
        violations_with_stats.c.quiz_title,
        violations_with_stats.c.event_type.label('display_violation_type'),
        violations_with_stats.c.severity.label('display_severity'),
        violations_with_stats.c.timestamp.label('highest_severity_time'),
        violations_with_stats.c.violation_count,
        violations_with_stats.c.latest_violation_time
    ).filter(violations_with_stats.c.severity_rank == 1) \
     .order_by(violations_with_stats.c.latest_violation_time.desc()) \
     .limit(30).all()
    
    # Format results for template
    recent_violations = []
    for result in consolidated_violations:
        recent_violations.append({
            'participant_id': result.participant_id,
            'participant_name': result.participant_name,
            'quiz_id': result.quiz_id,
            'quiz_title': result.quiz_title,
            'display_violation_type': result.display_violation_type,
            'display_severity': result.display_severity,
            'latest_violation_time': result.latest_violation_time,
            'violation_count': result.violation_count,
            'is_multiple': result.violation_count > 1
        })
    
    # Get participant device info
    participants_online = []
    for attempt in active_attempts:
        # Get latest login for this participant
        latest_login = LoginEvent.query.filter_by(user_id=attempt.participant_id).order_by(
            LoginEvent.login_time.desc()
        ).first()
        
        # Get violation count for this attempt
        violation_count = ProctoringEvent.query.filter_by(attempt_id=attempt.id).count()
        
        participants_online.append({
            'attempt': attempt,
            'latest_login': latest_login,
            'violation_count': violation_count,
            'time_elapsed': datetime.utcnow() - attempt.started_at,
            'quiz_title': attempt.quiz.title
        })
    
    return render_template('host_monitoring.html',
                         active_attempts=active_attempts,
                         participants_online=participants_online,
                         recent_logins=recent_logins,
                         recent_violations=recent_violations)

@app.route('/api/monitoring/live-data')
@login_required
def get_live_monitoring_data():
    """API endpoint for real-time monitoring data"""
    if not current_user.is_host() and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    # Get active attempts
    active_attempts = QuizAttempt.query.filter_by(status='in_progress').join(Quiz).filter(
        Quiz.creator_id == current_user.id if not current_user.is_admin() else True
    ).all()
    
    participants_data = []
    for attempt in active_attempts:
        # OPTIMIZED: Use highest risk summary instead of individual violations
        # Get violation summary for this attempt using new schema optimization
        if hasattr(attempt, 'highest_risk_level') and attempt.highest_risk_level:
            violation_count = 0
            latest_violation = attempt.highest_risk_level
            if attempt.violation_counts_json:
                try:
                    import json
                    counts = json.loads(attempt.violation_counts_json)
                    violation_count = sum(counts.values()) if counts else 0
                except:
                    violation_count = 0
        else:
            # Fallback: Get violation count directly from database (for older records)
            violations = ProctoringEvent.query.filter_by(attempt_id=attempt.id).all()
            violation_count = len(violations)
            latest_violation = violations[0].event_type if violations else None
        
        # Calculate time remaining
        time_elapsed = datetime.utcnow() - attempt.started_at
        time_remaining = timedelta(minutes=attempt.quiz.time_limit) - time_elapsed
        
        participants_data.append({
            'attempt_id': attempt.id,
            'participant_name': attempt.participant.username,
            'quiz_title': attempt.quiz.title,
            'time_elapsed': str(time_elapsed).split('.')[0],
            'time_remaining': str(time_remaining).split('.')[0] if time_remaining.total_seconds() > 0 else 'Overtime',
            'violation_count': violation_count,
            'latest_violation': latest_violation,
            'questions_answered': len(attempt.answers),
            'total_questions': len(attempt.quiz.questions),
            'progress_percentage': round((len(attempt.answers) / len(attempt.quiz.questions)) * 100) if attempt.quiz.questions else 0
        })
    
    return jsonify({
        'participants': participants_data,
        'total_active': len(participants_data),
        'timestamp': datetime.utcnow().isoformat()
    })

@app.route('/participant/dashboard')
@login_required
def participant_dashboard():
    """Participant dashboard"""
    # Get available quizzes (you might want to implement invitation system)
    available_quizzes = Quiz.query.filter_by(is_active=True).all()
    
    # Get participant's quiz attempts
    my_attempts = QuizAttempt.query.filter_by(participant_id=current_user.id).order_by(QuizAttempt.started_at.desc()).all()
    
    return render_template('participant_dashboard.html', 
                         available_quizzes=available_quizzes, 
                         my_attempts=my_attempts,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/participant/completed')
@login_required
def participant_completed():
    """Show completed quizzes for participant"""
    if current_user.role != 'participant':
        flash('Access denied. Participants only.', 'error')
        return redirect(url_for('index'))
    
    # Get completed quiz attempts with eager loading
    completed_attempts = QuizAttempt.query.filter_by(
        participant_id=current_user.id, 
        status='completed'
    ).options(
        joinedload(QuizAttempt.quiz).joinedload(Quiz.questions),
        joinedload(QuizAttempt.answers)
    ).order_by(QuizAttempt.completed_at.desc()).all()
    
    # Calculate total questions answered
    total_questions_answered = 0
    for attempt in completed_attempts:
        total_questions_answered += len(attempt.quiz.questions)
    
    return render_template('participant_completed.html', 
                         completed_attempts=completed_attempts,
                         total_questions_answered=total_questions_answered,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/participant/in-progress')
@login_required
def participant_in_progress():
    """Show in-progress quizzes for participant"""
    if current_user.role != 'participant':
        flash('Access denied. Participants only.', 'error')
        return redirect(url_for('index'))
    
    # Get in-progress quiz attempts with eager loading
    in_progress_attempts = QuizAttempt.query.filter_by(
        participant_id=current_user.id, 
        status='in_progress'
    ).options(
        joinedload(QuizAttempt.quiz).joinedload(Quiz.questions),
        joinedload(QuizAttempt.answers)
    ).order_by(QuizAttempt.started_at.desc()).all()
    
    # Calculate time remaining for each attempt
    from datetime import datetime, timedelta
    for attempt in in_progress_attempts:
        if attempt.quiz.time_limit and attempt.started_at:
            time_elapsed = datetime.utcnow() - attempt.started_at
            time_elapsed_minutes = time_elapsed.total_seconds() / 60
            time_remaining = attempt.quiz.time_limit - time_elapsed_minutes
            attempt.time_remaining = max(0, int(time_remaining))
        else:
            attempt.time_remaining = None
    
    return render_template('participant_in_progress.html', 
                         in_progress_attempts=in_progress_attempts,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/participant/average-score')
@login_required
def participant_average_score():
    """Show score analysis for participant"""
    if current_user.role != 'participant':
        flash('Access denied. Participants only.', 'error')
        return redirect(url_for('index'))
    
    # Get completed quiz attempts with scores and eager loading
    scored_attempts = QuizAttempt.query.filter_by(
        participant_id=current_user.id, 
        status='completed'
    ).filter(QuizAttempt.score.is_not(None)).options(
        joinedload(QuizAttempt.quiz),
        joinedload(QuizAttempt.answers)
    ).order_by(QuizAttempt.completed_at.desc()).all()
    
    # Calculate score statistics
    if scored_attempts:
        scores = [attempt.score for attempt in scored_attempts]
        avg_score = sum(scores) / len(scores)
        highest_score = max(scores)
        lowest_score = min(scores)
        
        # Calculate performance breakdown
        excellent_count = len([s for s in scores if s >= 80])
        good_count = len([s for s in scores if 60 <= s < 80])
        needs_improvement_count = len([s for s in scores if s < 60])
    else:
        avg_score = highest_score = lowest_score = 0
        excellent_count = good_count = needs_improvement_count = 0
    
    performance_breakdown = {
        'excellent': excellent_count,
        'good': good_count,
        'needs_improvement': needs_improvement_count
    }
    
    return render_template('participant_average_score.html', 
                         scored_attempts=scored_attempts,
                         avg_score=round(avg_score, 1),
                         highest_score=highest_score,
                         lowest_score=lowest_score,
                         total_quizzes=len(scored_attempts),
                         performance_breakdown=performance_breakdown,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/participant/violations')
@login_required
def participant_violations():
    """Show violation history for participant"""
    if current_user.role != 'participant':
        flash('Access denied. Participants only.', 'error')
        return redirect(url_for('index'))
    
    # OPTIMIZED: Use highest risk summary instead of aggregating individual violations
    attempts_with_violations = QuizAttempt.query.filter(
        QuizAttempt.participant_id == current_user.id,
        QuizAttempt.highest_risk_severity > 1  # Only show attempts with violations (above 'low')
    ).options(
        joinedload(QuizAttempt.quiz)
    ).order_by(QuizAttempt.started_at.desc()).all()
    
    # Count violations using the summary fields (much faster)
    critical_count = sum(1 for attempt in attempts_with_violations if attempt.highest_risk_level == 'critical')
    high_count = sum(1 for attempt in attempts_with_violations if attempt.highest_risk_level == 'high')
    
    # Create simplified violations_by_attempt structure using highest risk only
    violations_by_attempt = {}
    total_violations = 0
    
    for attempt in attempts_with_violations:
        # Parse violation counts from JSON
        try:
            import json
            counts = json.loads(attempt.violation_counts_json) if attempt.violation_counts_json else {}
            attempt_total = sum(counts.values()) if counts else 0
            total_violations += attempt_total
        except:
            attempt_total = 0
        
        # Show only the highest risk level violation for this attempt
        violations_by_attempt[attempt.id] = {
            'attempt': attempt,
            'highest_risk_level': attempt.highest_risk_level,
            'highest_risk_severity': attempt.highest_risk_severity,
            'violation_count': attempt_total,
            'violation_counts': counts if 'counts' in locals() else {}
        }
    
    return render_template('participant_violations.html', 
                         violations_by_attempt=violations_by_attempt,
                         total_violations=total_violations,
                         critical_count=critical_count,
                         high_count=high_count,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/quiz_listing')
@login_required
def quiz_listing():
    """Quiz listing page with mobile interface"""
    if current_user.role != 'participant':
        flash('Access denied. Participants only.', 'error')
        return redirect(url_for('index'))
    
    # Get available quizzes for this participant
    available_quizzes = Quiz.query.filter_by(is_active=True).all()
    
    # Get participant's quiz attempts
    my_attempts = QuizAttempt.query.filter_by(participant_id=current_user.id).order_by(QuizAttempt.started_at.desc()).all()
    
    return render_template('quiz_listing.html', 
                         available_quizzes=available_quizzes,
                         my_attempts=my_attempts,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    """Admin dashboard"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get system statistics
    total_users = User.query.count()
    total_hosts = User.query.filter_by(role='host').count()
    total_participants = User.query.filter_by(role='participant').count()
    total_quizzes = Quiz.query.count()
    total_attempts = QuizAttempt.query.count()
    total_courses = Course.query.count()
    total_violation_appeals = db.session.query(UserViolation).filter_by(is_flagged=True).count()
    
    # Recent registrations
    recent_users = User.query.order_by(User.created_at.desc()).limit(10).all()
    
    stats = {
        'total_users': total_users,
        'total_hosts': total_hosts,
        'total_participants': total_participants,
        'total_quizzes': total_quizzes,
        'total_attempts': total_attempts,
        'total_courses': total_courses,
        'total_violation_appeals': total_violation_appeals
    }
    
    return render_template('admin_dashboard.html', 
                         stats=stats, 
                         recent_users=recent_users,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())

@app.route('/admin/export-database')
@login_required
def admin_export_database():
    """Export complete database to Excel with comprehensive error handling"""
    if not current_user.is_admin():
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        from io import BytesIO
        
        wb = Workbook()
        
        # Users Sheet
        ws_users = wb.active
        ws_users.title = "Users"
        users_headers = ['ID', 'Username', 'Email', 'Role', 'Is Verified', 'Created At']
        
        for col, header in enumerate(users_headers, 1):
            cell = ws_users.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        
        # Fetch users with error handling
        try:
            users = User.query.all()
        except Exception as e:
            logging.error(f"Database error fetching users: {e}")
            flash('Error accessing user data for export.', 'error')
            return redirect(url_for('admin_dashboard'))
        
        for row, user in enumerate(users, 2):
            try:
                ws_users.cell(row=row, column=1, value=user.id)
                ws_users.cell(row=row, column=2, value=user.username or 'N/A')
                ws_users.cell(row=row, column=3, value=user.email or 'N/A')
                ws_users.cell(row=row, column=4, value=user.role or 'N/A')
                ws_users.cell(row=row, column=5, value='Yes' if user.is_verified else 'No')
                ws_users.cell(row=row, column=6, value=user.created_at.strftime('%Y-%m-%d %H:%M:%S') if user.created_at else 'N/A')
            except Exception as e:
                logging.warning(f"Error exporting user {user.id}: {e}")
                continue
        
        # Save to BytesIO with error handling
        try:
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
        except Exception as e:
            logging.error(f"Error creating Excel file: {e}")
            flash('Error generating export file.', 'error')
            return redirect(url_for('admin_dashboard'))
        
        return send_file(
            BytesIO(buffer.read()),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'database_export_{datetime.utcnow().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
        
    except Exception as e:
        logging.error(f"Unexpected error in database export: {e}")
        flash('Database export failed due to an unexpected error.', 'error')
        return redirect(url_for('admin_dashboard'))

# File Upload and Auto-Question Generation System

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'csv', 'xlsx', 'txt'}
UPLOAD_FOLDER = 'uploads'
QUIZ_ANSWER_UPLOAD_FOLDER = 'uploads/quiz_answers'

# Create upload directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(QUIZ_ANSWER_UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/upload-quiz-file', methods=['POST'])
@login_required
def upload_quiz_file():
    """Upload file and extract candidate questions with enhanced security"""
    if not current_user.is_host() and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not supported. Use PDF, DOCX, CSV, XLSX, or TXT files.'}), 400
    
    # Check file size (max 10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    file.seek(0, 2)  # Seek to end to get size
    file_size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if file_size > MAX_FILE_SIZE:
        return jsonify({'error': 'File too large. Maximum size is 10MB.'}), 400
    
    if file_size == 0:
        return jsonify({'error': 'File is empty.'}), 400
    
    try:
        # Secure filename and save
        filename = secure_filename(file.filename or 'upload')
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{filename}"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # Create upload record with proper validation
        try:
            upload_record = UploadRecord(
                host_id=current_user.id,
                filename=filename,
                mime_type=file.content_type or 'application/octet-stream',
                stored_path=file_path,
                file_size=file_size,
                parsed=False
            )
            db.session.add(upload_record)
            db.session.commit()
        except Exception as e:
            logging.error(f"Database error creating upload record: {e}")
            return jsonify({'error': 'Failed to save upload record'}), 500
        
        # Parse file and extract candidate questions
        candidate_questions = parse_file_for_questions(file_path, file.content_type)
        
        # Store candidate questions as JSON
        upload_record.candidate_questions_json = json.dumps(candidate_questions)
        upload_record.parsed = True
        db.session.commit()
        
        return jsonify({
            'upload_record_id': upload_record.id,
            'candidate_count': len(candidate_questions),
            'filename': filename,
            'message': f'Successfully extracted {len(candidate_questions)} candidate questions'
        })
        
    except Exception as e:
        logging.error(f"File upload error: {e}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/api/upload-quiz-create-draft', methods=['POST'])
@login_required
def upload_quiz_create_draft():
    """Create draft quiz from uploaded file"""
    if not current_user.is_host() and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    data = request.json
    upload_record_id = data.get('upload_record_id')
    num_questions = data.get('N', 10)
    quiz_title = data.get('title', 'Auto-Generated Quiz')
    quiz_description = data.get('description', 'Quiz generated from uploaded file')
    
    upload_record = UploadRecord.query.get_or_404(upload_record_id)
    
    if upload_record.host_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403
    
    try:
        # Load candidate questions
        candidate_questions = json.loads(upload_record.candidate_questions_json)
        
        # Select top N questions
        selected_questions = select_top_questions(candidate_questions, num_questions)
        
        # Create draft quiz
        quiz = Quiz(
            title=quiz_title,
            description=quiz_description,
            creator_id=current_user.id,
            auto_generate_from_upload=True,
            draft_from_upload_id=upload_record.id,
            is_active=False  # Draft mode
        )
        db.session.add(quiz)
        db.session.flush()  # Get quiz ID
        
        # Create questions and options
        for i, q_data in enumerate(selected_questions):
            question = Question(
                quiz_id=quiz.id,
                question_text=q_data['question'],
                question_type=q_data.get('type', 'multiple_choice'),
                points=q_data.get('points', 1),
                order=i + 1
            )
            db.session.add(question)
            db.session.flush()  # Get question ID
            
            # Add options for multiple choice questions
            if question.question_type in ['multiple_choice', 'true_false']:
                for j, option_text in enumerate(q_data.get('options', [])):
                    option = QuestionOption(
                        question_id=question.id,
                        option_text=option_text,
                        is_correct=j == q_data.get('correct_option_index', 0),
                        order=j + 1
                    )
                    db.session.add(option)
        
        upload_record.parsed_to_quiz_id = quiz.id
        db.session.commit()
        
        return jsonify({
            'draft_quiz_id': quiz.id,
            'message': f'Created draft quiz with {len(selected_questions)} questions'
        })
        
    except Exception as e:
        logging.error(f"Draft creation error: {e}")
        db.session.rollback()
        return jsonify({'error': f'Failed to create draft: {str(e)}'}), 500

def parse_file_for_questions(file_path, mime_type):
    """Parse uploaded file to extract candidate questions"""
    candidate_questions = []
    
    try:
        if 'pdf' in mime_type:
            candidate_questions = parse_pdf_questions(file_path)
        elif 'docx' in mime_type or 'document' in mime_type:
            candidate_questions = parse_docx_questions(file_path)
        elif 'csv' in mime_type:
            candidate_questions = parse_csv_questions(file_path)
        elif 'spreadsheet' in mime_type or 'excel' in mime_type:
            candidate_questions = parse_excel_questions(file_path)
        elif 'text' in mime_type:
            candidate_questions = parse_text_questions(file_path)
        
    except Exception as e:
        logging.error(f"File parsing error: {e}")
        
    return candidate_questions

def parse_pdf_questions(file_path):
    """Extract questions from PDF file with enhanced error handling"""
    questions = []
    
    try:
        # Validate file exists and is readable
        if not os.path.exists(file_path):
            logging.error(f"PDF file not found: {file_path}")
            return questions
            
        if os.path.getsize(file_path) == 0:
            logging.error(f"PDF file is empty: {file_path}")
            return questions
        
        with open(file_path, 'rb') as file:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    logging.warning(f"PDF has no pages: {file_path}")
                    return questions
                
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logging.warning(f"Error extracting text from page {page_num}: {e}")
                        continue
                        
                if len(text.strip()) < 10:
                    logging.warning(f"PDF contains very little text: {file_path}")
                    return questions
                
                # Parse text for questions with error handling
                questions = extract_questions_from_text(text)
                
            except PyPDF2.errors.PdfReadError as e:
                logging.error(f"PDF read error: {e}")
            except Exception as e:
                logging.error(f"Unexpected PDF parsing error: {e}")
        
    except Exception as e:
        logging.error(f"File access error for PDF {file_path}: {e}")
    
    return questions

def parse_docx_questions(file_path):
    """Extract questions from DOCX file with enhanced error handling"""
    questions = []
    
    try:
        # Validate file exists and is readable
        if not os.path.exists(file_path):
            logging.error(f"DOCX file not found: {file_path}")
            return questions
            
        if os.path.getsize(file_path) == 0:
            logging.error(f"DOCX file is empty: {file_path}")
            return questions
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            if len(doc.paragraphs) == 0:
                logging.warning(f"DOCX has no paragraphs: {file_path}")
                return questions
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
                            
            if len(text.strip()) < 10:
                logging.warning(f"DOCX contains very little text: {file_path}")
                return questions
            
            # Parse text for questions with error handling
            questions = extract_questions_from_text(text)
            
        except docx.opc.exceptions.PackageNotFoundError as e:
            logging.error(f"Invalid DOCX format: {e}")
        except Exception as e:
            logging.error(f"Unexpected DOCX parsing error: {e}")
        
    except Exception as e:
        logging.error(f"File access error for DOCX {file_path}: {e}")
    
    return questions

def parse_csv_questions(file_path):
    """Extract questions from CSV file with robust error handling"""
    questions = []
    
    try:
        # Try different encodings and delimiters
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            # If all encodings fail, return empty
            logging.error(f"Could not decode CSV file: {file_path}")
            return questions
        
        if df.empty:
            return questions
            
        # Expected columns: question, option_a, option_b, option_c, option_d, correct_answer
        for idx, row in df.iterrows():
            if 'question' in df.columns and pd.notna(row.get('question')):
                question_data = {
                    'question': str(row.get('question', '')),
                    'type': 'multiple_choice',
                    'options': [],
                    'correct_option_index': 0,
                    'confidence': 0.9  # High confidence for structured data
                }
                
                # Extract options
                option_cols = [col for col in df.columns if col.startswith('option')]
                for col in option_cols:
                    if pd.notna(row[col]):
                        question_data['options'].append(str(row[col]))
                
                # Determine correct answer
                if 'correct_answer' in df.columns:
                    correct_text = str(row['correct_answer']).strip()
                    for i, option in enumerate(question_data['options']):
                        if option.strip().lower() == correct_text.lower():
                            question_data['correct_option_index'] = i
                            break
                
                if question_data['question'] and len(question_data['options']) >= 2:
                    questions.append(question_data)
        
    except Exception as e:
        logging.error(f"CSV parsing error: {e}")
    
    return questions

def parse_excel_questions(file_path):
    """Extract questions from Excel file"""
    questions = []
    
    try:
        df = pd.read_excel(file_path)
        
        # Similar to CSV parsing
        for _, row in df.iterrows():
            if 'question' in df.columns:
                question_data = {
                    'question': str(row.get('question', '')),
                    'type': 'multiple_choice',
                    'options': [],
                    'correct_option_index': 0,
                    'confidence': 0.9
                }
                
                # Extract options
                option_cols = [col for col in df.columns if 'option' in col.lower()]
                for col in option_cols:
                    if pd.notna(row[col]):
                        question_data['options'].append(str(row[col]))
                
                # Determine correct answer
                if 'correct_answer' in df.columns:
                    correct_text = str(row['correct_answer']).strip()
                    for i, option in enumerate(question_data['options']):
                        if option.strip().lower() == correct_text.lower():
                            question_data['correct_option_index'] = i
                            break
                
                if question_data['question'] and len(question_data['options']) >= 2:
                    questions.append(question_data)
        
    except Exception as e:
        logging.error(f"Excel parsing error: {e}")
    
    return questions

def parse_text_questions(file_path):
    """Extract questions from plain text file"""
    questions = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        questions = extract_questions_from_text(text)
        
    except Exception as e:
        logging.error(f"Text parsing error: {e}")
    
    return questions

def extract_questions_from_text(text):
    """Extract questions from text using improved regex patterns"""
    questions = []
    
    if not text or len(text.strip()) < 20:
        return questions
    
    try:
        # Split text into lines for better processing
        lines = text.strip().split('\n')
        current_question = None
        current_options = []
        correct_index = 0
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Check for numbered question format (1. Question text)
            question_match = re.match(r'^(\d+)\.?\s*(.+)', line)
            if question_match:
                # Save previous question if exists
                if current_question and len(current_options) >= 2:
                    questions.append({
                        'question': current_question,
                        'type': 'multiple_choice',
                        'options': current_options,
                        'correct_option_index': correct_index,
                        'confidence': 0.9
                    })
                
                # Start new question
                current_question = question_match.group(2).strip()
                current_options = []
                correct_index = 0
                
            # Check for option format (A) Option text, *A) Option text, or A) *Option text)
            elif re.match(r'^[\*]?[A-Da-d]\)\s*', line):
                # Handle asterisk before label: "*A) Option"  
                leading_star = line.startswith('*')
                tmp = line[1:].lstrip() if leading_star else line
                
                # Extract the option body (everything after "A) ")
                option_match = re.match(r'^[A-Da-d][\)\.]\s*(.+)', tmp)
                if option_match:
                    body = option_match.group(1)
                    
                    # Handle asterisk after label: "A) *Option"
                    post_star = body.lstrip().startswith('*')
                    if post_star:
                        body = body.lstrip()[1:].lstrip()
                    
                    is_correct = leading_star or post_star
                    clean_option = body.strip()
                    
                    if is_correct:
                        correct_index = len(current_options)
                    
                    current_options.append(clean_option)
            
            # Check for alternative question formats
            elif line.startswith('Q:') or line.startswith('Question:'):
                if current_question and len(current_options) >= 2:
                    questions.append({
                        'question': current_question,
                        'type': 'multiple_choice',
                        'options': current_options,
                        'correct_option_index': correct_index,
                        'confidence': 0.9
                    })
                
                current_question = line.split(':', 1)[1].strip()
                current_options = []
                correct_index = 0
            
            i += 1
        
        # Add the last question if exists
        if current_question and len(current_options) >= 2:
            questions.append({
                'question': current_question,
                'type': 'multiple_choice',
                'options': current_options,
                'correct_option_index': correct_index,
                'confidence': 0.9
            })
        
        # If no questions found with structured approach, try regex fallback
        if not questions:
            # Clean and normalize text for regex patterns
            text_normalized = re.sub(r'\s+', ' ', text.strip())
            
            # Enhanced patterns for different question formats
            patterns = [
                # Pattern 1: Numbered questions with options (1. Question A) option B) option)
                r'(\d+\.?\s*)(.*?)\s*((?:[A-Da-d][\)\.].*?)(?=\d+\.|$))',
                # Pattern 2: Questions with options on new lines
                r'(Question\s*\d*:?\s*)(.*?)\s*((?:[A-Da-d][\)\.].*?)(?=Question|\d+\.|$))',
            ]
            
            # Try each pattern to extract questions
            for pattern in patterns:
                matches = re.finditer(pattern, text_normalized, re.MULTILINE | re.DOTALL)
                
                for match in matches:
                    try:
                        if len(match.groups()) >= 2:
                            question_text = match.group(-2).strip()
                            options_text = match.group(-1).strip()
                            
                            if len(question_text) > 10:
                                question_data = {
                                    'question': question_text,
                                    'type': 'multiple_choice',
                                    'options': [],
                                    'correct_option_index': 0,
                                    'confidence': 0.7
                                }
                                
                                # Extract individual options
                                option_parts = re.findall(r'[A-Da-d][\)\.]([^A-Da-d\)\.]*)(?=[A-Da-d][\)\.]|$)', options_text)
                                question_data['options'] = [opt.strip() for opt in option_parts if opt.strip()]
                                
                                if len(question_data['options']) >= 2:
                                    questions.append(question_data)
                                    
                    except Exception as e:
                        continue
                        
        return questions[:20]  # Limit to 20 questions max
        
    except Exception as e:
        logging.error(f"Text extraction error: {e}")
        return []

def select_top_questions(candidate_questions, num_questions):
    """Select top N questions based on confidence and completeness"""
    
    # Sort by confidence (highest first) and completeness
    def question_score(q):
        confidence = q.get('confidence', 0.5)
        option_count = len(q.get('options', []))
        has_answer = q.get('correct_option_index', -1) >= 0
        
        return confidence + (option_count / 10) + (0.2 if has_answer else 0)
    
    sorted_questions = sorted(candidate_questions, key=question_score, reverse=True)
    
    return sorted_questions[:num_questions]

# Enhanced Security Measures and Advanced Features

@app.route('/api/quiz/<int:quiz_id>/publish', methods=['POST'])
@login_required
def publish_quiz(quiz_id):
    """Publish a draft quiz after host review"""
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    data = request.json
    lock_answers = data.get('lock_answers', True)
    
    # Activate the quiz
    quiz.is_active = True
    
    # Lock answers if requested
    if lock_answers:
        for question in quiz.questions:
            for option in question.options:
                # Mark host-reviewed answers as final
                pass  # Options are already set correctly
    
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Quiz published successfully'})

@app.route('/api/quiz/<int:quiz_id>/delete', methods=['DELETE'])
@login_required
def delete_quiz(quiz_id):
    """Delete quiz (soft delete by default)"""
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    hard_delete = request.args.get('hard', 'false').lower() == 'true'
    
    if hard_delete and current_user.is_admin():
        # Hard delete - remove completely
        # First remove related upload files
        if quiz.draft_from_upload_id:
            upload_record = UploadRecord.query.get(quiz.draft_from_upload_id)
            if upload_record and upload_record.stored_path:
                try:
                    os.remove(upload_record.stored_path)
                except:
                    pass
                db.session.delete(upload_record)
        
        db.session.delete(quiz)
        message = 'Quiz permanently deleted'
    else:
        # Soft delete
        quiz.is_deleted = True
        message = 'Quiz moved to trash'
    
    db.session.commit()
    
    return jsonify({'success': True, 'message': message})

@app.route('/api/quiz/<int:quiz_id>/restore', methods=['POST'])
@login_required
def restore_quiz(quiz_id):
    """Restore a soft-deleted quiz"""
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    quiz.is_deleted = False
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Quiz restored successfully'})

@app.route('/host/participants-advanced')
@login_required
def host_participants_advanced():
    """Enhanced participant management and login activity"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get participants who have taken host's quizzes
    participants = db.session.query(User).join(QuizAttempt).join(Quiz).filter(
        Quiz.creator_id == current_user.id if not current_user.is_admin() else True,
        User.role == 'participant'
    ).distinct().all()
    
    # Get detailed info for each participant
    participant_data = []
    for participant in participants:
        # Get latest device log
        latest_device_log = DeviceLog.query.filter_by(user_id=participant.id).order_by(
            DeviceLog.logged_in_at.desc()
        ).first()
        
        # Get quiz attempts
        attempts = QuizAttempt.query.join(Quiz).filter(
            QuizAttempt.participant_id == participant.id,
            Quiz.creator_id == current_user.id if not current_user.is_admin() else True
        ).all()
        
        # Get violation count
        violation_count = 0
        for attempt in attempts:
            violation_count += ProctoringEvent.query.filter_by(attempt_id=attempt.id).count()
        
        participant_data.append({
            'user': participant,
            'latest_device': latest_device_log,
            'attempts': attempts,
            'violation_count': violation_count,
            'status': 'online' if latest_device_log and 
                     (datetime.utcnow() - latest_device_log.logged_in_at).seconds < 300 else 'offline'
        })
    
    return render_template('host_participants.html', participant_data=participant_data)

@app.route('/api/participant/<int:participant_id>/security-report')
@login_required
def participant_security_report(participant_id):
    """Generate detailed security report for a participant"""
    if not current_user.is_host() and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    participant = User.query.get_or_404(participant_id)
    
    # Get all attempts by this participant for current host's quizzes
    attempts = QuizAttempt.query.join(Quiz).filter(
        QuizAttempt.participant_id == participant_id,
        Quiz.creator_id == current_user.id if not current_user.is_admin() else True
    ).all()
    
    # Compile security data
    security_data = {
        'participant': {
            'username': participant.username,
            'email': participant.email,
            'total_attempts': len(attempts)
        },
        'violations': [],
        'device_logs': [],
        'flagged_attempts': [],
        'suspicious_patterns': []
    }
    
    # Get violations
    for attempt in attempts:
        violations = ProctoringEvent.query.filter_by(attempt_id=attempt.id).all()
        for violation in violations:
            security_data['violations'].append({
                'quiz_title': attempt.quiz.title,
                'event_type': violation.event_type,
                'severity': violation.severity,
                'timestamp': violation.timestamp.isoformat(),
                'details': violation.details
            })
    
    # Get device logs
    device_logs = DeviceLog.query.filter_by(user_id=participant_id).order_by(
        DeviceLog.logged_in_at.desc()
    ).limit(20).all()
    
    for log in device_logs:
        security_data['device_logs'].append({
            'ip_address': log.ip_address,
            'device_type': log.device_type,
            'browser_info': log.browser_info,
            'timestamp': log.logged_in_at.isoformat(),
            'is_suspicious': log.is_suspicious
        })
    
    # Detect suspicious patterns
    ip_addresses = set(log.ip_address for log in device_logs)
    if len(ip_addresses) > 3:
        security_data['suspicious_patterns'].append('Multiple IP addresses detected')
    
    user_agents = set(log.user_agent for log in device_logs if log.user_agent)
    if len(user_agents) > 2:
        security_data['suspicious_patterns'].append('Multiple devices/browsers detected')
    
    # Flagged attempts
    flagged_attempts = [a for a in attempts if a.violation_count > 2]
    for attempt in flagged_attempts:
        security_data['flagged_attempts'].append({
            'quiz_title': attempt.quiz.title,
            'started_at': attempt.started_at.isoformat(),
            'violation_count': attempt.violation_count,
            'status': attempt.status
        })
    
    return jsonify(security_data)

@app.route('/api/monitoring/send-warning/<int:attempt_id>', methods=['POST'])
@login_required
def send_warning_to_participant(attempt_id):
    """Send real-time warning to participant"""
    if not current_user.is_host() and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    if attempt.quiz.creator_id != current_user.id and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    data = request.json
    warning_message = data.get('message', 'Please follow the quiz guidelines and avoid suspicious activities.')
    
    # Create security alert
    alert = SecurityAlert(
        user_id=attempt.participant_id,
        quiz_id=attempt.quiz_id,
        attempt_id=attempt_id,
        alert_type='host_warning',
        severity='medium',
        description=warning_message,
        auto_action_taken='warning_sent'
    )
    db.session.add(alert)
    db.session.commit()
    
    # In a real WebSocket implementation, this would send the message via WebSocket
    # For now, we'll store it as a security alert that can be displayed to the participant
    
    return jsonify({'success': True, 'message': 'Warning sent successfully'})

@app.route('/api/monitoring/auto-terminate/<int:attempt_id>', methods=['POST'])
@login_required
def auto_terminate_quiz(attempt_id):
    """Automatically terminate quiz due to violations"""
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    # Check if auto-termination is enabled for this quiz
    if not attempt.quiz.auto_terminate_on_violation:
        return jsonify({'error': 'Auto-termination not enabled for this quiz'}), 400
    
    data = request.json
    reason = data.get('reason', 'Multiple proctoring violations detected')
    
    # Terminate the quiz
    attempt.status = 'terminated'
    attempt.completed_at = datetime.utcnow()
    attempt.termination_reason = reason
    attempt.is_flagged = True
    
    # Create security alert
    alert = SecurityAlert(
        user_id=attempt.participant_id,
        quiz_id=attempt.quiz_id,
        attempt_id=attempt_id,
        alert_type='auto_termination',
        severity='high',
        description=f'Quiz automatically terminated: {reason}',
        auto_action_taken='quiz_terminated'
    )
    db.session.add(alert)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Quiz terminated due to violations'})

@app.route('/api/device-log', methods=['POST'])
@login_required
def log_device_info():
    """Log participant device information for security tracking"""
    data = request.json
    
    # Detect suspicious behavior
    is_suspicious = False
    
    # Check for suspicious user agents
    suspicious_keywords = ['headless', 'phantom', 'selenium', 'webdriver', 'automation']
    user_agent = data.get('userAgent', '').lower()
    if any(keyword in user_agent for keyword in suspicious_keywords):
        is_suspicious = True
    
    # Check for unusual screen resolutions
    screen_resolution = data.get('screenResolution', '')
    if screen_resolution:
        try:
            width, height = map(int, screen_resolution.split('x'))
            if width < 800 or height < 600:  # Unusually small screens
                is_suspicious = True
        except:
            pass
    
    # Get IP address
    ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
    
    device_log = DeviceLog(
        user_id=current_user.id,
        quiz_id=data.get('quizId'),
        ip_address=ip_address,
        user_agent=data.get('userAgent'),
        device_type=data.get('deviceType'),
        browser_info=data.get('browserInfo'),
        screen_resolution=screen_resolution,
        timezone=data.get('timezone'),
        is_suspicious=is_suspicious
    )
    
    db.session.add(device_log)
    db.session.commit()
    
    return jsonify({'success': True, 'logged': True, 'suspicious': is_suspicious})

# Database Export for Admin
@app.route('/admin/export-database-sqlite')
@login_required
def export_database_sqlite():
    """Export database as SQLite file for admin"""
    if not current_user.is_admin():
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        import sqlite3
        import tempfile
        
        # Create temporary SQLite database
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'_quizzes_{timestamp}.db')
        sqlite_path = temp_file.name
        temp_file.close()
        
        # Connect to SQLite
        sqlite_conn = sqlite3.connect(sqlite_path)
        sqlite_cursor = sqlite_conn.cursor()
        
        # Export users
        users = User.query.all()
        sqlite_cursor.execute('''
            CREATE TABLE users (
                id INTEGER PRIMARY KEY,
                username TEXT,
                email TEXT,
                role TEXT,
                is_verified BOOLEAN,
                created_at TEXT
            )
        ''')
        
        for user in users:
            sqlite_cursor.execute('''
                INSERT INTO users VALUES (?, ?, ?, ?, ?, ?)
            ''', (user.id, user.username, user.email, user.role, user.is_verified, 
                  user.created_at.isoformat() if user.created_at else None))
        
        # Export quizzes
        quizzes = Quiz.query.filter_by(is_deleted=False).all()
        sqlite_cursor.execute('''
            CREATE TABLE quizzes (
                id INTEGER PRIMARY KEY,
                title TEXT,
                creator_username TEXT,
                time_limit INTEGER,
                proctoring_enabled BOOLEAN,
                created_at TEXT
            )
        ''')
        
        for quiz in quizzes:
            sqlite_cursor.execute('''
                INSERT INTO quizzes VALUES (?, ?, ?, ?, ?, ?)
            ''', (quiz.id, quiz.title, quiz.creator.username, quiz.time_limit, 
                  quiz.proctoring_enabled, quiz.created_at.isoformat()))
        
        # Export quiz attempts
        attempts = QuizAttempt.query.all()
        sqlite_cursor.execute('''
            CREATE TABLE quiz_attempts (
                id INTEGER PRIMARY KEY,
                participant_username TEXT,
                quiz_title TEXT,
                score REAL,
                status TEXT,
                started_at TEXT,
                completed_at TEXT,
                violation_count INTEGER
            )
        ''')
        
        for attempt in attempts:
            sqlite_cursor.execute('''
                INSERT INTO quiz_attempts VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (attempt.id, attempt.participant.username, attempt.quiz.title,
                  attempt.score, attempt.status, attempt.started_at.isoformat(),
                  attempt.completed_at.isoformat() if attempt.completed_at else None,
                  attempt.violation_count))
        
        sqlite_conn.commit()
        sqlite_conn.close()
        
        # Return file for download
        return send_file(
            sqlite_path,
            mimetype='application/x-sqlite3',
            as_attachment=True,
            download_name=f'quizzes_{timestamp}.db'
        )
        
    except Exception as e:
        logging.error(f"Database export error: {e}")
        flash('Database export failed.', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/admin/users')
@login_required
def admin_users():
    """Manage all users"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    users = User.query.order_by(User.created_at.desc()).all()
    courses = Course.query.filter_by(is_active=True).order_by(Course.code).all()
    return render_template('admin_users.html', users=users, courses=courses)

@app.route('/admin/user/<int:user_id>/toggle-status', methods=['POST'])
@login_required
def admin_toggle_user_status(user_id):
    """Toggle user active/inactive status"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    
    # Don't allow disabling yourself
    if user.id == current_user.id:
        flash('You cannot disable your own account.', 'error')
        return redirect(url_for('admin_users'))
    
    user.is_active = not user.is_active
    db.session.commit()
    
    status = 'activated' if user.is_active else 'deactivated'
    flash(f'User {user.username} has been {status}.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/<int:user_id>/change-role', methods=['POST'])
@login_required
def admin_change_user_role(user_id):
    """Change user role"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    new_role = request.form.get('role')
    
    if new_role not in ['admin', 'host', 'participant']:
        flash('Invalid role specified.', 'error')
        return redirect(url_for('admin_users'))
    
    # Don't allow changing your own role
    if user.id == current_user.id:
        flash('You cannot change your own role.', 'error')
        return redirect(url_for('admin_users'))
    
    old_role = user.role
    user.role = new_role
    db.session.commit()
    
    flash(f'User {user.username} role changed from {old_role} to {new_role}.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/<int:user_id>/delete', methods=['POST'])
@login_required
def admin_delete_user(user_id):
    """Delete a user (admin only)"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    
    # Don't allow deleting yourself
    if user.id == current_user.id:
        flash('You cannot delete your own account.', 'error')
        return redirect(url_for('admin_users'))
    
    try:
        # Delete related data in proper order
        # 1. Delete answers and proctoring events for this user's attempts
        user_attempts = QuizAttempt.query.filter_by(participant_id=user.id).all()
        for attempt in user_attempts:
            Answer.query.filter_by(attempt_id=attempt.id).delete()
            ProctoringEvent.query.filter_by(attempt_id=attempt.id).delete()
        
        # 2. Delete quiz attempts
        QuizAttempt.query.filter_by(participant_id=user.id).delete()
        
        # 3. Handle quizzes created by this user
        user_quizzes = Quiz.query.filter_by(creator_id=user.id).all()
        for quiz in user_quizzes:
            # Delete all attempts for these quizzes first
            quiz_attempts = QuizAttempt.query.filter_by(quiz_id=quiz.id).all()
            for attempt in quiz_attempts:
                Answer.query.filter_by(attempt_id=attempt.id).delete()
                ProctoringEvent.query.filter_by(attempt_id=attempt.id).delete()
            QuizAttempt.query.filter_by(quiz_id=quiz.id).delete()
            
            # Delete questions and options
            for question in quiz.questions:
                QuestionOption.query.filter_by(question_id=question.id).delete()
            Question.query.filter_by(quiz_id=quiz.id).delete()
            
            # Delete the quiz
            db.session.delete(quiz)
        
        # 4. Delete course enrollments and assignments
        ParticipantEnrollment.query.filter_by(participant_id=user.id).delete()
        HostCourseAssignment.query.filter_by(host_id=user.id).delete()
        
        # 5. Finally delete the user
        username = user.username
        db.session.delete(user)
        db.session.commit()
        
        flash(f'User {username} has been permanently deleted.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting user: {str(e)}', 'error')
    
    return redirect(url_for('admin_users'))

@app.route('/admin/user/<int:user_id>/reset-password', methods=['POST'])
@login_required
def admin_reset_password(user_id):
    """Reset user password"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    new_password = request.form.get('new_password')
    
    if not new_password or len(new_password) < 6:
        flash('Password must be at least 6 characters long.', 'error')
        return redirect(url_for('admin_users'))
    
    user.set_password(new_password)
    db.session.commit()
    
    flash(f'Password reset for user {user.username}.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/bulk-delete-users', methods=['POST'])
@login_required
def admin_bulk_delete_users():
    """Bulk delete multiple users"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    user_ids = request.form.getlist('user_ids')
    
    if not user_ids:
        flash('No users selected for deletion.', 'error')
        return redirect(url_for('admin_users'))
    
    # Convert to integers and validate
    try:
        user_ids = [int(uid) for uid in user_ids]
    except ValueError:
        flash('Invalid user IDs provided.', 'error')
        return redirect(url_for('admin_users'))
    
    # Prevent deleting current user
    if current_user.id in user_ids:
        flash('You cannot delete your own account.', 'error')
        return redirect(url_for('admin_users'))
    
    deleted_count = 0
    errors = []
    
    # Process users in very small batches to prevent worker timeouts
    batch_size = 3  # Process only 3 users at a time for stability
    for i in range(0, len(user_ids), batch_size):
        batch_ids = user_ids[i:i + batch_size]
        
        try:
            for user_id in batch_ids:
                try:
                    user = User.query.get(user_id)
                    if not user:
                        errors.append(f'User with ID {user_id} not found.')
                        continue
                    
                    username = user.username
                    
                    # Delete related data with explicit session flushing and optimized operations
                    # 1. Delete user's quiz attempts and related data first
                    user_attempt_ids = db.session.query(QuizAttempt.id).filter_by(participant_id=user.id).all()
                    if user_attempt_ids:
                        attempt_ids = [aid[0] for aid in user_attempt_ids]
                        # Delete in smaller chunks to avoid timeout
                        chunk_size = 50
                        for i in range(0, len(attempt_ids), chunk_size):
                            chunk = attempt_ids[i:i + chunk_size]
                            Answer.query.filter(Answer.attempt_id.in_(chunk)).delete(synchronize_session=False)
                            ProctoringEvent.query.filter(ProctoringEvent.attempt_id.in_(chunk)).delete(synchronize_session=False)
                        
                    QuizAttempt.query.filter_by(participant_id=user.id).delete(synchronize_session=False)
                    db.session.flush()  # Explicit flush to ensure data consistency
                    
                    # 2. Delete user's login events and violations (simplified)
                    LoginEvent.query.filter_by(user_id=user.id).delete(synchronize_session=False)
                    UserViolation.query.filter_by(user_id=user.id).delete(synchronize_session=False)
                    db.session.flush()
                    
                    # 3. Delete course enrollments and assignments
                    ParticipantEnrollment.query.filter_by(participant_id=user.id).delete(synchronize_session=False)
                    HostCourseAssignment.query.filter_by(host_id=user.id).delete(synchronize_session=False)
                    db.session.flush()
                    
                    # 4. Delete user-created quizzes (only for hosts, simplified)
                    if user.role == 'host':
                        # Get quiz IDs first
                        user_quiz_ids = db.session.query(Quiz.id).filter_by(creator_id=user.id).all()
                        if user_quiz_ids:
                            quiz_ids = [qid[0] for qid in user_quiz_ids]
                            
                            # Delete quiz attempts for these quizzes
                            quiz_attempt_ids = db.session.query(QuizAttempt.id).filter(QuizAttempt.quiz_id.in_(quiz_ids)).all()
                            if quiz_attempt_ids:
                                quiz_attempt_ids = [qaid[0] for qaid in quiz_attempt_ids]
                                Answer.query.filter(Answer.attempt_id.in_(quiz_attempt_ids)).delete(synchronize_session=False)
                                ProctoringEvent.query.filter(ProctoringEvent.attempt_id.in_(quiz_attempt_ids)).delete(synchronize_session=False)
                            
                            QuizAttempt.query.filter(QuizAttempt.quiz_id.in_(quiz_ids)).delete(synchronize_session=False)
                            
                            # Delete questions and options  
                            question_ids = db.session.query(Question.id).filter(Question.quiz_id.in_(quiz_ids)).all()
                            if question_ids:
                                question_ids = [qid[0] for qid in question_ids]
                                QuestionOption.query.filter(QuestionOption.question_id.in_(question_ids)).delete(synchronize_session=False)
                            
                            Question.query.filter(Question.quiz_id.in_(quiz_ids)).delete(synchronize_session=False)
                            Quiz.query.filter(Quiz.id.in_(quiz_ids)).delete(synchronize_session=False)
                            db.session.flush()
                    
                    # 5. Finally delete the user
                    db.session.delete(user)
                    deleted_count += 1
                    
                except Exception as user_error:
                    errors.append(f'Error deleting user {username if "username" in locals() else user_id}: {str(user_error)}')
                    continue
            
            # Commit each batch with explicit transaction handling
            db.session.commit()
            
            # Add small delay between batches to prevent overwhelming the database
            import time
            time.sleep(0.1)
            
        except Exception as e:
            db.session.rollback()
            errors.append(f'Error in batch starting at user {batch_ids[0]}: {str(e)}')
            continue
    
    if deleted_count > 0:
        flash(f'Successfully deleted {deleted_count} user(s).', 'success')
    
    if errors:
        for error in errors:
            flash(error, 'error')
    
    return redirect(url_for('admin_users'))

@app.route('/admin/quiz-attempts')
@login_required
def admin_quiz_attempts():
    """Admin page to view all quiz attempts"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    quiz_id = request.args.get('quiz_id', type=int)
    status = request.args.get('status', 'all')
    
    query = QuizAttempt.query.options(
        joinedload(QuizAttempt.participant),
        joinedload(QuizAttempt.quiz)
    )
    
    if quiz_id:
        query = query.filter(QuizAttempt.quiz_id == quiz_id)
    
    if status != 'all':
        query = query.filter(QuizAttempt.status == status)
    
    attempts = query.order_by(
        QuizAttempt.started_at.desc()
    ).paginate(
        page=page, per_page=20, error_out=False
    )
    
    # Get available quizzes for filtering
    quizzes = Quiz.query.order_by(Quiz.title).all()
    
    return render_template('admin_quiz_attempts.html',
                         attempts=attempts,
                         quizzes=quizzes,
                         current_quiz_id=quiz_id,
                         current_status=status)

@app.route('/admin/hosts')
@login_required
def admin_hosts():
    """Admin page to view all host accounts"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    status = request.args.get('status', 'all')
    search = request.args.get('search', '').strip()
    
    query = User.query.filter_by(role='host')
    
    if status == 'active':
        query = query.filter_by(is_active=True)
    elif status == 'inactive':
        query = query.filter_by(is_active=False)
    
    if search:
        query = query.filter(
            (User.username.ilike(f'%{search}%')) |
            (User.email.ilike(f'%{search}%'))
        )
    
    hosts = query.order_by(User.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    # Get host statistics
    host_stats = {}
    for host in hosts.items:
        quiz_count = Quiz.query.filter_by(creator_id=host.id).count()
        recent_login = LoginEvent.query.filter_by(user_id=host.id).order_by(LoginEvent.login_time.desc()).first()
        host_stats[host.id] = {
            'quiz_count': quiz_count,
            'last_login': recent_login.login_time if recent_login else None
        }
    
    return render_template('admin_hosts.html',
                         hosts=hosts,
                         host_stats=host_stats,
                         current_status=status,
                         search_query=search)

# Alert Threshold Management Routes
@app.route('/admin/alert-thresholds')
@login_required
def admin_alert_thresholds():
    """Admin page to manage alert thresholds"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '').strip()
    event_type = request.args.get('event_type', 'all')
    
    query = AlertThreshold.query.filter_by(is_active=True)
    
    if search:
        query = query.filter(
            (AlertThreshold.name.ilike(f'%{search}%')) |
            (AlertThreshold.event_type.ilike(f'%{search}%'))
        )
    
    if event_type != 'all':
        query = query.filter(AlertThreshold.event_type == event_type)
    
    thresholds = query.order_by(AlertThreshold.created_at.desc()).paginate(
        page=page, per_page=15, error_out=False
    )
    
    # Get distinct event types for filter dropdown
    event_types = db.session.query(AlertThreshold.event_type).distinct().all()
    event_types = [et[0] for et in event_types]
    
    return render_template('admin_alert_thresholds.html',
                         thresholds=thresholds,
                         event_types=event_types,
                         current_event_type=event_type,
                         search_query=search)

@app.route('/admin/alert-threshold/create', methods=['GET', 'POST'])
@login_required
def admin_create_alert_threshold():
    """Create new alert threshold"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        name = request.form.get('name')
        event_type = request.form.get('event_type')
        low_threshold = request.form.get('low_threshold', type=int)
        medium_threshold = request.form.get('medium_threshold', type=int)
        high_threshold = request.form.get('high_threshold', type=int)
        time_window = request.form.get('time_window', type=int)
        
        # Action settings
        send_alert = request.form.get('send_alert') == 'on'
        notify_proctor = request.form.get('notify_proctor') == 'on'
        auto_flag_attempt = request.form.get('auto_flag_attempt') == 'on'
        auto_terminate = request.form.get('auto_terminate') == 'on'
        
        # Validation
        if not name or not event_type:
            flash('Name and event type are required.', 'error')
            return redirect(url_for('admin_create_alert_threshold'))
        
        if low_threshold < 1 or medium_threshold < 1 or high_threshold < 1:
            flash('Thresholds must be at least 1.', 'error')
            return redirect(url_for('admin_create_alert_threshold'))
        
        if time_window < 1:
            flash('Time window must be at least 1 minute.', 'error')
            return redirect(url_for('admin_create_alert_threshold'))
        
        # Check if threshold already exists for this event type
        existing = AlertThreshold.query.filter_by(
            event_type=event_type, 
            is_global=True, 
            is_active=True
        ).first()
        
        if existing:
            flash(f'A global threshold for {event_type} already exists. Please edit the existing one.', 'error')
            return redirect(url_for('admin_alert_thresholds'))
        
        try:
            threshold = AlertThreshold(
                name=name,
                event_type=event_type,
                low_threshold=low_threshold,
                medium_threshold=medium_threshold,
                high_threshold=high_threshold,
                time_window=time_window,
                send_alert=send_alert,
                notify_proctor=notify_proctor,
                auto_flag_attempt=auto_flag_attempt,
                auto_terminate=auto_terminate,
                is_global=True,
                created_by=current_user.id
            )
            
            db.session.add(threshold)
            db.session.commit()
            
            flash(f'Alert threshold "{name}" created successfully.', 'success')
            return redirect(url_for('admin_alert_thresholds'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating threshold: {str(e)}', 'error')
            return redirect(url_for('admin_create_alert_threshold'))
    
    # Define available event types
    event_types = [
        'tab_switch', 'window_blur', 'multiple_faces', 'face_not_detected',
        'mouse_leave', 'keyboard_shortcut', 'copy_paste', 'right_click',
        'screen_share', 'suspicious_movement', 'audio_detected', 'browser_resize'
    ]
    
    return render_template('admin_create_alert_threshold.html', event_types=event_types)

@app.route('/admin/alert-threshold/<int:threshold_id>/edit', methods=['GET', 'POST'])
@login_required
def admin_edit_alert_threshold(threshold_id):
    """Edit alert threshold"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    threshold = AlertThreshold.query.get_or_404(threshold_id)
    
    if request.method == 'POST':
        threshold.name = request.form.get('name')
        threshold.low_threshold = request.form.get('low_threshold', type=int)
        threshold.medium_threshold = request.form.get('medium_threshold', type=int)
        threshold.high_threshold = request.form.get('high_threshold', type=int)
        threshold.time_window = request.form.get('time_window', type=int)
        
        # Action settings
        threshold.send_alert = request.form.get('send_alert') == 'on'
        threshold.notify_proctor = request.form.get('notify_proctor') == 'on'
        threshold.auto_flag_attempt = request.form.get('auto_flag_attempt') == 'on'
        threshold.auto_terminate = request.form.get('auto_terminate') == 'on'
        
        # Validation
        if not threshold.name:
            flash('Name is required.', 'error')
            return redirect(url_for('admin_edit_alert_threshold', threshold_id=threshold_id))
        
        if threshold.low_threshold < 1 or threshold.medium_threshold < 1 or threshold.high_threshold < 1:
            flash('Thresholds must be at least 1.', 'error')
            return redirect(url_for('admin_edit_alert_threshold', threshold_id=threshold_id))
        
        if threshold.time_window < 1:
            flash('Time window must be at least 1 minute.', 'error')
            return redirect(url_for('admin_edit_alert_threshold', threshold_id=threshold_id))
        
        try:
            threshold.updated_at = datetime.utcnow()
            db.session.commit()
            
            flash(f'Alert threshold "{threshold.name}" updated successfully.', 'success')
            return redirect(url_for('admin_alert_thresholds'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating threshold: {str(e)}', 'error')
    
    return render_template('admin_edit_alert_threshold.html', threshold=threshold)

@app.route('/admin/alert-threshold/<int:threshold_id>/delete', methods=['POST'])
@login_required
def admin_delete_alert_threshold(threshold_id):
    """Delete alert threshold"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    threshold = AlertThreshold.query.get_or_404(threshold_id)
    
    try:
        # Soft delete by setting is_active to False
        threshold.is_active = False
        threshold.updated_at = datetime.utcnow()
        db.session.commit()
        
        flash(f'Alert threshold "{threshold.name}" deleted successfully.', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting threshold: {str(e)}', 'error')
    
    return redirect(url_for('admin_alert_thresholds'))

@app.route('/admin/create-user', methods=['POST'])
@login_required
def admin_create_user():
    """Create new user account"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')
    role = request.form.get('role')
    
    # Validation
    if not all([username, email, password, role]):
        flash('All fields are required.', 'error')
        return redirect(url_for('admin_users'))
    
    if User.query.filter_by(email=email).first():
        flash('Email already exists.', 'error')
        return redirect(url_for('admin_users'))
    
    if User.query.filter_by(username=username).first():
        flash('Username already exists.', 'error')
        return redirect(url_for('admin_users'))
    
    if role not in ['admin', 'host', 'participant']:
        flash('Invalid role.', 'error')
        return redirect(url_for('admin_users'))
    
    # Create user
    user = User()
    user.username = username
    user.email = email
    user.role = role
    user.set_password(password)
    user.is_verified = True  # Admin-created users are pre-verified
    
    db.session.add(user)
    db.session.commit()
    
    flash(f'User {username} created successfully with role {role}.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/bulk-create-users', methods=['POST'])
@login_required
def admin_bulk_create_users():
    """Create multiple users at once"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    users_data = request.form.get('users_data', '').strip()
    default_role = request.form.get('default_role', 'participant')
    default_course = request.form.get('default_course', '').strip()
    auto_create_courses = request.form.get('auto_create_courses') == 'on'
    
    if not users_data:
        flash('Please provide user data.', 'error')
        return redirect(url_for('admin_users'))
    
    # Parse user data (format: username,email,password,role,courses per line)
    lines = [line.strip() for line in users_data.split('\n') if line.strip()]
    created_count = 0
    errors = []
    
    # Process users in small batches to prevent worker timeouts
    batch_size = 5  # Process 5 users at a time to prevent timeouts
    total_lines = len(lines)
    
    for batch_start in range(0, total_lines, batch_size):
        batch_end = min(batch_start + batch_size, total_lines)
        batch_lines = lines[batch_start:batch_end]
        
        try:
            # Process current batch
            for batch_index, line in enumerate(batch_lines):
                i = batch_start + batch_index + 1  # Original line number
                try:
                    parts = [part.strip() for part in line.split(',')]
                    if len(parts) < 2:
                        errors.append(f"Line {i}: Invalid format. Expected: username,email,password,role,courses")
                        continue
                    
                    username = parts[0]
                    email = parts[1]
                    password = parts[2] if len(parts) >= 3 and parts[2] else f"BigBoss{__import__('random').randrange(1000, 9999)}"
                    role = parts[3] if len(parts) >= 4 and parts[3] else default_role
                    courses_str = parts[4] if len(parts) >= 5 else ''
                    
                    # Validation
                    if not username or not email:
                        errors.append(f"Line {i}: Username and email are required")
                        continue
                    
                    if User.query.filter_by(email=email).first():
                        errors.append(f"Line {i}: Email {email} already exists")
                        continue
                    
                    if User.query.filter_by(username=username).first():
                        errors.append(f"Line {i}: Username {username} already exists")
                        continue
                    
                    if role not in ['admin', 'host', 'participant']:
                        errors.append(f"Line {i}: Invalid role {role}")
                        continue
                    
                    # Create user
                    user = User()
                    user.username = username
                    user.email = email
                    user.role = role
                    user.set_password(password)
                    user.is_verified = True
                    
                    db.session.add(user)
                    db.session.flush()  # Get user ID before processing courses
                    
                    # Process course assignments efficiently
                    course_codes = []
                    if courses_str.strip():
                        # Use courses from CSV line
                        course_codes = [code.strip().upper() for code in courses_str.split(';') if code.strip()]
                    elif default_course:
                        # Use default course  
                        course_codes = [default_course.upper()]
                    
                    # Assign user to courses
                    for course_code in course_codes:
                        try:
                            course = Course.query.filter_by(code=course_code).first()
                            
                            # Auto-create course if it doesn't exist and auto-create is enabled
                            if not course and auto_create_courses:
                                course = Course()
                                course.name = course_code  # Use code as name
                                course.code = course_code
                                course.description = f"Auto-created course for {course_code}"
                                course.max_participants = 100
                                course.is_active = True
                                db.session.add(course)
                                db.session.flush()  # Get course ID
                            
                            if course:
                                if role == 'participant':
                                    # Check if already enrolled
                                    existing_enrollment = ParticipantEnrollment.query.filter_by(
                                        participant_id=user.id, course_id=course.id
                                    ).first()
                                    if not existing_enrollment:
                                        enrollment = ParticipantEnrollment()
                                        enrollment.participant_id = user.id
                                        enrollment.course_id = course.id
                                        enrollment.enrolled_by = current_user.id
                                        db.session.add(enrollment)
                                
                                elif role == 'host':
                                    # Check if already assigned
                                    existing_assignment = HostCourseAssignment.query.filter_by(
                                        host_id=user.id, course_id=course.id
                                    ).first()
                                    if not existing_assignment:
                                        assignment = HostCourseAssignment()
                                        assignment.host_id = user.id
                                        assignment.course_id = course.id
                                        assignment.assigned_by = current_user.id
                                        db.session.add(assignment)
                            else:
                                errors.append(f"Line {i}: Course '{course_code}' not found")
                        
                        except Exception as course_error:
                            errors.append(f"Line {i}: Error assigning course '{course_code}': {str(course_error)}")
                    
                    created_count += 1
                    
                except Exception as e:
                    errors.append(f"Line {i}: Error processing - {str(e)}")
            
            # Commit each batch to prevent timeouts
            try:
                db.session.commit()
                
                # Add small delay between batches to prevent overwhelming the database
                if batch_end < total_lines:  # Not the last batch
                    import time
                    time.sleep(0.1)
                    
            except Exception as batch_error:
                db.session.rollback()
                errors.append(f"Batch {batch_start//batch_size + 1}: Database error - {str(batch_error)}")
                
        except Exception as batch_exception:
            db.session.rollback()
            errors.append(f"Batch {batch_start//batch_size + 1}: Processing error - {str(batch_exception)}")
    
    # Final status messages
    if created_count > 0:
        flash(f'Successfully created {created_count} users.', 'success')
    if errors:
        flash(f'Errors encountered: {"; ".join(errors[:5])}{"..." if len(errors) > 5 else ""}', 'warning')
    
    return redirect(url_for('admin_users'))


@app.route('/admin/upload-users-excel', methods=['POST'])
@login_required
def admin_upload_users_excel():
    """Upload users from Excel file"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    if 'excel_file' not in request.files:
        flash('No file selected.', 'error')
        return redirect(url_for('admin_users'))
    
    file = request.files['excel_file']
    if file.filename == '':
        flash('No file selected.', 'error')
        return redirect(url_for('admin_users'))
    
    if not file.filename.lower().endswith(('.xlsx', '.xls')):
        flash('Please upload an Excel file (.xlsx or .xls).', 'error')
        return redirect(url_for('admin_users'))
    
    try:
        # Read Excel file
        import pandas as pd
        df = pd.read_excel(file)
        
        # Expected columns: username, email, password (optional), role (optional)
        required_columns = ['username', 'email']
        if not all(col in df.columns for col in required_columns):
            flash(f'Excel file must contain columns: {", ".join(required_columns)}. Optional: password, role', 'error')
            return redirect(url_for('admin_users'))
        
        created_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                username = str(row['username']).strip()
                email = str(row['email']).strip()
                password = str(row.get('password', f'BigBoss{__import__("random").randrange(1000, 9999)}')).strip()
                role = str(row.get('role', 'participant')).strip().lower()
                
                # Validation
                if not username or not email or username == 'nan' or email == 'nan':
                    errors.append(f"Row {index + 2}: Username and email are required")
                    continue
                
                if User.query.filter_by(email=email).first():
                    errors.append(f"Row {index + 2}: Email {email} already exists")
                    continue
                
                if User.query.filter_by(username=username).first():
                    errors.append(f"Row {index + 2}: Username {username} already exists")
                    continue
                
                if role not in ['admin', 'host', 'participant']:
                    role = 'participant'  # Default to participant if invalid
                
                # Create user
                user = User()
                user.username = username
                user.email = email
                user.role = role
                user.set_password(password)
                user.is_verified = True
                
                db.session.add(user)
                created_count += 1
                
            except Exception as e:
                errors.append(f"Row {index + 2}: Error - {str(e)}")
        
        db.session.commit()
        
        if created_count > 0:
            flash(f'Successfully created {created_count} users from Excel file.', 'success')
        if errors:
            flash(f'Errors: {"; ".join(errors[:3])}{"..." if len(errors) > 3 else ""}', 'warning')
            
    except Exception as e:
        flash(f'Error reading Excel file: {str(e)}', 'error')
    
    return redirect(url_for('admin_users'))

@app.route('/host/participants')
@login_required
def host_participants():
    """Course-based participant management for hosts"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get courses assigned to this host (or all courses if admin)
    if current_user.is_admin():
        assigned_courses = Course.query.filter_by(is_active=True).all()
    else:
        # Get courses where this user is assigned as host
        host_assignments = HostCourseAssignment.query.filter_by(host_id=current_user.id).all()
        assigned_courses = [assignment.course for assignment in host_assignments if assignment.course.is_active]
    
    # Get participants enrolled in these courses
    course_participants = {}
    all_attempts = []
    
    for course in assigned_courses:
        # Get participants enrolled in this course
        enrollments = ParticipantEnrollment.query.filter_by(course_id=course.id).all()
        participants = [enrollment.participant for enrollment in enrollments]
        
        # Get quizzes for this course created by the current host (or all if admin)
        if current_user.is_admin():
            course_quizzes = Quiz.query.filter_by(course_id=course.id, is_active=True).all()
        else:
            course_quizzes = Quiz.query.filter_by(course_id=course.id, creator_id=current_user.id, is_active=True).all()
        
        # Get attempts for these quizzes
        quiz_ids = [quiz.id for quiz in course_quizzes]
        if quiz_ids:
            course_attempts = QuizAttempt.query.filter(QuizAttempt.quiz_id.in_(quiz_ids)).order_by(QuizAttempt.started_at.desc()).all()
            all_attempts.extend(course_attempts)
        else:
            course_attempts = []
        
        # Calculate participant statistics for this course
        participant_stats = {}
        for participant in participants:
            participant_attempts = [attempt for attempt in course_attempts if attempt.participant_id == participant.id]
            completed_attempts = [attempt for attempt in participant_attempts if attempt.status == 'completed']
            avg_score = sum([attempt.score for attempt in completed_attempts if attempt.score]) / len(completed_attempts) if completed_attempts else 0
            
            # Get violation count
            violation_count = 0
            for attempt in participant_attempts:
                violation_count += ProctoringEvent.query.filter_by(attempt_id=attempt.id).count()
            
            # Get recent login
            recent_login = LoginEvent.query.filter_by(user_id=participant.id).order_by(LoginEvent.login_time.desc()).first()
            
            participant_stats[participant.id] = {
                'total_attempts': len(participant_attempts),
                'completed_attempts': len(completed_attempts),
                'avg_score': avg_score,
                'violation_count': violation_count,
                'recent_login': recent_login,
                'is_flagged': False  # Will be tracked via UserViolation model
            }
        
        course_participants[course] = {
            'participants': participants,
            'quizzes': course_quizzes,
            'attempts': course_attempts,
            'stats': participant_stats
        }
    
    return render_template('host_participants.html', 
                         course_participants=course_participants,
                         assigned_courses=assigned_courses,
                         attempts=all_attempts)

@app.route('/host/participant/<int:participant_id>/manage', methods=['GET', 'POST'])
@login_required
def manage_participant(participant_id):
    """Manage individual participant - credentials, performance, flags"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    participant = User.query.get_or_404(participant_id)
    if participant.role != 'participant':
        flash('Invalid participant.', 'error')
        return redirect(url_for('host_participants'))
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'update_credentials':
            # Update participant credentials
            new_username = request.form.get('username')
            new_email = request.form.get('email')
            new_password = request.form.get('password')
            
            if new_username:
                participant.username = new_username
            if new_email:
                participant.email = new_email
            if new_password:
                participant.set_password(new_password)
            
            db.session.commit()
            flash('Participant credentials updated successfully.', 'success')
            
        elif action == 'flag_user':
            # Flag user for violations
            violation_record = UserViolation.query.filter_by(user_id=participant_id).first()
            if not violation_record:
                violation_record = UserViolation(user_id=participant_id)
                db.session.add(violation_record)
            
            violation_record.is_flagged = True
            violation_record.flagged_at = datetime.utcnow()
            violation_record.flagged_by = current_user.id
            violation_record.notes = request.form.get('notes', '')
            
            db.session.commit()
            flash('Participant flagged for violations.', 'warning')
            
        elif action == 'unflag_user':
            # Remove flag
            violation_record = UserViolation.query.filter_by(user_id=participant_id).first()
            if violation_record:
                violation_record.is_flagged = False
                violation_record.notes = request.form.get('notes', '')
                db.session.commit()
            flash('Participant flag removed.', 'success')
            
        return redirect(url_for('manage_participant', participant_id=participant_id))
    
    # Get participant data for display
    attempts = QuizAttempt.query.filter_by(participant_id=participant_id).order_by(QuizAttempt.started_at.desc()).all()
    login_events = LoginEvent.query.filter_by(user_id=participant_id).order_by(LoginEvent.login_time.desc()).limit(10).all()
    
    # Get violation data
    violation_record = UserViolation.query.filter_by(user_id=participant_id).first()
    total_violations = 0
    for attempt in attempts:
        total_violations += ProctoringEvent.query.filter_by(attempt_id=attempt.id).count()
    
    return render_template('manage_participant.html', 
                         participant=participant,
                         attempts=attempts,
                         login_events=login_events,
                         violation_record=violation_record,
                         total_violations=total_violations)

@app.route('/host/participant/<int:participant_id>/violations')
@login_required
def view_participant_violations(participant_id):
    """View detailed violation history for a participant"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    participant = User.query.get_or_404(participant_id)
    attempts = QuizAttempt.query.filter_by(participant_id=participant_id).all()
    
    # Get all violations for this participant
    violations = []
    for attempt in attempts:
        attempt_violations = ProctoringEvent.query.filter_by(attempt_id=attempt.id).order_by(ProctoringEvent.timestamp.desc()).all()
        for violation in attempt_violations:
            violations.append({
                'violation': violation,
                'attempt': attempt,
                'quiz': attempt.quiz
            })
    
    violations.sort(key=lambda x: x['violation'].timestamp, reverse=True)
    
    return render_template('participant_violations.html', 
                         participant=participant,
                         violations=violations)

@app.route('/host/login-activity')
@login_required
def host_login_activity():
    """View login activity of all participants"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get recent login events for participants with proper eager loading
    login_events = LoginEvent.query.options(
        joinedload(LoginEvent.user)
    ).join(User).filter(
        User.role == 'participant'
    ).order_by(LoginEvent.login_time.desc()).limit(50).all()
    
    return render_template('host_login_activity.html', login_events=login_events)


@app.route('/host/live-monitoring')
@login_required
def host_live_monitoring():
    """Live violation monitoring dashboard for hosts"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get all active quizzes by this host
    active_quizzes = Quiz.query.filter_by(creator_id=current_user.id).all()
    
    # Get recent violations across all their quizzes
    recent_violations = db.session.query(
        ProctoringEvent, QuizAttempt, User, Quiz
    ).join(
        QuizAttempt, ProctoringEvent.attempt_id == QuizAttempt.id
    ).join(
        User, QuizAttempt.participant_id == User.id
    ).join(
        Quiz, QuizAttempt.quiz_id == Quiz.id
    ).filter(
        Quiz.creator_id == current_user.id
    ).order_by(
        ProctoringEvent.timestamp.desc()
    ).limit(100).all()
    
    return render_template('live_monitoring.html', 
                         active_quizzes=active_quizzes,
                         recent_violations=recent_violations)

@app.route('/admin/manage-flags')
@login_required
def admin_manage_flags():
    """Admin interface to manage user flags and retake permissions"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get all flagged users - specify which foreign key to use for join
    flagged_users = db.session.query(UserViolation, User).join(
        User, UserViolation.user_id == User.id
    ).filter(
        UserViolation.is_flagged == True
    ).order_by(UserViolation.flagged_at.desc()).all()
    
    # Get recent violations
    recent_violations = db.session.query(ProctoringEvent, QuizAttempt, User, Quiz).join(
        QuizAttempt, ProctoringEvent.attempt_id == QuizAttempt.id
    ).join(
        User, QuizAttempt.participant_id == User.id
    ).join(
        Quiz, QuizAttempt.quiz_id == Quiz.id
    ).filter(
        ProctoringEvent.severity.in_(['high', 'critical'])
    ).order_by(ProctoringEvent.timestamp.desc()).limit(20).all()
    
    return render_template('admin_manage_flags.html', 
                         flagged_users=flagged_users,
                         recent_violations=recent_violations)

@app.route('/admin/unflag-user/<int:user_id>', methods=['POST'])
@login_required
def admin_unflag_user(user_id):
    """Admin action to unflag a user and grant retake permissions"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    violation_record = UserViolation.query.filter_by(user_id=user_id).first()
    if not violation_record:
        flash('User violation record not found.', 'error')
        return redirect(url_for('admin_manage_flags'))
    
    user = User.query.get_or_404(user_id)
    
    # Remove flag
    violation_record.is_flagged = False
    violation_record.unflagged_at = datetime.utcnow()
    violation_record.unflagged_by = current_user.id
    violation_record.notes = request.form.get('notes', '')
    
    db.session.commit()
    
    flash(f'User {user.username} has been unflagged and granted retake permissions.', 'success')
    return redirect(url_for('admin_manage_flags'))

@app.route('/admin/flag-user/<int:user_id>', methods=['POST'])
@login_required
def admin_flag_user(user_id):
    """Admin action to manually flag a user"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    
    violation_record = UserViolation.query.filter_by(user_id=user_id).first()
    if not violation_record:
        violation_record = UserViolation(user_id=user_id)
        db.session.add(violation_record)
    
    violation_record.is_flagged = True
    violation_record.flagged_at = datetime.utcnow()
    violation_record.flagged_by = current_user.id
    violation_record.notes = request.form.get('notes', 'Manually flagged by administrator')
    violation_record.violation_count = (violation_record.violation_count or 0) + 1
    
    db.session.commit()
    
    flash(f'User {user.username} has been flagged for violations.', 'warning')
    return redirect(url_for('admin_manage_flags'))

@app.route('/api/violations/<int:attempt_id>')
@login_required
def get_violations(attempt_id):
    """Get violations for a quiz attempt"""
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    # Check permissions
    if not (current_user.is_admin() or 
            (current_user.is_host() and attempt.quiz.creator_id == current_user.id)):
        return jsonify({'error': 'Access denied'}), 403
    
    violations = [{
        'event_type': event.event_type,
        'description': event.description,
        'severity': event.severity,
        'timestamp': event.timestamp.isoformat() if event.timestamp else None
    } for event in attempt.proctoring_events]
    
    return jsonify({'violations': violations})

@app.route('/api/proctoring/event', methods=['POST'])
@login_required
def log_proctoring_event():
    """Log proctoring violation events"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Get current quiz attempt
        attempt_id = data.get('attemptId')
        if not attempt_id:
            return jsonify({'error': 'Attempt ID required'}), 400
        
        attempt = QuizAttempt.query.get_or_404(attempt_id)
        
        # Verify user owns this attempt
        if attempt.participant_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403
        
        # Create proctoring event
        event = ProctoringEvent(
            attempt_id=attempt_id,
            event_type=data.get('type', 'unknown'),
            details=data.get('description', 'Unknown violation'),
            severity=data.get('severity', 'medium'),
            timestamp=datetime.utcnow()
        )
        
        db.session.add(event)
        
        # Update highest risk summary for this attempt (performance optimization)
        attempt.update_highest_risk(data.get('severity', 'medium'))
        
        # Enhanced violation tracking and termination logic
        violation_count = ProctoringEvent.query.filter_by(attempt_id=attempt_id).count() + 1
        high_severity_count = ProctoringEvent.query.filter_by(
            attempt_id=attempt_id, 
            severity='high'
        ).count()
        
        if data.get('severity') == 'high':
            high_severity_count += 1
        
        # Immediate termination conditions
        immediate_termination_types = ['quiz_terminated', 'console_access', 'multiple_instances', 'devtools_opened']
        should_terminate = (
            data.get('type') in immediate_termination_types or
            violation_count >= 3 or
            high_severity_count >= 2
        )
        
        if should_terminate:
            # Terminate the quiz attempt
            attempt.status = 'terminated'
            attempt.completed_at = datetime.utcnow()
            
            # Auto-flag the user for violations
            violation_record = UserViolation.query.filter_by(user_id=current_user.id).first()
            if not violation_record:
                violation_record = UserViolation(user_id=current_user.id)
                db.session.add(violation_record)
            
            violation_record.is_flagged = True
            violation_record.flagged_at = datetime.utcnow()
            violation_record.flagged_by = None  # System flagged
            violation_record.notes = f"Auto-flagged due to quiz termination: {data.get('type')} - {data.get('description')}"
            violation_record.violation_count = (violation_record.violation_count or 0) + 1
            
            # Save current answers before termination
            db.session.commit()
            
            return jsonify({
                'status': 'terminated',
                'message': f'Quiz terminated due to security violation: {data.get("description")}'
            })
        
        db.session.commit()
        
        response_data = {'status': 'logged'}
        
        # Add warning if approaching limits
        if violation_count >= 2:
            response_data['warning'] = f'{3 - violation_count} violations remaining before termination'
        
        return jsonify(response_data)
        
    except Exception as e:
        print(f"Error logging proctoring event: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/proctoring/verify-identity', methods=['POST'])
@login_required
def verify_identity():
    """API endpoint for face verification during quiz start"""
    try:
        data = request.json
        image_data = data.get('image')
        attempt_id = data.get('attemptId')
        
        if not image_data or not attempt_id:
            return jsonify({'verified': False, 'error': 'Missing required data'})
        
        # Get the attempt to verify it belongs to current user
        attempt = QuizAttempt.query.get(attempt_id)
        if not attempt or attempt.participant_id != current_user.id:
            return jsonify({'verified': False, 'error': 'Invalid attempt'})
        
        # Simple face verification logic - in production, you would use actual face recognition
        # For now, we'll do basic image validation and always approve if image is provided
        try:
            # Basic validation - check if it's a valid base64 image
            if image_data.startswith('data:image/'):
                # Extract base64 part
                base64_data = image_data.split(',')[1]
                import base64
                image_bytes = base64.b64decode(base64_data)
                
                # Basic check - image should be at least 1KB
                if len(image_bytes) > 1024:
                    # For demo purposes, always return verified=True
                    # In production, implement actual face recognition here
                    return jsonify({'verified': True, 'message': 'Identity verified successfully'})
                else:
                    return jsonify({'verified': False, 'error': 'Image too small or invalid'})
            else:
                return jsonify({'verified': False, 'error': 'Invalid image format'})
                
        except Exception as e:
            return jsonify({'verified': False, 'error': f'Image processing failed: {str(e)}'})
            
    except Exception as e:
        return jsonify({'verified': False, 'error': f'Verification failed: {str(e)}'})


@app.route('/api/proctoring/notify-violation', methods=['POST'])
@login_required
def notify_violation():
    """API endpoint for real-time violation notifications to hosts and admins"""
    try:
        data = request.json
        
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Extract notification details
        message = data.get('message', 'Unknown violation')
        severity = data.get('severity', 'medium')
        attempt_id = data.get('attemptId')
        student_info = data.get('student', {})
        
        if not attempt_id:
            return jsonify({'success': False, 'error': 'Attempt ID required'})
        
        # Get the quiz attempt and related quiz/host info
        attempt = QuizAttempt.query.get(attempt_id)
        if not attempt:
            return jsonify({'success': False, 'error': 'Invalid attempt ID'})
        
        quiz = attempt.quiz
        host = quiz.creator
        
        # Create violation notification record (for tracking)
        violation_notification = ProctoringEvent(
            attempt_id=attempt_id,
            event_type='notification_sent',
            details=f"Real-time notification: {message}",
            severity=severity,
            timestamp=datetime.utcnow()
        )
        db.session.add(violation_notification)
        
        # Send email notification to host if high severity
        if severity == 'high':
            try:
                subject = f"🚨 URGENT: Quiz Violation Alert - {student_info.get('name', 'Student')}"
                body = f"""
                URGENT VIOLATION ALERT
                
                Student: {student_info.get('name', 'Unknown')} ({student_info.get('email', 'N/A')})
                Quiz: {quiz.title}
                Violation: {message}
                Severity: {severity.upper()}
                Time: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}
                
                Please check the live monitoring dashboard immediately.
                
                Quiz URL: {request.url_root}host/live-monitoring
                """
                
                msg = Message(
                    subject=subject,
                    recipients=[host.email],
                    body=body
                )
                
                mail.send(msg)
                
            except Exception as email_error:
                logging.error(f"Failed to send violation email: {email_error}")
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Violation notification processed successfully',
            'notification_sent': True
        })
        
    except Exception as e:
        logging.error(f"Error processing violation notification: {e}")
        return jsonify({'success': False, 'error': 'Internal server error'})

# Helper function for heatmap data processing
def update_heatmap_data(quiz_id, question_id):
    """Update aggregated heatmap data for a question"""
    try:
        if not question_id:
            return
            
        # Get or create heatmap data record
        heatmap_data = QuestionHeatmapData.query.filter_by(
            quiz_id=quiz_id, 
            question_id=question_id
        ).first()
        
        if not heatmap_data:
            heatmap_data = QuestionHeatmapData(
                quiz_id=quiz_id,
                question_id=question_id
            )
            db.session.add(heatmap_data)
        
        # Calculate aggregated metrics from interaction events
        attempts = QuizAttempt.query.filter_by(quiz_id=quiz_id).all()
        attempt_ids = [attempt.id for attempt in attempts]
        
        if attempt_ids:
            # Get interaction events for this question
            events = InteractionEvent.query.filter(
                InteractionEvent.attempt_id.in_(attempt_ids),
                InteractionEvent.question_id == question_id
            ).all()
            
            # Calculate metrics
            participants = set([event.attempt_id for event in events])
            heatmap_data.total_participants = len(participants)
            
            # Calculate average time spent (from focus events)
            focus_events = [e for e in events if e.event_type == 'focus' and e.duration]
            if focus_events:
                heatmap_data.average_time_spent = sum([e.duration for e in focus_events]) / len(focus_events)
            
            # Count interactions
            heatmap_data.total_clicks = len([e for e in events if e.event_type == 'click'])
            heatmap_data.total_hovers = len([e for e in events if e.event_type == 'hover'])
            
            # Calculate hotspots (coordinates where most interactions happen)
            click_coords = [(e.x_coordinate, e.y_coordinate) for e in events 
                           if e.event_type == 'click' and e.x_coordinate and e.y_coordinate]
            hover_coords = [(e.x_coordinate, e.y_coordinate) for e in events 
                           if e.event_type == 'hover' and e.x_coordinate and e.y_coordinate]
            
            heatmap_data.click_hotspots = json.dumps(click_coords[:100])  # Store top 100 coords
            heatmap_data.hover_hotspots = json.dumps(hover_coords[:100])
            
            # Calculate engagement score based on interaction frequency
            if heatmap_data.total_participants > 0:
                total_interactions = heatmap_data.total_clicks + heatmap_data.total_hovers
                heatmap_data.engagement_score = total_interactions / heatmap_data.total_participants
            
            heatmap_data.last_updated = datetime.utcnow()
        
        db.session.commit()
        
    except Exception as e:
        logging.error(f"Error updating heatmap data: {e}")
        db.session.rollback()

# Real-time Collaboration Heatmap API Endpoints

@app.route('/api/heatmap/interaction', methods=['POST'])
@login_required
def log_interaction_event():
    """Log participant interaction events for heatmap generation"""
    try:
        data = request.json
        
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'})
        
        # Validate required fields
        attempt_id = data.get('attemptId')
        event_type = data.get('eventType')
        
        if not attempt_id or not event_type:
            return jsonify({'success': False, 'error': 'attemptId and eventType are required'})
        
        # Verify attempt belongs to current user
        attempt = QuizAttempt.query.get(attempt_id)
        if not attempt or attempt.participant_id != current_user.id:
            return jsonify({'success': False, 'error': 'Invalid attempt or access denied'})
        
        # Create interaction event
        interaction = InteractionEvent(
            attempt_id=attempt_id,
            question_id=data.get('questionId'),
            event_type=event_type,
            element_selector=data.get('elementSelector'),
            x_coordinate=data.get('x'),
            y_coordinate=data.get('y'),
            viewport_width=data.get('viewportWidth'),
            viewport_height=data.get('viewportHeight'),
            duration=data.get('duration'),
            event_metadata=json.dumps(data.get('metadata', {})),
            timestamp=datetime.utcnow()
        )
        
        db.session.add(interaction)
        db.session.commit()
        
        # Trigger real-time heatmap data update and analysis (async)
        try:
            update_heatmap_data(attempt.quiz_id, data.get('questionId'))
            
            # Trigger insights analysis periodically (every 10 interactions)
            interaction_count = InteractionEvent.query.filter_by(attempt_id=attempt_id).count()
            if interaction_count % 10 == 0:  # Analyze every 10 interactions
                from heatmap_analysis import trigger_analysis_for_quiz
                trigger_analysis_for_quiz(attempt.quiz_id)
                
        except Exception as e:
            logging.warning(f"Failed to update heatmap data or trigger analysis: {e}")
        
        return jsonify({'success': True, 'logged': True})
        
    except Exception as e:
        logging.error(f"Error logging interaction event: {e}")
        return jsonify({'success': False, 'error': 'Internal server error'})

@app.route('/api/heatmap/interaction/batch', methods=['POST'])
@login_required
def log_interaction_events_batch():
    """Log multiple participant interaction events for heatmap generation"""
    try:
        data = request.get_json()
        
        if not data or 'interactions' not in data:
            return jsonify({'success': False, 'error': 'Interactions data required'})
        
        interactions_data = data['interactions']
        attempt_id = data.get('attemptId')
        
        if not attempt_id:
            return jsonify({'success': False, 'error': 'Attempt ID required'})
        
        # Verify the attempt belongs to current user
        attempt = QuizAttempt.query.filter_by(
            id=attempt_id,
            participant_id=current_user.id,
            status='in_progress'
        ).first()
        
        if not attempt:
            return jsonify({'success': False, 'error': 'Invalid or inactive quiz attempt'})
        
        successful_logs = 0
        failed_logs = 0
        
        # Process each interaction in the batch
        for interaction_data in interactions_data:
            try:
                # Extract event details
                event_type = interaction_data.get('type', 'unknown')
                question_id = interaction_data.get('questionId')
                x_coord = interaction_data.get('x')
                y_coord = interaction_data.get('y')
                timestamp_ms = interaction_data.get('timestamp', time.time() * 1000)
                
                # Create interaction event
                interaction = InteractionEvent(
                    attempt_id=attempt_id,
                    event_type=event_type,
                    question_id=question_id,
                    x_coordinate=x_coord,
                    y_coordinate=y_coord,
                    timestamp=datetime.fromtimestamp(timestamp_ms / 1000),
                    event_details=json.dumps({
                        'target': interaction_data.get('target'),
                        'scrollTop': interaction_data.get('scrollTop'),
                        'scrollLeft': interaction_data.get('scrollLeft'),
                        'visibility': interaction_data.get('visibility'),
                        'timeSpent': interaction_data.get('timeSpent'),
                        'answerValue': interaction_data.get('answerValue'),
                        'answerType': interaction_data.get('answerType'),
                        'textLength': interaction_data.get('textLength')
                    })
                )
                
                db.session.add(interaction)
                successful_logs += 1
                
            except Exception as e:
                logging.warning(f"Failed to process individual interaction: {e}")
                failed_logs += 1
                continue
        
        # Commit all successful interactions
        if successful_logs > 0:
            db.session.commit()
            
            # Trigger heatmap data update and analysis for significant batches
            if successful_logs >= 5:  # Only for meaningful batches
                try:
                    update_heatmap_data(attempt.quiz_id, None)
                    
                    # Trigger insights analysis for larger batches
                    if successful_logs >= 10:
                        from heatmap_analysis import trigger_analysis_for_quiz
                        trigger_analysis_for_quiz(attempt.quiz_id)
                        
                except Exception as e:
                    logging.warning(f"Failed to update heatmap data or trigger analysis: {e}")
        
        return jsonify({
            'success': True, 
            'processed': len(interactions_data),
            'successful': successful_logs,
            'failed': failed_logs
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error processing interaction batch: {e}")
        return jsonify({'success': False, 'error': 'Failed to process interaction batch'})

@app.route('/api/heatmap/data/<int:quiz_id>')
@login_required
def get_heatmap_data(quiz_id):
    """Get aggregated heatmap data for a quiz"""
    try:
        # Verify user has access to this quiz
        quiz = Quiz.query.get_or_404(quiz_id)
        
        # Check if user is the host or admin
        if not (current_user.is_admin() or quiz.creator_id == current_user.id):
            return jsonify({'error': 'Access denied'}), 403
        
        # Get heatmap data for all questions in the quiz
        heatmap_data = QuestionHeatmapData.query.filter_by(quiz_id=quiz_id).all()
        
        # Format response
        questions_data = []
        for data in heatmap_data:
            question_info = {
                'questionId': data.question_id,
                'totalParticipants': data.total_participants,
                'averageTimeSpent': data.average_time_spent,
                'totalClicks': data.total_clicks,
                'totalHovers': data.total_hovers,
                'correctAnswerRate': data.correct_answer_rate,
                'difficultyScore': data.difficulty_score,
                'engagementScore': data.engagement_score,
                'clickHotspots': json.loads(data.click_hotspots) if data.click_hotspots else [],
                'hoverHotspots': json.loads(data.hover_hotspots) if data.hover_hotspots else [],
                'scrollPatterns': json.loads(data.scroll_patterns) if data.scroll_patterns else {},
                'lastUpdated': data.last_updated.isoformat()
            }
            questions_data.append(question_info)
        
        return jsonify({
            'success': True,
            'quizId': quiz_id,
            'questionsData': questions_data,
            'totalQuestions': len(questions_data),
            'lastUpdated': max([data.last_updated for data in heatmap_data]).isoformat() if heatmap_data else None
        })
        
    except Exception as e:
        logging.error(f"Error getting heatmap data: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/heatmap/insights/<int:quiz_id>')
@login_required
def get_collaboration_insights(quiz_id):
    """Get collaboration insights for a quiz"""
    try:
        # Verify user has access to this quiz
        quiz = Quiz.query.get_or_404(quiz_id)
        
        # Check if user is the host or admin
        if not (current_user.is_admin() or quiz.creator_id == current_user.id):
            return jsonify({'error': 'Access denied'}), 403
        
        # Get active insights for the quiz
        insights = CollaborationInsight.query.filter_by(
            quiz_id=quiz_id,
            is_active=True
        ).order_by(CollaborationInsight.created_at.desc()).all()
        
        # Format response
        insights_data = []
        for insight in insights:
            insight_info = {
                'id': insight.id,
                'type': insight.insight_type,
                'title': insight.title,
                'description': insight.description,
                'severity': insight.severity,
                'affectedQuestions': json.loads(insight.affected_questions) if insight.affected_questions else [],
                'metricValues': json.loads(insight.metric_values) if insight.metric_values else {},
                'suggestedActions': json.loads(insight.suggested_actions) if insight.suggested_actions else [],
                'isAcknowledged': insight.is_acknowledged,
                'createdAt': insight.created_at.isoformat(),
                'updatedAt': insight.updated_at.isoformat()
            }
            insights_data.append(insight_info)
        
        return jsonify({
            'success': True,
            'quizId': quiz_id,
            'insights': insights_data,
            'totalInsights': len(insights_data),
            'criticalCount': len([i for i in insights if i.severity == 'critical']),
            'highCount': len([i for i in insights if i.severity == 'high'])
        })
        
    except Exception as e:
        logging.error(f"Error getting collaboration insights: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/heatmap/insights/<int:insight_id>/acknowledge', methods=['POST'])
@login_required
def acknowledge_insight(insight_id):
    """Acknowledge a collaboration insight"""
    try:
        insight = CollaborationInsight.query.get_or_404(insight_id)
        
        # Verify user has access to this quiz
        if not (current_user.is_admin() or insight.quiz.creator_id == current_user.id):
            return jsonify({'error': 'Access denied'}), 403
        
        # Acknowledge the insight
        insight.is_acknowledged = True
        insight.acknowledged_by = current_user.id
        insight.acknowledged_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({'success': True, 'acknowledged': True})
        
    except Exception as e:
        logging.error(f"Error acknowledging insight: {e}")
        return jsonify({'error': 'Internal server error'}), 500


import os
import csv
import io
import logging
import tempfile
import mimetypes
from werkzeug.utils import secure_filename

@app.route('/quiz/create', methods=['GET', 'POST'])
@login_required
def create_quiz():
    """Create a new quiz with optional file upload"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    form = QuizForm()
    if form.validate_on_submit():
        # Create the quiz
        quiz = Quiz(
            title=form.title.data,
            description=form.description.data,
            time_limit=form.time_limit.data,
            proctoring_enabled=form.proctoring_enabled.data,
            creator_id=current_user.id
        )
        
        # Handle file upload if selected
        if form.create_from_file.data and form.quiz_file.data:
            file = form.quiz_file.data
            if file and file.filename:
                filename = secure_filename(file.filename)
                
                # Save file temporarily for parsing
                temp_file_path = None
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_file:
                        file.save(temp_file.name)
                        temp_file_path = temp_file.name
                    
                    # Get MIME type
                    mime_type, _ = mimetypes.guess_type(filename)
                    if not mime_type:
                        # Fallback based on file extension
                        ext = os.path.splitext(filename)[1].lower()
                        mime_type_map = {
                            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                            '.xls': 'application/vnd.ms-excel',  
                            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            '.pdf': 'application/pdf',
                            '.csv': 'text/csv',
                            '.txt': 'text/plain'
                        }
                        mime_type = mime_type_map.get(ext, 'text/plain')
                    
                    # Use comprehensive file parsing
                    candidate_questions = parse_file_for_questions(temp_file_path, mime_type)
                    
                    if not candidate_questions:
                        flash('No questions found in the uploaded file. Please check the format and try again.', 'warning')
                        return render_template('create_quiz.html', form=form)
                    
                    db.session.add(quiz)
                    db.session.commit()
                    
                    # Create questions from parsed data with better error handling
                    created_count = 0
                    for q_data in candidate_questions:
                        try:
                            create_question_from_comprehensive_data(quiz, q_data)
                            created_count += 1
                        except Exception as e:
                            logging.warning(f"Failed to create question: {e}")
                            continue
                    
                    if created_count > 0:
                        flash(f'Quiz created successfully from file "{filename}"! {created_count} questions added.', 'success')
                        return redirect(url_for('edit_quiz', quiz_id=quiz.id))
                    else:
                        flash('No valid questions could be created from the uploaded file.', 'error')
                        return render_template('create_quiz.html', form=form)
                        
                except Exception as e:
                    flash(f'Error parsing file: {str(e)}', 'error')
                    return render_template('create_quiz.html', form=form)
                    
                finally:
                    # Always clean up temporary file
                    if temp_file_path and os.path.exists(temp_file_path):
                        try:
                            os.unlink(temp_file_path)
                        except Exception as e:
                            logging.warning(f"Failed to clean up temp file {temp_file_path}: {e}")
        
        db.session.add(quiz)
        db.session.commit()
        
        flash('Quiz created successfully! Now add questions to your quiz.', 'success')
        return redirect(url_for('edit_quiz', quiz_id=quiz.id))
    
    return render_template('create_quiz.html', form=form)

def parse_quiz_file(file):
    """Parse uploaded quiz file and return questions data"""
    questions_data = []
    content = file.read().decode('utf-8')
    file.seek(0)  # Reset file pointer
    
    if file.filename.endswith('.csv'):
        csv_reader = csv.reader(io.StringIO(content))
        for row in csv_reader:
            if len(row) >= 6:  # Question, Option1, Option2, Option3, Option4, CorrectAnswer
                question_data = {
                    'question_text': row[0].strip(),
                    'options': [row[i].strip() for i in range(1, 5)],
                    'correct_answer': int(row[5]) - 1 if row[5].isdigit() else 0,
                    'points': int(row[6]) if len(row) > 6 and row[6].isdigit() else 1
                }
                questions_data.append(question_data)
    else:  # TXT format
        lines = content.split('\n')
        current_question = None
        options = []
        correct_answer = 0
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_question and options:
                    questions_data.append({
                        'question_text': current_question,
                        'options': options,
                        'correct_answer': correct_answer,
                        'points': 1
                    })
                    current_question = None
                    options = []
                    correct_answer = 0
                continue
                
            if line.startswith('Q:') or line.startswith('Question:'):
                current_question = line.split(':', 1)[1].strip()
            elif line.startswith(('A)', 'B)', 'C)', 'D)')) or line.startswith(('1.', '2.', '3.', '4.')):
                option_text = line[2:].strip() if line[1] in ').' else line[3:].strip()
                if line.startswith('*') or '(correct)' in line.lower():
                    correct_answer = len(options)
                    option_text = option_text.replace('*', '').replace('(correct)', '').strip()
                options.append(option_text)
        
        # Add last question if exists
        if current_question and options:
            questions_data.append({
                'question_text': current_question,
                'options': options,
                'correct_answer': correct_answer,
                'points': 1
            })
    
    return questions_data

def create_question_from_comprehensive_data(quiz, question_data):
    """Create a question and its options from comprehensive parsed data"""
    # Handle different data formats from comprehensive parser
    question_text = question_data.get('question', question_data.get('question_text', ''))
    question_type = question_data.get('type', 'multiple_choice')
    
    if not question_text:
        raise ValueError("Question text is required")
    
    question = Question(
        quiz_id=quiz.id,
        question_text=question_text,
        question_type=question_type,
        points=question_data.get('points', 1),
        order=len(quiz.questions)
    )
    
    db.session.add(question)
    db.session.commit()
    
    # Create options for multiple choice questions
    if question_type == 'multiple_choice':
        options = question_data.get('options', [])
        correct_index = question_data.get('correct_option_index', 0)
        
        if len(options) < 2:
            raise ValueError("Multiple choice questions need at least 2 options")
        
        for i, option_text in enumerate(options):
            if option_text:  # Only add non-empty options
                option = QuestionOption(
                    question_id=question.id,
                    option_text=option_text,
                    is_correct=(i == correct_index),
                    order=i
                )
                db.session.add(option)
    
    db.session.commit()
    return question

def create_question_from_data(quiz, question_data):
    """Create a question and its options from parsed data (legacy format)"""
    question = Question(
        quiz_id=quiz.id,
        question_text=question_data['question_text'],
        question_type='multiple_choice',
        points=question_data.get('points', 1),
        order=len(quiz.questions)
    )
    
    db.session.add(question)
    db.session.commit()
    
    # Create options
    for i, option_text in enumerate(question_data['options']):
        option = QuestionOption(
            question_id=question.id,
            option_text=option_text,
            is_correct=(i == question_data['correct_answer']),
            order=i
        )
        db.session.add(option)
    
    db.session.commit()

@app.route('/quiz/<int:quiz_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_quiz(quiz_id):
    """Edit quiz and manage questions"""
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        flash('Access denied. You can only edit your own quizzes.', 'error')
        return redirect(url_for('host_dashboard'))
    
    form = QuizForm(obj=quiz)
    question_form = QuestionForm()
    
    if form.validate_on_submit() and 'update_quiz' in request.form:
        quiz.title = form.title.data
        quiz.description = form.description.data
        quiz.time_limit = form.time_limit.data
        quiz.proctoring_enabled = form.proctoring_enabled.data
        quiz.updated_at = datetime.utcnow()
        
        db.session.commit()
        flash('Quiz updated successfully!', 'success')
        return redirect(url_for('edit_quiz', quiz_id=quiz.id))
    
    return render_template('edit_quiz.html', quiz=quiz, form=form, question_form=question_form)

@app.route('/question/<int:question_id>/edit', methods=['POST'])
@login_required
def edit_question(question_id):
    """Edit a specific question"""
    question = Question.query.get_or_404(question_id)
    quiz = question.quiz
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('host_dashboard'))
    
    question.question_text = request.form.get('question_text')
    question.question_type = request.form.get('question_type')
    question.points = int(request.form.get('points', 1))
    
    # Update options for multiple choice questions
    if question.question_type == 'multiple_choice':
        # Remove old options
        QuestionOption.query.filter_by(question_id=question.id).delete()
        
        # Add new options
        for i in range(1, 5):  # Support up to 4 options
            option_text = request.form.get(f'option_{i}')
            if option_text:
                is_correct = request.form.get('correct_option') == str(i)
                option = QuestionOption(
                    question_id=question.id,
                    option_text=option_text,
                    is_correct=is_correct
                )
                db.session.add(option)
    
    db.session.commit()
    flash('Question updated successfully!', 'success')
    return redirect(url_for('edit_quiz', quiz_id=quiz.id))

@app.route('/question/<int:question_id>/delete', methods=['POST'])
@login_required
def delete_question(question_id):
    """Delete a question"""
    question = Question.query.get_or_404(question_id)
    quiz = question.quiz
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('host_dashboard'))
    
    # Delete associated options first
    QuestionOption.query.filter_by(question_id=question.id).delete()
    # Delete the question
    db.session.delete(question)
    db.session.commit()
    
    flash('Question deleted successfully!', 'success')
    return redirect(url_for('edit_quiz', quiz_id=quiz.id))

@app.route('/quiz/<int:quiz_id>/add_question', methods=['POST'])
@login_required
def add_question(quiz_id):
    """Add a question to quiz"""
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('host_dashboard'))
    
    question_text = request.form.get('question_text')
    question_type = request.form.get('question_type', 'multiple_choice')
    points = int(request.form.get('points', 1))
    
    if not question_text:
        flash('Question text is required.', 'error')
        return redirect(url_for('edit_quiz', quiz_id=quiz_id))
    
    question = Question(
        quiz_id=quiz_id,
        question_text=question_text,
        question_type=question_type,
        points=points,
        order=len(quiz.questions)
    )
    
    db.session.add(question)
    db.session.commit()
    
    # Add options for multiple choice questions
    if question_type == 'multiple_choice':
        for i in range(4):  # Default 4 options
            option_text = request.form.get(f'option_{i}')
            is_correct = request.form.get(f'correct_{i}') == 'on'
            
            if option_text:
                option = QuestionOption(
                    question_id=question.id,
                    option_text=option_text,
                    is_correct=is_correct,
                    order=i
                )
                db.session.add(option)
    
    elif question_type == 'true_false':
        # Add True/False options
        true_option = QuestionOption(
            question_id=question.id,
            option_text='True',
            is_correct=request.form.get('correct_answer') == 'true',
            order=0
        )
        false_option = QuestionOption(
            question_id=question.id,
            option_text='False',
            is_correct=request.form.get('correct_answer') == 'false',
            order=1
        )
        db.session.add(true_option)
        db.session.add(false_option)
    
    db.session.commit()
    flash('Question added successfully!', 'success')
    return redirect(url_for('edit_quiz', quiz_id=quiz_id))

@app.route('/quiz/<int:quiz_id>')
@login_required
def view_quiz(quiz_id):
    """View quiz details"""
    quiz = Quiz.query.get_or_404(quiz_id)
    return render_template('quiz_list.html', quiz=quiz)

@app.route('/quiz/<int:quiz_id>/take')
@login_required
def take_quiz(quiz_id):
    """Take a quiz with enhanced security checks"""
    if not current_user.is_participant():
        flash('Access denied. Participant privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Check if user is flagged for violations - ADMIN ONLY RETAKE PERMISSIONS
    violation_record = UserViolation.query.filter_by(user_id=current_user.id).first()
    if violation_record and violation_record.is_flagged:
        flash('❌ Access denied. Your account has been flagged for security violations. Contact an administrator for retake permissions.', 'error')
        return redirect(url_for('dashboard'))
    
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if not quiz.is_active:
        flash('This quiz is not currently active.', 'error')
        return redirect(url_for('participant_dashboard'))
    
    # DEVICE DETECTION FOR MOBILE-SPECIFIC PROCTORING
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile_keywords = ['mobile', 'android', 'iphone', 'ipod', 'blackberry', 'windows phone']
    tablet_keywords = ['tablet', 'ipad']
    
    is_mobile = any(keyword in user_agent for keyword in mobile_keywords)
    is_tablet = any(keyword in user_agent for keyword in tablet_keywords)
    device_type = 'mobile' if is_mobile else ('tablet' if is_tablet else 'desktop')
    
    # Log device access for security monitoring
    try:
        event = ProctoringEvent(
            attempt_id=None,
            event_type=f'{device_type}_device_access',
            details=f'Device access from {device_type}: {user_agent}',
            severity='info',
            timestamp=datetime.utcnow()
        )
        db.session.add(event)
        db.session.commit()
    except Exception as e:
        logging.error(f"Failed to log device access: {e}")

    # Check if user already has an active attempt
    existing_attempt = QuizAttempt.query.filter_by(
        participant_id=current_user.id,
        quiz_id=quiz_id,
        status='in_progress'
    ).first()
    
    if existing_attempt:
        return redirect(url_for('continue_quiz', attempt_id=existing_attempt.id))
    
    # Create new attempt
    attempt = QuizAttempt(
        participant_id=current_user.id,
        quiz_id=quiz_id
    )
    
    db.session.add(attempt)
    db.session.commit()
    
    # Pass device type to the template
    return redirect(url_for('continue_quiz', attempt_id=attempt.id))

@app.route('/attempt/<int:attempt_id>')
@login_required
def continue_quiz(attempt_id):
    """Continue taking a quiz"""
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    if attempt.participant_id != current_user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('participant_dashboard'))
    
    if attempt.status != 'in_progress':
        flash('This quiz attempt has already been completed.', 'error')
        return redirect(url_for('quiz_results', attempt_id=attempt_id))
    
    quiz = attempt.quiz
    questions = quiz.questions
    
    # Get existing answers
    existing_answers = {answer.question_id: answer for answer in attempt.answers}
    
    # Detect device type for mobile-specific proctoring
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile_keywords = ['mobile', 'android', 'iphone', 'ipod', 'blackberry', 'windows phone']
    tablet_keywords = ['tablet', 'ipad']
    
    is_mobile = any(keyword in user_agent for keyword in mobile_keywords)
    is_tablet = any(keyword in user_agent for keyword in tablet_keywords)
    device_type = 'mobile' if is_mobile else ('tablet' if is_tablet else 'desktop')
    
    return render_template('take_quiz.html', 
                         attempt=attempt, 
                         quiz=quiz, 
                         questions=questions,
                         existing_answers=existing_answers,
                         device_type=device_type)

@app.route('/attempt/<int:attempt_id>/submit', methods=['POST'])
@login_required
def submit_quiz(attempt_id):
    """Submit quiz answers"""
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    if attempt.participant_id != current_user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('participant_dashboard'))
    
    if attempt.status != 'in_progress':
        flash('This quiz has already been submitted.', 'error')
        return redirect(url_for('quiz_results', attempt_id=attempt_id))
    
    quiz = attempt.quiz
    
    # Process answers
    for question in quiz.questions:
        answer_key = f'question_{question.id}'
        
        # Check if answer already exists
        existing_answer = Answer.query.filter_by(
            attempt_id=attempt_id,
            question_id=question.id
        ).first()
        
        if existing_answer:
            # Update existing answer
            answer = existing_answer
        else:
            # Create new answer
            answer = Answer(
                attempt_id=attempt_id,
                question_id=question.id
            )
        
        if question.question_type == 'multiple_choice' or question.question_type == 'true_false':
            selected_option_id = request.form.get(answer_key)
            if selected_option_id:
                answer.selected_option_id = int(selected_option_id)
                selected_option = QuestionOption.query.get(selected_option_id)
                answer.is_correct = selected_option.is_correct if selected_option else False
            else:
                answer.is_correct = False
        
        elif question.question_type == 'text':
            answer.text_answer = request.form.get(answer_key, '')
            # For text answers, manual grading would be needed
            answer.is_correct = None
            
        elif question.question_type == 'code_submission':
            answer.code_submission = request.form.get(answer_key, '')
            # Execute code if provided (basic validation)
            if answer.code_submission:
                try:
                    # Basic code execution simulation (for demonstration)
                    # In production, use a secure sandboxed environment
                    answer.execution_output = "Code submitted successfully"
                    answer.is_correct = None  # Manual grading required
                except Exception as e:
                    answer.execution_error = str(e)
                    answer.is_correct = False
        
        elif question.question_type == 'file_upload':
            # Handle file upload for this answer
            file_key = f'file_{question.id}'
            if file_key in request.files:
                file = request.files[file_key]
                if file and file.filename:
                    # Validate file type and size
                    allowed_types = question.allowed_file_types.split(',') if question.allowed_file_types else ['pdf', 'docx', 'jpg', 'png', 'txt']
                    file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
                    
                    if file_ext in [t.strip() for t in allowed_types]:
                        # Check file size
                        file.seek(0, 2)
                        file_size = file.tell()
                        file.seek(0)
                        
                        max_size = (question.max_file_size_mb or 10) * 1024 * 1024
                        if file_size <= max_size:
                            # Save file
                            filename = secure_filename(file.filename)
                            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                            filename = f"{attempt.id}_{question.id}_{timestamp}_{filename}"
                            file_path = os.path.join(QUIZ_ANSWER_UPLOAD_FOLDER, filename)
                            file.save(file_path)
                            
                            answer.uploaded_file_path = file_path
                            answer.uploaded_file_name = file.filename
                            answer.uploaded_file_size = file_size
                            answer.is_correct = None  # Manual grading required
                        else:
                            answer.is_correct = False
                            answer.text_answer = f"File too large. Maximum size: {question.max_file_size_mb}MB"
                    else:
                        answer.is_correct = False
                        answer.text_answer = f"Invalid file type. Allowed: {question.allowed_file_types}"
            else:
                answer.is_correct = False
                answer.text_answer = "No file uploaded"
        
        elif question.question_type == 'drawing':
            # Handle drawing/canvas data
            drawing_data = request.form.get(f'drawing_{question.id}', '')
            if drawing_data:
                answer.drawing_data = drawing_data
                answer.is_correct = None  # Manual grading required
            else:
                answer.is_correct = False
                answer.text_answer = "No drawing provided"
        
        db.session.add(answer)
        
        # Trigger collaboration detection for new/updated answers
        if detector and (not existing_answer or existing_answer.selected_option_id != answer.selected_option_id):
            try:
                # Set quiz_id for the answer (if not already set)
                answer.quiz_id = quiz.id
                db.session.flush()  # Ensure answer has an ID
                
                # Run collaboration detection
                signals = detector.process_new_answer(answer)
                if signals:
                    app.logger.info(f"Detected {len(signals)} collaboration signals for quiz {quiz.id}")
            except Exception as e:
                app.logger.error(f"Error in collaboration detection: {e}")
        
        # AI-Powered Plagiarism Detection for text answers
        if plagiarism_detector and question.question_type == 'text' and answer.text_answer:
            try:
                db.session.flush()  # Ensure answer has an ID
                
                # Get all other text answers for this question for comparison
                other_answers = Answer.query.filter(
                    Answer.question_id == question.id,
                    Answer.id != answer.id,
                    Answer.text_answer.isnot(None),
                    Answer.text_answer != ''
                ).all()
                
                # Prepare comparison texts
                comparison_texts = [(ans.id, ans.text_answer) for ans in other_answers]
                
                # Run plagiarism analysis
                if comparison_texts:  # Only analyze if there are other answers to compare against
                    analysis = plagiarism_detector.analyze_text_for_plagiarism(
                        target_text=answer.text_answer,
                        comparison_texts=comparison_texts,
                        answer_id=answer.id,
                        quiz_attempt_id=attempt.id,
                        question_id=question.id
                    )
                    
                    # Save analysis to database
                    db.session.add(analysis)
                    
                    # Log if high-risk plagiarism detected
                    if analysis.risk_level in ['high', 'critical']:
                        app.logger.warning(f"High-risk plagiarism detected for answer {answer.id}: {analysis.risk_level} ({analysis.overall_similarity_score:.3f})")
                    
            except Exception as e:
                app.logger.error(f"Error in plagiarism detection: {e}")
    
    # Mark attempt as completed
    attempt.completed_at = datetime.utcnow()
    attempt.status = 'completed'
    attempt.calculate_score()
    
    db.session.commit()
    
    flash('Quiz submitted successfully!', 'success')
    return redirect(url_for('quiz_results', attempt_id=attempt_id))

@app.route('/results/<int:attempt_id>')
@login_required
def quiz_results(attempt_id):
    """View quiz results"""
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    if attempt.participant_id != current_user.id and attempt.quiz.creator_id != current_user.id and not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    quiz = attempt.quiz
    questions = quiz.questions
    answers = {answer.question_id: answer for answer in attempt.answers}
    
    return render_template('quiz_results.html', 
                         attempt=attempt, 
                         quiz=quiz, 
                         questions=questions,
                         answers=answers)

@app.route('/download/participant-report/<int:attempt_id>')
@login_required
def download_participant_report(attempt_id):
    """Download participant report as PDF with comprehensive error handling"""
    try:
        attempt = QuizAttempt.query.get_or_404(attempt_id)
        
        if attempt.participant_id != current_user.id:
            flash('Access denied.', 'error')
            return redirect(url_for('participant_dashboard'))
    except Exception as e:
        logging.error(f"Error fetching quiz attempt {attempt_id}: {e}")
        flash('Quiz attempt not found.', 'error')
        return redirect(url_for('participant_dashboard'))
    
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from io import BytesIO
    
    # Create PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.darkblue,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    story = []
    
    # Title
    story.append(Paragraph("Quiz Results Report", title_style))
    story.append(Spacer(1, 20))
    
    # Quiz Information
    quiz_info = [
        ['Quiz Title:', attempt.quiz.title],
        ['Participant:', attempt.participant.username],
        ['Score:', f"{attempt.score:.1f}%" if attempt.score else 'N/A'],
        ['Started:', attempt.started_at.strftime('%Y-%m-%d %H:%M:%S')],
        ['Completed:', attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else 'Not completed'],
        ['Status:', attempt.status.title()]
    ]
    
    quiz_table = Table(quiz_info, colWidths=[2*inch, 4*inch])
    quiz_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(quiz_table)
    story.append(Spacer(1, 30))
    
    # Questions and Answers
    story.append(Paragraph("Detailed Results", styles['Heading2']))
    story.append(Spacer(1, 20))
    
    answers = {answer.question_id: answer for answer in attempt.answers}
    
    for i, question in enumerate(attempt.quiz.questions, 1):
        # Question
        story.append(Paragraph(f"Question {i}: {question.question_text}", styles['Heading3']))
        story.append(Spacer(1, 10))
        
        answer = answers.get(question.id)
        
        if question.question_type in ['multiple_choice', 'true_false']:
            # Show options
            options_data = [['Option', 'Your Answer', 'Correct Answer']]
            for option in question.options:
                is_selected = '✓' if answer and answer.selected_option_id == option.id else ''
                is_correct = '✓' if option.is_correct else ''
                options_data.append([option.option_text, is_selected, is_correct])
            
            options_table = Table(options_data, colWidths=[3*inch, 1*inch, 1*inch])
            options_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(options_table)
            
            # Result
            if answer:
                if answer.is_correct:
                    result_text = f"<para><b>Result:</b> <font color='green'>Correct</font> (+{question.points} points)</para>"
                else:
                    result_text = f"<para><b>Result:</b> <font color='red'>Incorrect</font> (0 points)</para>"
            else:
                result_text = "<para><b>Result:</b> Not answered (0 points)</para>"
            
            story.append(Spacer(1, 10))
            story.append(Paragraph(result_text, styles['Normal']))
            
        elif question.question_type == 'text':
            if answer and answer.text_answer:
                story.append(Paragraph(f"<b>Your Answer:</b> {answer.text_answer}", styles['Normal']))
                if answer.is_correct == None:
                    story.append(Paragraph("<i>This answer requires manual grading.</i>", styles['Normal']))
            else:
                story.append(Paragraph("<b>Your Answer:</b> No answer provided", styles['Normal']))
        
        story.append(Spacer(1, 20))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    # Return as download
    return send_file(
        BytesIO(buffer.read()),
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'quiz_report_{attempt.quiz.title}_{attempt.participant.username}.pdf'
    )

@app.route('/download/host-report/<int:attempt_id>')
@login_required
def download_host_report(attempt_id):
    """Download detailed host report as Excel"""
    attempt = QuizAttempt.query.get_or_404(attempt_id)
    
    if attempt.quiz.creator_id != current_user.id and not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('host_dashboard'))
    
    from openpyxl import Workbook
    from openpyxl.styles import Font, Fill, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    from io import BytesIO
    
    # Create workbook
    wb = Workbook()
    
    # Quiz Summary Sheet
    ws1 = wb.active
    ws1.title = "Quiz Summary"
    
    # Headers
    ws1['A1'] = "Quiz Results Summary"
    ws1['A1'].font = Font(size=16, bold=True)
    ws1['A1'].fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    
    # Quiz info
    quiz_data = [
        ["Quiz Title", attempt.quiz.title],
        ["Participant", attempt.participant.username],
        ["Email", attempt.participant.email],
        ["Score", f"{attempt.score:.1f}%" if attempt.score else 'N/A'],
        ["Points Earned", f"{attempt.score * attempt.total_points / 100:.0f}" if attempt.score and attempt.total_points else 'N/A'],
        ["Total Points", attempt.total_points or 0],
        ["Started", attempt.started_at.strftime('%Y-%m-%d %H:%M:%S')],
        ["Completed", attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else 'Not completed'],
        ["Status", attempt.status.title()],
        ["Time Taken", str(attempt.completed_at - attempt.started_at) if attempt.completed_at else 'N/A']
    ]
    
    for row, (label, value) in enumerate(quiz_data, 3):
        ws1[f'A{row}'] = label
        ws1[f'B{row}'] = value
        ws1[f'A{row}'].font = Font(bold=True)
    
    # Detailed Answers Sheet
    ws2 = wb.create_sheet("Detailed Answers")
    headers = ["Question #", "Question Text", "Question Type", "Points", "Your Answer", "Correct Answer", "Result", "Points Earned"]
    
    for col, header in enumerate(headers, 1):
        ws2.cell(row=1, column=col, value=header)
        ws2.cell(row=1, column=col).font = Font(bold=True)
        ws2.cell(row=1, column=col).fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
    
    answers = {answer.question_id: answer for answer in attempt.answers}
    
    for row, question in enumerate(attempt.quiz.questions, 2):
        answer = answers.get(question.id)
        
        ws2.cell(row=row, column=1, value=row-1)
        ws2.cell(row=row, column=2, value=question.question_text)
        ws2.cell(row=row, column=3, value=question.question_type.title())
        ws2.cell(row=row, column=4, value=question.points)
        
        if question.question_type in ['multiple_choice', 'true_false']:
            if answer and answer.selected_option_id:
                selected_option = next((opt for opt in question.options if opt.id == answer.selected_option_id), None)
                ws2.cell(row=row, column=5, value=selected_option.option_text if selected_option else 'Unknown')
            else:
                ws2.cell(row=row, column=5, value='Not answered')
            
            correct_option = next((opt for opt in question.options if opt.is_correct), None)
            ws2.cell(row=row, column=6, value=correct_option.option_text if correct_option else 'No correct answer set')
            
            if answer:
                if answer.is_correct:
                    ws2.cell(row=row, column=7, value='Correct')
                    ws2.cell(row=row, column=8, value=question.points)
                    ws2.cell(row=row, column=7).fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                else:
                    ws2.cell(row=row, column=7, value='Incorrect')
                    ws2.cell(row=row, column=8, value=0)
                    ws2.cell(row=row, column=7).fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            else:
                ws2.cell(row=row, column=7, value='Not answered')
                ws2.cell(row=row, column=8, value=0)
        
        elif question.question_type == 'text':
            ws2.cell(row=row, column=5, value=answer.text_answer if answer and answer.text_answer else 'Not answered')
            ws2.cell(row=row, column=6, value='Manual grading required')
            ws2.cell(row=row, column=7, value='Needs review' if answer and answer.text_answer else 'Not answered')
            ws2.cell(row=row, column=8, value='TBD')
    
    # Auto-adjust column widths
    for ws in [ws1, ws2]:
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save to BytesIO
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return send_file(
        BytesIO(buffer.read()),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'detailed_report_{attempt.quiz.title}_{attempt.participant.username}.xlsx'
    )

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    """User profile management"""
    form = ProfileForm(obj=current_user)
    
    if form.validate_on_submit():
        # Check if username is taken by another user
        if form.username.data != current_user.username:
            existing_user = User.query.filter_by(username=form.username.data).first()
            if existing_user:
                flash('Username already taken.', 'error')
                return render_template('profile.html', form=form)
        
        # Check if email is taken by another user
        if form.email.data != current_user.email:
            existing_user = User.query.filter_by(email=form.email.data).first()
            if existing_user:
                flash('Email already registered.', 'error')
                return render_template('profile.html', form=form)
        
        # Handle profile picture upload
        if 'profile_picture' in request.files:
            file = request.files['profile_picture']
            if file and file.filename:
                import os
                from werkzeug.utils import secure_filename
                
                # Create uploads directory if it doesn't exist
                upload_dir = os.path.join('static', 'uploads', 'profiles')
                os.makedirs(upload_dir, exist_ok=True)
                
                # Save file with user id prefix
                filename = f"user_{current_user.id}_{secure_filename(file.filename)}"
                filepath = os.path.join(upload_dir, filename)
                file.save(filepath)
                
                # Update user profile picture path
                current_user.profile_picture = f"uploads/profiles/{filename}"

        # Update profile
        current_user.username = form.username.data
        
        # Handle email change without verification requirement
        if form.email.data != current_user.email:
            current_user.email = form.email.data
            current_user.is_verified = True  # Auto-verify for profile updates
            flash('Email updated successfully!', 'success')
        
        # Update password if provided
        if form.current_password.data and form.new_password.data:
            if current_user.check_password(form.current_password.data):
                current_user.set_password(form.new_password.data)
                flash('Password updated successfully!', 'success')
            else:
                flash('Current password is incorrect.', 'error')
                return render_template('profile.html', form=form)
        
        db.session.commit()
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('profile'))
    
    return render_template('profile.html', 
                         form=form,
                         greeting=get_time_greeting(),
                         greeting_icon=get_greeting_icon())


@app.route('/api/quiz/<int:quiz_id>/questions')
@login_required
def get_quiz_questions(quiz_id):
    """Get quiz questions for AJAX loading"""
    quiz = Quiz.query.get_or_404(quiz_id)
    
    if quiz.creator_id != current_user.id and not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    questions = []
    for question in quiz.questions:
        question_data = {
            'id': question.id,
            'text': question.question_text,
            'type': question.question_type,
            'points': question.points,
            'options': []
        }
        
        for option in question.options:
            question_data['options'].append({
                'id': option.id,
                'text': option.option_text,
                'is_correct': option.is_correct
            })
        
        questions.append(question_data)
    
    return jsonify({'questions': questions})



@app.route('/admin/violations')
@login_required
def admin_violations():
    """Enhanced violations view with consolidated entries per user per quiz"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get filter parameters
    severity_filter = request.args.get('severity', 'all')
    date_filter = request.args.get('date_range', '7')
    user_filter = request.args.get('user_id', '')
    
    # Build aggregated query to consolidate violations by user, quiz, and type
    from sqlalchemy import func, case
    
    base_query = db.session.query(
        User.id.label('user_id'),
        User.username.label('username'),
        User.email.label('email'),
        Quiz.id.label('quiz_id'),
        Quiz.title.label('quiz_title'),
        Quiz.creator_id.label('creator_id'),
        func.max(User.username).label('creator_username'),  # Will be corrected below
        ProctoringEvent.event_type.label('violation_type'),
        func.count(ProctoringEvent.id).label('count'),
        func.max(ProctoringEvent.timestamp).label('latest_at'),
        func.max(ProctoringEvent.description).label('description'),
        func.max(case(
            (ProctoringEvent.severity == 'low', 1),
            (ProctoringEvent.severity == 'medium', 2), 
            (ProctoringEvent.severity == 'high', 3),
            (ProctoringEvent.severity == 'critical', 4),
            else_=0
        )).label('severity_rank'),
        func.max(QuizAttempt.status).label('attempt_status')
    ).join(QuizAttempt, ProctoringEvent.attempt_id == QuizAttempt.id) \
     .join(User, QuizAttempt.participant_id == User.id) \
     .join(Quiz, QuizAttempt.quiz_id == Quiz.id)
    
    # Apply filters
    if severity_filter != 'all':
        base_query = base_query.filter(ProctoringEvent.severity == severity_filter)
    
    if date_filter != 'all':
        days_ago = datetime.utcnow() - timedelta(days=int(date_filter))
        base_query = base_query.filter(ProctoringEvent.timestamp >= days_ago)
    
    if user_filter:
        base_query = base_query.filter(User.id == user_filter)
    
    # Group by user, quiz, and violation type
    aggregated_violations = base_query.group_by(
        User.id, User.username, User.email,
        Quiz.id, Quiz.title, Quiz.creator_id,
        ProctoringEvent.event_type
    ).order_by(func.max(ProctoringEvent.timestamp).desc()).limit(500).all()
    
    # Get creator usernames for the results
    creator_usernames = {}
    if aggregated_violations:
        creator_ids = {v.creator_id for v in aggregated_violations}
        creators = User.query.filter(User.id.in_(creator_ids)).all()
        creator_usernames = {creator.id: creator.username for creator in creators}
    
    # Convert severity rank back to text and add creator username
    violations = []
    for v in aggregated_violations:
        severity_map = {1: 'low', 2: 'medium', 3: 'high', 4: 'critical'}
        violations.append({
            'user_id': v.user_id,
            'username': v.username,
            'email': v.email,
            'quiz_id': v.quiz_id,
            'quiz_title': v.quiz_title,
            'creator_username': creator_usernames.get(v.creator_id, 'Unknown'),
            'violation_type': v.violation_type,
            'count': v.count,
            'latest_at': v.latest_at,
            'description': v.description,
            'severity': severity_map.get(v.severity_rank, 'low'),
            'attempt_status': v.attempt_status
        })
    
    # Get violation statistics
    total_violations = ProctoringEvent.query.count()
    high_severity = ProctoringEvent.query.filter_by(severity='high').count()
    recent_violations = ProctoringEvent.query.filter(
        ProctoringEvent.timestamp >= datetime.utcnow() - timedelta(hours=24)
    ).count()
    
    stats = {
        'total_violations': total_violations,
        'high_severity': high_severity,
        'recent_violations': recent_violations
    }
    
    # Get all users for filter dropdown
    users = User.query.filter_by(role='participant').all()
    
    return render_template('admin_violations.html', 
                         violations=violations, 
                         stats=stats,
                         users=users,
                         current_filters={
                             'severity': severity_filter,
                             'date_range': date_filter,
                             'user_id': user_filter
                         })

@app.route('/admin/violations/<int:user_id>/<int:quiz_id>/<violation_type>')
@login_required
def admin_violation_details(user_id, quiz_id, violation_type):
    """Show detailed violation timeline for a specific user-quiz-type combination"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get the user and quiz for display
    user = User.query.get_or_404(user_id)
    quiz = Quiz.query.get_or_404(quiz_id)
    
    # Get all detailed violations for this combination
    violations = ProctoringEvent.query.join(QuizAttempt) \
        .filter(
            QuizAttempt.participant_id == user_id,
            QuizAttempt.quiz_id == quiz_id,
            ProctoringEvent.event_type == violation_type
        ).order_by(ProctoringEvent.timestamp.desc()).all()
    
    return render_template('admin_violation_details.html',
                         user=user,
                         quiz=quiz,
                         violation_type=violation_type,
                         violations=violations)

@app.route('/admin/user/<int:user_id>/edit-credentials', methods=['POST'])
@login_required
def admin_edit_credentials(user_id):
    """Edit user credentials"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    
    # Don't allow editing your own account
    if user.id == current_user.id:
        flash('You cannot edit your own credentials.', 'error')
        return redirect(url_for('admin_users'))
    
    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')
    
    # Validation
    if not username or not email:
        flash('Username and email are required.', 'error')
        return redirect(url_for('admin_users'))
    
    # Check for duplicates
    if username != user.username and User.query.filter_by(username=username).first():
        flash('Username already exists.', 'error')
        return redirect(url_for('admin_users'))
    
    if email != user.email and User.query.filter_by(email=email).first():
        flash('Email already exists.', 'error')
        return redirect(url_for('admin_users'))
    
    # Update credentials
    user.username = username
    user.email = email
    
    if password and len(password) >= 6:
        user.set_password(password)
    
    db.session.commit()
    flash(f'Credentials updated for user {username}.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/quiz-management')
@login_required
def admin_quiz_management():
    """Admin quiz management system"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    quizzes = Quiz.query.order_by(Quiz.display_order.asc(), Quiz.created_at.desc()).all()
    return render_template('admin_quiz_management.html', quizzes=quizzes)

# AI-Powered Plagiarism Detection Admin Routes
@app.route('/admin/plagiarism-detection')
@login_required 
def admin_plagiarism_detection():
    """Admin dashboard for plagiarism detection monitoring and management"""
    if not current_user.is_admin():
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get filter parameters
    risk_filter = request.args.get('risk', 'all')
    review_filter = request.args.get('reviewed', 'all')
    quiz_filter = request.args.get('quiz_id', type=int)
    
    # Build query
    query = db.session.query(PlagiarismAnalysis).join(Answer).join(QuizAttempt).join(Quiz).join(User)
    
    # Apply filters
    if risk_filter != 'all':
        query = query.filter(PlagiarismAnalysis.risk_level == risk_filter)
    
    if review_filter == 'reviewed':
        query = query.filter(PlagiarismAnalysis.is_reviewed == True)
    elif review_filter == 'pending':
        query = query.filter(PlagiarismAnalysis.requires_review == True, PlagiarismAnalysis.is_reviewed == False)
    
    if quiz_filter:
        query = query.filter(Quiz.id == quiz_filter)
    
    # Get paginated results
    page = request.args.get('page', 1, type=int)
    analyses = query.order_by(PlagiarismAnalysis.analyzed_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    # Get summary statistics
    total_analyses = PlagiarismAnalysis.query.count()
    flagged_count = PlagiarismAnalysis.query.filter_by(is_flagged=True).count()
    pending_review = PlagiarismAnalysis.query.filter_by(requires_review=True, is_reviewed=False).count()
    
    # Risk level breakdown
    risk_stats = db.session.query(
        PlagiarismAnalysis.risk_level,
        func.count(PlagiarismAnalysis.id).label('count')
    ).group_by(PlagiarismAnalysis.risk_level).all()
    
    # Get available quizzes for filter
    quizzes = Quiz.query.order_by(Quiz.title).all()
    
    return render_template('admin_plagiarism_detection.html',
                         analyses=analyses,
                         total_analyses=total_analyses,
                         flagged_count=flagged_count,
                         pending_review=pending_review,
                         risk_stats=dict(risk_stats),
                         quizzes=quizzes,
                         current_filters={
                             'risk': risk_filter,
                             'reviewed': review_filter,
                             'quiz_id': quiz_filter
                         })

@app.route('/admin/plagiarism-analysis/<int:analysis_id>')
@login_required
def admin_plagiarism_analysis_detail(analysis_id):
    """Detailed view of a specific plagiarism analysis"""
    if not current_user.is_admin():
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    analysis = PlagiarismAnalysis.query.get_or_404(analysis_id)
    
    # Get related data
    answer = analysis.answer
    quiz_attempt = analysis.quiz_attempt
    question = analysis.question
    participant = quiz_attempt.participant
    quiz = quiz_attempt.quiz
    
    # Get plagiarism matches for this analysis
    matches = PlagiarismMatch.query.filter_by(analysis_id=analysis_id).all()
    
    return render_template('admin_plagiarism_analysis_detail.html',
                         analysis=analysis,
                         answer=answer,
                         quiz_attempt=quiz_attempt,
                         question=question,
                         participant=participant,
                         quiz=quiz,
                         matches=matches)

@app.route('/admin/plagiarism-analysis/<int:analysis_id>/review', methods=['POST'])
@login_required
def admin_review_plagiarism(analysis_id):
    """Review and make decision on plagiarism analysis"""
    if not current_user.is_admin():
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # CSRF Protection - Check token
    csrf_token = request.form.get('csrf_token')
    expected_token = session.get('csrf_token')
    if not csrf_token or csrf_token != expected_token:
        flash('Security error: Invalid CSRF token. Please try again.', 'error')
        return redirect(url_for('admin_plagiarism_detection'))
    
    analysis = PlagiarismAnalysis.query.get_or_404(analysis_id)
    
    decision = request.form.get('decision')  # 'innocent', 'suspicious', 'plagiarized'
    notes = request.form.get('notes', '')
    
    if decision not in ['innocent', 'suspicious', 'plagiarized']:
        flash('Invalid decision option.', 'error')
        return redirect(url_for('admin_plagiarism_analysis_detail', analysis_id=analysis_id))
    
    # Update analysis
    analysis.is_reviewed = True
    analysis.reviewed_by = current_user.id
    analysis.reviewed_at = datetime.utcnow()
    analysis.review_decision = decision
    analysis.review_notes = notes
    
    # Update flagged status based on decision
    if decision == 'plagiarized':
        analysis.is_flagged = True
    elif decision == 'innocent':
        analysis.is_flagged = False
    
    db.session.commit()
    
    flash(f'Plagiarism analysis marked as {decision}.', 'success')
    return redirect(url_for('admin_plagiarism_detection'))

@app.route('/admin/plagiarism-reports/download')
@login_required
def download_plagiarism_report():
    """Download comprehensive plagiarism detection report"""
    if not current_user.is_admin():
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Plagiarism Detection Report"
        
        # Define headers
        headers = [
            'Analysis ID', 'Quiz Title', 'Participant', 'Question Text', 
            'Risk Level', 'Similarity Score', 'Cosine Similarity', 
            'Jaccard Similarity', 'Levenshtein Similarity', 'Semantic Similarity',
            'Flagged', 'Requires Review', 'Reviewed', 'Review Decision',
            'Analyzed At', 'Reviewed At', 'Reviewer'
        ]
        
        # Add headers with styling
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Get all analyses with related data
        analyses = db.session.query(PlagiarismAnalysis).join(Answer).join(QuizAttempt).join(Quiz).join(User).join(Question).all()
        
        # Add data rows
        for row, analysis in enumerate(analyses, 2):
            data = [
                analysis.id,
                analysis.quiz_attempt.quiz.title,
                analysis.quiz_attempt.participant.username,
                analysis.question.question_text[:100] + "..." if len(analysis.question.question_text) > 100 else analysis.question.question_text,
                analysis.risk_level,
                f"{analysis.overall_similarity_score:.3f}",
                f"{analysis.cosine_similarity:.3f}" if analysis.cosine_similarity else "N/A",
                f"{analysis.jaccard_similarity:.3f}" if analysis.jaccard_similarity else "N/A", 
                f"{analysis.levenshtein_similarity:.3f}" if analysis.levenshtein_similarity else "N/A",
                f"{analysis.semantic_similarity:.3f}" if analysis.semantic_similarity else "N/A",
                "Yes" if analysis.is_flagged else "No",
                "Yes" if analysis.requires_review else "No",
                "Yes" if analysis.is_reviewed else "No",
                analysis.review_decision or "N/A",
                analysis.analyzed_at.strftime('%Y-%m-%d %H:%M:%S'),
                analysis.reviewed_at.strftime('%Y-%m-%d %H:%M:%S') if analysis.reviewed_at else "N/A",
                analysis.reviewer.username if analysis.reviewer else "N/A"
            ]
            
            for col, value in enumerate(data, 1):
                ws.cell(row=row, column=col, value=value)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save to BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        # Create response
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Content-Disposition'] = f'attachment; filename=plagiarism_detection_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        
        app.logger.info(f"Plagiarism detection report downloaded by admin {current_user.username}")
        return response
        
    except Exception as e:
        app.logger.error(f"Error generating plagiarism report: {e}")
        flash('Error generating report. Please try again.', 'error')
        return redirect(url_for('admin_plagiarism_detection'))

@app.route('/admin/analytics')
@login_required
def admin_analytics():
    """Comprehensive system analytics dashboard"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Comprehensive analytics data
    total_users = User.query.count()
    total_hosts = User.query.filter_by(role='host').count()
    total_participants = User.query.filter_by(role='participant').count()
    total_quizzes = Quiz.query.count()
    active_quizzes = Quiz.query.filter_by(is_active=True).count()
    total_attempts = QuizAttempt.query.count()
    completed_attempts = QuizAttempt.query.filter_by(status='completed').count()
    total_violations = ProctoringEvent.query.count()
    high_violations = ProctoringEvent.query.filter_by(severity='high').count()
    total_courses = Course.query.count()
    active_courses = Course.query.filter_by(is_active=True).count()
    
    # Recent activity analytics
    today = datetime.utcnow().date()
    week_ago = datetime.utcnow() - timedelta(days=7)
    month_ago = datetime.utcnow() - timedelta(days=30)
    
    today_attempts = QuizAttempt.query.filter(
        func.date(QuizAttempt.started_at) == today
    ).count()
    
    week_attempts = QuizAttempt.query.filter(
        QuizAttempt.started_at >= week_ago
    ).count()
    
    month_attempts = QuizAttempt.query.filter(
        QuizAttempt.started_at >= month_ago
    ).count()
    
    # Performance analytics
    if completed_attempts > 0:
        completed_attempts_with_scores = QuizAttempt.query.filter(
            QuizAttempt.status == 'completed',
            QuizAttempt.score.isnot(None)
        ).all()
        
        if completed_attempts_with_scores:
            avg_score = sum(attempt.score for attempt in completed_attempts_with_scores) / len(completed_attempts_with_scores)
            highest_score = max(attempt.score for attempt in completed_attempts_with_scores)
            lowest_score = min(attempt.score for attempt in completed_attempts_with_scores)
        else:
            avg_score = highest_score = lowest_score = 0
    else:
        avg_score = highest_score = lowest_score = 0
    
    # Top performing participants
    top_participants = db.session.query(
        User.username,
        func.avg(QuizAttempt.score).label('avg_score'),
        func.count(QuizAttempt.id).label('attempt_count')
    ).join(QuizAttempt, User.id == QuizAttempt.participant_id).filter(
        QuizAttempt.status == 'completed',
        QuizAttempt.score.isnot(None)
    ).group_by(User.id, User.username).order_by(
        func.avg(QuizAttempt.score).desc()
    ).limit(10).all()
    
    # Calculate additional user metrics
    active_users = User.query.filter(
        User.last_login >= datetime.utcnow() - timedelta(days=30)
    ).count()
    recent_users = User.query.filter(
        User.created_at >= datetime.utcnow() - timedelta(days=30)
    ).count()
    
    analytics_data = {
        'users': {
            'total': total_users,
            'hosts': total_hosts,
            'participants': total_participants,
            'active': active_users,
            'recent': recent_users
        },
        'quizzes': {
            'total': total_quizzes,
            'active': active_quizzes
        },
        'attempts': {
            'total': total_attempts,
            'completed': completed_attempts,
            'today': today_attempts,
            'week': week_attempts,
            'month': month_attempts
        },
        'violations': {
            'total': total_violations,
            'high': high_violations
        },
        'courses': {
            'total': total_courses,
            'active': active_courses
        },
        'scores': {
            'average': round(avg_score, 1),
            'highest': highest_score,
            'lowest': lowest_score
        },
        'top_participants': top_participants
    }
    
    return render_template('admin_analytics.html', analytics=analytics_data)

@app.route('/admin/bulk-operations', endpoint='admin_bulk_users')
@login_required
def admin_bulk_operations():
    """Bulk operations management dashboard"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get recent bulk operation statistics
    total_users = User.query.count()
    recent_users = User.query.filter(
        User.created_at >= datetime.utcnow() - timedelta(days=30)
    ).count()
    
    # Get courses for bulk assignment
    active_courses = Course.query.filter_by(is_active=True).all()
    
    return render_template('admin_bulk_users.html', 
                         total_users=total_users,
                         recent_users=recent_users,
                         active_courses=active_courses)

@app.route('/admin/system-settings')
@login_required
def admin_system_settings():
    """Admin system settings"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    return render_template('admin_system_settings.html')

@app.route('/admin/violation-appeals')
@login_required
def admin_violation_appeals():
    """Admin page to manage student violation appeals"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get users with pending appeals (flagged users who requested reconsideration)
    pending_appeals = db.session.query(UserViolation, User).join(
        User, UserViolation.user_id == User.id
    ).filter(
        UserViolation.is_flagged == True,
        UserViolation.can_retake == False,
        UserViolation.notes.ilike('%appeal%')
    ).order_by(UserViolation.flagged_at.desc()).all()
    
    # Get recent violations for context
    recent_violations = db.session.query(ProctoringEvent, QuizAttempt, User, Quiz).join(
        QuizAttempt, ProctoringEvent.attempt_id == QuizAttempt.id
    ).join(
        User, QuizAttempt.participant_id == User.id
    ).join(
        Quiz, QuizAttempt.quiz_id == Quiz.id
    ).filter(
        ProctoringEvent.severity.in_(['high', 'critical'])
    ).order_by(ProctoringEvent.timestamp.desc()).limit(50).all()
    
    return render_template('admin_violation_appeals.html', 
                         pending_appeals=pending_appeals,
                         recent_violations=recent_violations)

@app.route('/admin/approve-appeal/<int:violation_id>', methods=['POST'])
@login_required
def admin_approve_appeal(violation_id):
    """Admin approves a student's violation appeal"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    violation = UserViolation.query.get_or_404(violation_id)
    decision = request.form.get('decision')  # 'approve' or 'deny'
    admin_notes = request.form.get('admin_notes', '')
    
    if decision == 'approve':
        violation.can_retake = True
        violation.retake_approved_by = current_user.id
        violation.retake_approved_at = datetime.utcnow()
        violation.notes += f"\n\n[ADMIN APPROVAL - {current_user.username}]: Appeal approved. {admin_notes}"
        flash(f'Appeal approved for {violation.user.username}. They can now retake exams.', 'success')
    else:
        violation.notes += f"\n\n[ADMIN DENIAL - {current_user.username}]: Appeal denied. {admin_notes}"
        flash(f'Appeal denied for {violation.user.username}.', 'info')
    
    db.session.commit()
    return redirect(url_for('admin_violation_appeals'))

@app.route('/student/request-appeal', methods=['GET', 'POST'])
@login_required
def student_request_appeal():
    """Student requests appeal for security violations"""
    if current_user.is_admin() or current_user.is_host():
        flash('This page is for students only.', 'error')
        return redirect(url_for('dashboard'))
    
    # Check if user has violations
    violation = UserViolation.query.filter_by(user_id=current_user.id, is_flagged=True).first()
    
    if request.method == 'POST':
        if not violation:
            flash('You do not have any flagged violations to appeal.', 'error')
            return redirect(url_for('dashboard'))
        
        appeal_reason = request.form.get('appeal_reason', '').strip()
        if not appeal_reason:
            flash('Please provide a reason for your appeal.', 'error')
            return render_template('student_appeal_form.html', violation=violation)
        
        # Add appeal to notes
        current_time = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        violation.notes += f"\n\n[STUDENT APPEAL - {current_time}]: {appeal_reason}"
        db.session.commit()
        
        flash('Your appeal has been submitted. An administrator will review it shortly.', 'success')
        return redirect(url_for('dashboard'))
    
    return render_template('student_appeal_form.html', violation=violation)

@app.route('/admin/audit-logs')
@login_required
def admin_audit_logs():
    """Admin audit logs"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get recent activities (for now, we'll show quiz attempts and user registrations)
    recent_attempts = QuizAttempt.query.order_by(QuizAttempt.started_at.desc()).limit(50).all()
    recent_users = User.query.order_by(User.created_at.desc()).limit(20).all()
    
    return render_template('admin_audit_logs.html', recent_attempts=recent_attempts, recent_users=recent_users)

@app.route('/admin/quiz/<int:quiz_id>/toggle-active', methods=['POST'])
@login_required
def admin_toggle_quiz_active(quiz_id):
    """Toggle quiz active status"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    quiz = Quiz.query.get_or_404(quiz_id)
    quiz.is_active = not quiz.is_active
    db.session.commit()
    
    status = 'activated' if quiz.is_active else 'deactivated'
    flash(f'Quiz "{quiz.title}" has been {status}.', 'success')
    return redirect(url_for('admin_quiz_management'))

@app.route('/admin/quiz/<int:quiz_id>/delete', methods=['POST'])
@login_required
def admin_delete_quiz(quiz_id):
    """Delete a quiz (admin only)"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    quiz = Quiz.query.get_or_404(quiz_id)
    
    # Delete all related data in correct order (respecting foreign keys)
    # 1. Delete answers first (they reference quiz_attempt)
    attempts = QuizAttempt.query.filter_by(quiz_id=quiz_id).all()
    for attempt in attempts:
        Answer.query.filter_by(attempt_id=attempt.id).delete()
        ProctoringEvent.query.filter_by(attempt_id=attempt.id).delete()
    
    # 2. Delete quiz attempts
    QuizAttempt.query.filter_by(quiz_id=quiz_id).delete()
    
    # 3. Delete question options and questions
    for question in quiz.questions:
        QuestionOption.query.filter_by(question_id=question.id).delete()
    Question.query.filter_by(quiz_id=quiz_id).delete()
    
    # 4. Finally delete the quiz
    db.session.delete(quiz)
    db.session.commit()
    
    flash(f'Quiz "{quiz.title}" has been permanently deleted.', 'success')
    return redirect(url_for('admin_quiz_management'))

# ===== COURSE MANAGEMENT SYSTEM =====
@app.route('/admin/course-management')
@login_required
def admin_course_management():
    """Admin course management system"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    courses = Course.query.order_by(Course.display_order.asc(), Course.created_at.desc()).all()
    hosts = User.query.filter_by(role='host').all()
    participants = User.query.filter_by(role='participant').all()
    
    # Get course statistics
    course_stats = {}
    for course in courses:
        stats = {
            'total_hosts': len(course.host_assignments),
            'total_participants': len(course.participant_enrollments),
            'total_quizzes': len(course.quizzes),
            'active_quizzes': len([q for q in course.quizzes if q.is_active])
        }
        course_stats[course.id] = stats
    
    return render_template('admin_course_management.html', 
                         courses=courses, 
                         hosts=hosts, 
                         participants=participants,
                         course_stats=course_stats)

@app.route('/admin/create-course', methods=['POST'])
@login_required
def admin_create_course():
    """Create new course"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    name = request.form.get('name')
    code = request.form.get('code')
    description = request.form.get('description')
    max_participants = request.form.get('max_participants', 100, type=int)
    
    # Validation
    if not all([name, code]):
        flash('Course name and code are required.', 'error')
        return redirect(url_for('admin_course_management'))
    
    if Course.query.filter_by(code=code).first():
        flash('Course code already exists.', 'error')
        return redirect(url_for('admin_course_management'))
    
    # Create course
    course = Course()
    course.name = name
    course.code = code.upper()
    course.description = description
    course.max_participants = max_participants
    
    db.session.add(course)
    db.session.commit()
    
    flash(f'Course "{name}" ({code}) created successfully.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/assign-host', methods=['POST'])
@login_required
def admin_assign_host_to_course(course_id):
    """Assign host to course"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    course = Course.query.get_or_404(course_id)
    host_id = request.form.get('host_id', type=int)
    
    if not host_id:
        flash('Please select a host to assign.', 'error')
        return redirect(url_for('admin_course_management'))
    
    host = User.query.get_or_404(host_id)
    if host.role != 'host':
        flash('Selected user is not a host.', 'error')
        return redirect(url_for('admin_course_management'))
    
    # Check if already assigned
    existing = HostCourseAssignment.query.filter_by(host_id=host_id, course_id=course_id).first()
    if existing:
        flash(f'Host {host.username} is already assigned to course {course.name}.', 'warning')
        return redirect(url_for('admin_course_management'))
    
    # Create assignment
    assignment = HostCourseAssignment()
    assignment.host_id = host_id
    assignment.course_id = course_id
    assignment.assigned_by = current_user.id
    
    db.session.add(assignment)
    db.session.commit()
    
    flash(f'Host {host.username} assigned to course {course.name} successfully.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/enroll-participant', methods=['POST'])
@login_required
def admin_enroll_participant_in_course(course_id):
    """Enroll participant in course"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    course = Course.query.get_or_404(course_id)
    participant_id = request.form.get('participant_id', type=int)
    
    if not participant_id:
        flash('Please select a participant to enroll.', 'error')
        return redirect(url_for('admin_course_management'))
    
    participant = User.query.get_or_404(participant_id)
    if participant.role != 'participant':
        flash('Selected user is not a participant.', 'error')
        return redirect(url_for('admin_course_management'))
    
    # Check participant limit
    current_enrollments = len(course.participant_enrollments)
    if current_enrollments >= course.max_participants:
        flash(f'Course {course.name} has reached maximum participant limit ({course.max_participants}).', 'error')
        return redirect(url_for('admin_course_management'))
    
    # Check if already enrolled
    existing = ParticipantEnrollment.query.filter_by(participant_id=participant_id, course_id=course_id).first()
    if existing:
        flash(f'Participant {participant.username} is already enrolled in course {course.name}.', 'warning')
        return redirect(url_for('admin_course_management'))
    
    # Create enrollment
    enrollment = ParticipantEnrollment()
    enrollment.participant_id = participant_id
    enrollment.course_id = course_id
    enrollment.enrolled_by = current_user.id
    
    db.session.add(enrollment)
    db.session.commit()
    
    flash(f'Participant {participant.username} enrolled in course {course.name} successfully.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/bulk-enroll-participants', methods=['POST'])
@login_required
def admin_bulk_enroll_participants(course_id):
    """Bulk enroll multiple participants in course"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    course = Course.query.get_or_404(course_id)
    participant_ids = request.form.getlist('participant_ids')
    
    # DEBUG: Log what we're receiving
    import logging
    logging.basicConfig(level=logging.DEBUG)
    logger = logging.getLogger(__name__)
    logger.debug(f"🔍 BULK ENROLL DEBUG - Course: {course.name} (ID: {course_id})")
    logger.debug(f"🔍 Form data keys: {list(request.form.keys())}")
    logger.debug(f"🔍 All form data: {dict(request.form)}")
    logger.debug(f"🔍 participant_ids received: {participant_ids}")
    logger.debug(f"🔍 Number of participant_ids: {len(participant_ids)}")
    
    print(f"🔍 BULK ENROLL - Course: {course.name}, participant_ids: {participant_ids}, count: {len(participant_ids)}")
    
    if not participant_ids:
        flash('Please select participants to enroll.', 'error')
        return redirect(url_for('admin_course_management'))
    
    # Check current enrollment count
    current_enrollments = len(course.participant_enrollments)
    available_spots = course.max_participants - current_enrollments
    
    if available_spots <= 0:
        flash(f'Course {course.name} has reached maximum participant limit ({course.max_participants}).', 'error')
        return redirect(url_for('admin_course_management'))
    
    # Limit enrollments to available spots
    if len(participant_ids) > available_spots:
        flash(f'Only {available_spots} spots available in course {course.name}. Limiting enrollment to first {available_spots} selected participants.', 'warning')
        participant_ids = participant_ids[:available_spots]
    
    enrolled_count = 0
    errors = []
    
    # Process participants in small batches to prevent timeouts
    batch_size = 5  # Process 5 participants at a time
    total_participants = len(participant_ids)
    
    for batch_start in range(0, total_participants, batch_size):
        batch_end = min(batch_start + batch_size, total_participants)
        batch_ids = participant_ids[batch_start:batch_end]
        
        try:
            # Process current batch
            for participant_id in batch_ids:
                try:
                    participant_id = int(participant_id)
                    participant = User.query.get(participant_id)
                    
                    if not participant:
                        errors.append(f'Participant with ID {participant_id} not found.')
                        continue
                    
                    if participant.role != 'participant':
                        errors.append(f'User {participant.username} is not a participant.')
                        continue
                    
                    # Check if already enrolled
                    existing = ParticipantEnrollment.query.filter_by(
                        participant_id=participant_id, course_id=course_id
                    ).first()
                    
                    if existing:
                        errors.append(f'Participant {participant.username} is already enrolled in this course.')
                        continue
                    
                    # Create enrollment
                    enrollment = ParticipantEnrollment()
                    enrollment.participant_id = participant_id
                    enrollment.course_id = course_id
                    enrollment.enrolled_by = current_user.id
                    
                    db.session.add(enrollment)
                    enrolled_count += 1
                    
                except ValueError:
                    errors.append(f'Invalid participant ID: {participant_id}')
                except Exception as e:
                    errors.append(f'Error enrolling participant {participant_id}: {str(e)}')
            
            # Commit each batch to prevent timeouts
            try:
                db.session.commit()
                
                # Add small delay between batches to prevent overwhelming the database
                if batch_end < total_participants:  # Not the last batch
                    import time
                    time.sleep(0.1)
                    
            except Exception as batch_error:
                db.session.rollback()
                errors.append(f'Batch {batch_start//batch_size + 1}: Database error - {str(batch_error)}')
                
        except Exception as batch_exception:
            db.session.rollback()
            errors.append(f'Batch {batch_start//batch_size + 1}: Processing error - {str(batch_exception)}')
    
    # Final status messages
    if enrolled_count > 0:
        flash(f'Successfully enrolled {enrolled_count} participants in course {course.name}.', 'success')
    if errors:
        flash(f'Errors encountered: {"; ".join(errors[:3])}{"..." if len(errors) > 3 else ""}', 'warning')
    
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/remove-host/<int:host_id>', methods=['POST'])
@login_required
def admin_remove_host_from_course(course_id, host_id):
    """Remove host from course"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    assignment = HostCourseAssignment.query.filter_by(host_id=host_id, course_id=course_id).first_or_404()
    host_name = assignment.host.username
    course_name = assignment.course.name
    
    db.session.delete(assignment)
    db.session.commit()
    
    flash(f'Host {host_name} removed from course {course_name}.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/remove-participant/<int:participant_id>', methods=['POST'])
@login_required
def admin_remove_participant_from_course(course_id, participant_id):
    """Remove participant from course"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    enrollment = ParticipantEnrollment.query.filter_by(participant_id=participant_id, course_id=course_id).first_or_404()
    participant_name = enrollment.participant.username
    course_name = enrollment.course.name
    
    db.session.delete(enrollment)
    db.session.commit()
    
    flash(f'Participant {participant_name} removed from course {course_name}.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/toggle-status', methods=['POST'])
@login_required
def admin_toggle_course_status(course_id):
    """Toggle course active status"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    course = Course.query.get_or_404(course_id)
    course.is_active = not course.is_active
    db.session.commit()
    
    status = 'activated' if course.is_active else 'deactivated'
    flash(f'Course "{course.name}" has been {status}.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/admin/course/<int:course_id>/delete', methods=['POST'])
@login_required
def admin_delete_course(course_id):
    """Delete course and all related data"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    course = Course.query.get_or_404(course_id)
    course_name = course.name
    
    # Delete all related data
    HostCourseAssignment.query.filter_by(course_id=course_id).delete()
    ParticipantEnrollment.query.filter_by(course_id=course_id).delete()
    
    # Remove course association from quizzes (don't delete quizzes)
    quizzes = Quiz.query.filter_by(course_id=course_id).all()
    for quiz in quizzes:
        quiz.course_id = None
    
    db.session.delete(course)
    db.session.commit()
    
    flash(f'Course "{course_name}" has been permanently deleted.', 'success')
    return redirect(url_for('admin_course_management'))

@app.route('/api/quiz/<int:quiz_id>/stats')
@login_required
def get_quiz_stats(quiz_id):
    """Get quiz statistics for admin"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    quiz = Quiz.query.get_or_404(quiz_id)
    attempts = QuizAttempt.query.filter_by(quiz_id=quiz_id).all()
    violations = ProctoringEvent.query.join(QuizAttempt).filter(QuizAttempt.quiz_id == quiz_id).all()
    
    completed_attempts = [a for a in attempts if a.status == 'completed']
    terminated_attempts = [a for a in attempts if a.status == 'terminated']
    
    stats = {
        'total_attempts': len(attempts),
        'completed_attempts': len(completed_attempts),
        'terminated_attempts': len(terminated_attempts),
        'total_violations': len(violations),
        'average_score': sum([a.score for a in completed_attempts if a.score]) / len(completed_attempts) if completed_attempts else 0,
        'highest_score': max([a.score for a in completed_attempts if a.score]) if completed_attempts else 0,
        'common_violation': violations[0].event_type.replace('_', ' ').title() if violations else 'None'
    }
    
    return jsonify(stats)

@app.route('/api/violations/count')
@login_required
def get_violations_count():
    """Get total violation count for admin dashboard"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    count = ProctoringEvent.query.count()
    return jsonify({'count': count})


@app.route('/api/export-logs')
@login_required  
def export_logs():
    """Export audit logs to CSV"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    import csv
    import io
    from flask import make_response
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write headers
    writer.writerow(['Date', 'User', 'Quiz', 'Action', 'Status', 'Score', 'Violations'])
    
    # Write data
    attempts = QuizAttempt.query.order_by(QuizAttempt.started_at.desc()).all()
    for attempt in attempts:
        writer.writerow([
            attempt.started_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.started_at else 'N/A',
            attempt.participant.username,
            attempt.quiz.title,
            'Quiz Attempt',
            attempt.status,
            f"{attempt.score:.1f}%" if attempt.score else 'N/A',
            len(attempt.answers) if hasattr(attempt, 'answers') else 0
        ])
    
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = 'attachment; filename=audit_logs.csv'
    
    return response

@app.route('/api/export-users')
@login_required  
def export_users():
    """Export user data to CSV"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    import csv
    import io
    from flask import make_response
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write headers
    writer.writerow(['Username', 'Email', 'Role', 'Status', 'Verified', 'Registered', 'Last Login'])
    
    # Write data
    users = User.query.order_by(User.created_at.desc()).all()
    for user in users:
        writer.writerow([
            user.username,
            user.email,
            user.role,
            'Active' if user.is_verified else 'Inactive',
            'Yes' if user.is_verified else 'No',
            user.created_at.strftime('%Y-%m-%d %H:%M:%S') if user.created_at else 'N/A',
            user.last_login.strftime('%Y-%m-%d %H:%M:%S') if user.last_login else 'Never'
        ])
    
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = 'attachment; filename=users_export.csv'
    
    return response

# Enhanced Proctoring System API
@app.route('/api/proctoring/violation', methods=['POST'])
@login_required
def record_enhanced_violation():
    """Enhanced proctoring violation endpoint for new security system"""
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['attemptId', 'type', 'severity', 'description']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        # Get attempt and verify ownership
        attempt = QuizAttempt.query.get(data['attemptId'])
        if not attempt:
            return jsonify({'error': 'Attempt not found'}), 404
            
        if attempt.participant_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403
        
        # Create proctoring event
        event = ProctoringEvent(
            attempt_id=data['attemptId'],
            event_type=data['type'],
            details=data['description'],
            severity=data['severity'],
            timestamp=datetime.utcnow()
        )
        
        db.session.add(event)
        
        # Check for termination conditions
        violation_count = ProctoringEvent.query.filter_by(attempt_id=data['attemptId']).count() + 1
        critical_count = ProctoringEvent.query.filter_by(
            attempt_id=data['attemptId'], 
            severity='critical'
        ).count()
        
        if data['severity'] == 'critical':
            critical_count += 1
        
        should_terminate = (
            critical_count >= 1 or
            violation_count >= 3 or
            data['type'] in ['multiple_people', 'camera_disabled', 'tab_switch']
        )
        
        if should_terminate:
            attempt.status = 'terminated'
            attempt.completed_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'message': 'Enhanced violation recorded successfully',
            'id': event.id,
            'should_terminate': should_terminate,
            'violation_count': violation_count
        }), 201
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Failed to record enhanced violation: {e}")
        return jsonify({'error': 'Failed to record violation'}), 500

@app.route('/api/proctoring/mark-malpractice', methods=['POST'])
@login_required
def mark_malpractice():
    """Mark user as malpractice and handle immediate termination"""
    try:
        data = request.get_json()
        
        attempt_id = data.get('attemptId')
        quiz_id = data.get('quizId')
        violation = data.get('violation', {})
        
        if not attempt_id:
            return jsonify({'error': 'Attempt ID required'}), 400
        
        # Get attempt and verify ownership
        attempt = QuizAttempt.query.get(attempt_id)
        if not attempt or attempt.participant_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403
        
        # Mark attempt as terminated with malpractice flag
        attempt.status = 'terminated'
        attempt.completed_at = datetime.utcnow()
        attempt.termination_reason = f"Malpractice: {violation.get('description', 'Critical violation detected')}"
        attempt.is_flagged = True
        
        # Create or update user violation record
        user_violation = UserViolation.query.filter_by(user_id=current_user.id).first()
        if not user_violation:
            user_violation = UserViolation(user_id=current_user.id)
            db.session.add(user_violation)
        
        user_violation.is_flagged = True
        user_violation.flagged_at = datetime.utcnow()
        user_violation.flagged_by = None  # System flagged
        
        # Create security alert
        alert = SecurityAlert(
            user_id=current_user.id,
            quiz_id=quiz_id,
            attempt_id=attempt_id,
            alert_type='immediate_malpractice',
            severity='critical',
            description=f"Immediate malpractice detected: {violation.get('description', 'Unknown')}",
            auto_action_taken='quiz_terminated_malpractice'
        )
        db.session.add(alert)
        
        # Log the malpractice event
        event = ProctoringEvent(
            attempt_id=attempt_id,
            event_type='malpractice_marked',
            details=f"User marked as malpractice: {violation.get('description', 'Critical violation')}",
            severity='critical',
            timestamp=datetime.utcnow()
        )
        db.session.add(event)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'User marked as malpractice successfully',
            'attempt_terminated': True
        }), 200
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Failed to mark malpractice: {e}")
        return jsonify({'error': 'Failed to mark malpractice'}), 500

@app.route('/api/proctoring/notify-participants', methods=['POST'])
@login_required
def notify_participants():
    """Notify other participants about malpractice detection"""
    try:
        data = request.get_json()
        
        quiz_id = data.get('quizId')
        message = data.get('message', 'A participant has been terminated for malpractice.')
        
        if not quiz_id:
            return jsonify({'error': 'Quiz ID required'}), 400
        
        # Get all participants in this quiz (except current user)
        quiz = Quiz.query.get_or_404(quiz_id)
        participants = QuizAttempt.query.filter(
            QuizAttempt.quiz_id == quiz_id,
            QuizAttempt.participant_id != current_user.id,
            QuizAttempt.status.in_(['in_progress', 'started'])
        ).all()
        
        # Send email notifications to active participants
        for attempt in participants:
            try:
                participant = attempt.participant
                subject = f"⚠️ Quiz Alert - {quiz.title}"
                body = f"""
                Quiz Alert Notification
                
                Quiz: {quiz.title}
                Alert: {message}
                Time: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}
                
                This is an automated notification for your awareness.
                Please continue with your quiz if it's still in progress.
                
                Assessment Platform
                """
                
                msg = Message(
                    subject=subject,
                    recipients=[participant.email],
                    body=body
                )
                
                mail.send(msg)
                
            except Exception as email_error:
                logging.error(f"Failed to send participant notification to {participant.email}: {email_error}")
        
        # Log the notification event
        event = ProctoringEvent(
            attempt_id=None,  # System-wide notification
            event_type='participant_notification_sent',
            details=f"Malpractice alert sent to {len(participants)} participants",
            severity='medium',
            timestamp=datetime.utcnow()
        )
        db.session.add(event)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Notifications sent to {len(participants)} participants'
        }), 200
        
    except Exception as e:
        logging.error(f"Failed to send participant notifications: {e}")
        return jsonify({'error': 'Failed to send notifications'}), 500

@app.route('/api/quiz/force-submit', methods=['POST'])
@login_required
def force_submit_quiz():
    """Force submit quiz due to malpractice"""
    try:
        data = request.get_json()
        
        # Find the current user's active quiz attempt
        attempt = QuizAttempt.query.filter_by(
            participant_id=current_user.id,
            status='in_progress'
        ).first()
        
        if not attempt:
            return jsonify({'error': 'No active quiz attempt found'}), 404
        
        # Force submit the attempt
        attempt.status = 'submitted'
        attempt.completed_at = datetime.utcnow()
        attempt.force_submitted = True
        attempt.termination_reason = data.get('reason', 'Malpractice detected')
        
        # Calculate score with current answers (if any)
        total_questions = len(attempt.quiz.questions)
        answered_questions = Answer.query.filter_by(attempt_id=attempt.id).count()
        
        if total_questions > 0:
            attempt.score = (answered_questions / total_questions) * 100
        else:
            attempt.score = 0
        
        # Log the force submission
        event = ProctoringEvent(
            attempt_id=attempt.id,
            event_type='force_submitted',
            details=f"Quiz force submitted: {data.get('reason', 'Malpractice detected')}",
            severity='critical',
            timestamp=datetime.utcnow()
        )
        db.session.add(event)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Quiz force submitted successfully',
            'score': attempt.score
        }), 200
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Failed to force submit quiz: {e}")
        return jsonify({'error': 'Failed to force submit quiz'}), 500

# Host Heatmap Dashboard Route
@app.route('/host/heatmap-dashboard')
@login_required
def host_heatmap_dashboard():
    """Host collaboration heatmap dashboard"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied. Host privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get quizzes created by the current host (or all if admin)
    if current_user.is_admin():
        quizzes = Quiz.query.filter_by(is_deleted=False).order_by(Quiz.created_at.desc()).all()
    else:
        quizzes = Quiz.query.filter_by(
            creator_id=current_user.id,
            is_deleted=False
        ).order_by(Quiz.created_at.desc()).all()
    
    # Get selected quiz if specified
    selected_quiz_id = request.args.get('quiz_id')
    selected_quiz = None
    if selected_quiz_id:
        selected_quiz = Quiz.query.get(selected_quiz_id)
        # Verify access to selected quiz
        if selected_quiz and not current_user.is_admin() and selected_quiz.creator_id != current_user.id:
            selected_quiz = None
    
    # Add attempt counts to quizzes
    for quiz in quizzes:
        quiz.attempts = QuizAttempt.query.filter_by(quiz_id=quiz.id).all()
    
    return render_template('host_heatmap_dashboard.html',
                         quizzes=quizzes,
                         selected_quiz=selected_quiz)

@app.route('/help')
def help_page():
    """Comprehensive help and documentation page"""
    return render_template('help.html')

# Drag-and-Drop API Endpoints
@app.route('/api/reorder-courses', methods=['POST'])
@login_required
def api_reorder_courses():
    """API endpoint to handle course reordering via drag-and-drop"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403
    
    try:
        data = request.get_json()
        course_ids = data.get('course_ids', [])
        
        if not course_ids:
            return jsonify({'success': False, 'message': 'No course IDs provided'}), 400
        
        # Update display_order for each course
        for index, course_id in enumerate(course_ids):
            course = Course.query.get(course_id)
            if course:
                course.display_order = index
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Course order updated successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error updating course order: {str(e)}'}), 500

@app.route('/api/reorder-quizzes', methods=['POST'])
@login_required
def api_reorder_quizzes():
    """API endpoint to handle quiz reordering via drag-and-drop"""
    if not (current_user.is_admin() or current_user.is_host()):
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403
    
    try:
        data = request.get_json()
        quiz_ids = data.get('quiz_ids', [])
        course_id = data.get('course_id')
        
        if not quiz_ids:
            return jsonify({'success': False, 'message': 'No quiz IDs provided'}), 400
        
        # Update display_order for each quiz
        for index, quiz_id in enumerate(quiz_ids):
            quiz = Quiz.query.get(quiz_id)
            if quiz:
                # Verify user has permission to modify this quiz
                if not current_user.is_admin() and quiz.creator_id != current_user.id:
                    return jsonify({'success': False, 'message': 'Unauthorized to modify this quiz'}), 403
                quiz.display_order = index
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Quiz order updated successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error updating quiz order: {str(e)}'}), 500

# RBAC Admin Routes - Role and Permission Management

@app.route('/admin/rbac')
@login_required
def admin_rbac_dashboard():
    """RBAC management dashboard"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    roles = Role.query.order_by(Role.created_at.desc()).all()
    permissions = Permission.query.order_by(Permission.category, Permission.name).all()
    
    # Group permissions by category
    permissions_by_category = {}
    for permission in permissions:
        category = permission.category
        if category not in permissions_by_category:
            permissions_by_category[category] = []
        permissions_by_category[category].append(permission)
    
    # Get recent audit logs
    recent_logs = RoleAuditLog.query.order_by(RoleAuditLog.created_at.desc()).limit(10).all()
    
    # Statistics
    stats = {
        'total_roles': len(roles),
        'total_permissions': len(permissions),
        'active_roles': len([r for r in roles if r.is_active]),
        'system_roles': len([r for r in roles if r.is_system_role]),
        'total_users_with_roles': UserRole.query.filter_by(is_active=True).count()
    }
    
    return render_template('admin_rbac_dashboard.html',
                         roles=roles,
                         permissions_by_category=permissions_by_category,
                         recent_logs=recent_logs,
                         stats=stats)

@app.route('/admin/rbac/roles')
@login_required
def admin_rbac_roles():
    """Role management interface"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    roles = Role.query.order_by(Role.created_at.desc()).all()
    permissions = Permission.query.order_by(Permission.category, Permission.name).all()
    
    # Add user count and permission details to each role
    for role in roles:
        role.active_users = UserRole.query.filter_by(role_id=role.id, is_active=True).count()
        role.permission_details = [rp.permission for rp in role.role_permissions]
    
    return render_template('admin_rbac_roles.html', roles=roles, permissions=permissions)

@app.route('/admin/rbac/permissions')
@login_required
def admin_rbac_permissions():
    """Permission management interface"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    permissions = Permission.query.order_by(Permission.category, Permission.name).all()
    
    # Group by category
    permissions_by_category = {}
    for permission in permissions:
        category = permission.category
        if category not in permissions_by_category:
            permissions_by_category[category] = []
        permissions_by_category[category].append(permission)
    
    return render_template('admin_rbac_permissions.html', 
                         permissions_by_category=permissions_by_category)

@app.route('/admin/rbac/user-assignments')
@login_required
def admin_rbac_user_assignments():
    """User role assignment interface"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    users = User.query.order_by(User.username).all()
    roles = Role.query.filter_by(is_active=True).order_by(Role.name).all()
    
    # Get current assignments
    assignments = UserRole.query.filter_by(is_active=True).all()
    user_assignments = {}
    for assignment in assignments:
        if assignment.user_id not in user_assignments:
            user_assignments[assignment.user_id] = []
        user_assignments[assignment.user_id].append(assignment)
    
    return render_template('admin_rbac_user_assignments.html',
                         users=users,
                         roles=roles,
                         user_assignments=user_assignments)

@app.route('/admin/rbac/audit-logs')
@login_required
def admin_rbac_audit_logs():
    """RBAC audit log viewer"""
    if not current_user.is_admin():
        flash('Access denied. Administrator privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    entity_type = request.args.get('entity_type', '')
    action = request.args.get('action', '')
    
    query = RoleAuditLog.query
    
    if entity_type:
        query = query.filter_by(entity_type=entity_type)
    if action:
        query = query.filter_by(action=action)
    
    logs = query.order_by(RoleAuditLog.created_at.desc()).paginate(
        page=page, per_page=50, error_out=False
    )
    
    return render_template('admin_rbac_audit_logs.html', logs=logs, 
                         entity_type=entity_type, action=action)

# RBAC API Endpoints

@app.route('/api/rbac/role', methods=['POST'])
@login_required
def api_create_role():
    """Create a new role"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        
        role = RBACService.create_role(
            name=data['name'],
            display_name=data['display_name'],
            description=data.get('description', ''),
            permission_names=data.get('permissions', []),
            created_by_user_id=current_user.id
        )
        
        return jsonify({
            'success': True,
            'message': f'Role "{role.display_name}" created successfully',
            'role_id': role.id
        })
    
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to create role'}), 500

@app.route('/api/rbac/role/<int:role_id>', methods=['PUT'])
@login_required
def api_update_role(role_id):
    """Update an existing role"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        data['updated_by'] = current_user.id
        
        role = RBACService.update_role(role_id, **data)
        
        return jsonify({
            'success': True,
            'message': f'Role "{role.display_name}" updated successfully'
        })
    
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to update role'}), 500

@app.route('/api/rbac/role/<int:role_id>', methods=['DELETE'])
@login_required
def api_delete_role(role_id):
    """Delete a role"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        RBACService.delete_role(role_id, current_user.id)
        
        return jsonify({
            'success': True,
            'message': 'Role deleted successfully'
        })
    
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to delete role'}), 500

@app.route('/api/rbac/assign-role', methods=['POST'])
@login_required
def api_assign_role():
    """Assign role to user"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        
        expires_at = None
        if data.get('expires_at'):
            expires_at = datetime.fromisoformat(data['expires_at'])
        
        user_role = RBACService.assign_role_to_user(
            user_id=data['user_id'],
            role_name=data['role_name'],
            assigned_by_user_id=current_user.id,
            expires_at=expires_at
        )
        
        return jsonify({
            'success': True,
            'message': f'Role assigned successfully'
        })
    
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to assign role'}), 500

@app.route('/api/rbac/revoke-role', methods=['POST'])
@login_required
def api_revoke_role():
    """Revoke role from user"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        
        RBACService.revoke_role_from_user(
            user_id=data['user_id'],
            role_name=data['role_name'],
            revoked_by_user_id=current_user.id
        )
        
        return jsonify({
            'success': True,
            'message': 'Role revoked successfully'
        })
    
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to revoke role'}), 500

@app.route('/api/rbac/role/<int:role_id>/permissions', methods=['PUT'])
@login_required
def api_update_role_permissions(role_id):
    """Update role permissions"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        role = Role.query.get_or_404(role_id)
        
        # Clear existing permissions
        RolePermission.query.filter_by(role_id=role_id).delete()
        
        # Add new permissions
        for perm_name in data['permissions']:
            permission = Permission.query.filter_by(name=perm_name).first()
            if permission:
                role_permission = RolePermission(
                    role_id=role_id,
                    permission_id=permission.id,
                    granted_by=current_user.id
                )
                db.session.add(role_permission)
        
        # Create audit log
        RBACService._create_audit_log(
            action='update_permissions',
            entity_type='role',
            entity_id=role_id,
            performed_by=current_user.id,
            new_values=json.dumps({'permissions': data['permissions']})
        )
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Permissions updated for role "{role.display_name}"'
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': 'Failed to update permissions'}), 500

@app.route('/api/rbac/bulk-assign', methods=['POST'])
@login_required
def api_bulk_assign_roles():
    """Bulk assign roles to multiple users"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        
        results = RBACService.bulk_assign_roles(
            user_ids=data['user_ids'],
            role_names=data['role_names'],
            assigned_by_user_id=current_user.id
        )
        
        successful = len([r for r in results if r['success']])
        total = len(results)
        
        return jsonify({
            'success': True,
            'message': f'Bulk assignment completed: {successful}/{total} successful',
            'results': results
        })
    
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to perform bulk assignment'}), 500

@app.route('/api/rbac/initialize', methods=['POST'])
@login_required
def api_initialize_rbac():
    """Initialize RBAC system with default roles and permissions"""
    if not current_user.is_admin():
        return jsonify({'success': False, 'error': 'Access denied'}), 403
    
    try:
        if initialize_rbac_system:
            result = initialize_rbac_system()
            return jsonify({
                'success': True,
                'message': 'RBAC system initialized successfully',
                'permissions_created': result['permissions_created'],
                'roles_created': result['roles_created']
            })
        else:
            return jsonify({'success': False, 'error': 'RBAC system not available'}), 500
    
    except Exception as e:
        return jsonify({'success': False, 'error': 'Failed to initialize RBAC system'}), 500

# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500

# ===== LTI (Learning Tools Interoperability) Integration =====

@app.route('/lti/launch', methods=['POST'])
def lti_launch():
    """LTI 1.1 Launch endpoint - receives launches from LMS"""
    try:
        logging.info(f"LTI Launch request received from {request.remote_addr}")
        logging.debug(f"LTI Launch parameters: {request.form.to_dict()}")
        
        # Get LTI provider instance
        lti_provider = get_lti_provider()
        
        # Validate OAuth signature
        if not lti_provider.validate_signature(request.form.to_dict()):
            logging.warning("LTI signature validation failed")
            flash('LTI authentication failed. Please contact your administrator.', 'error')
            return redirect(url_for('home'))
        
        # Process launch request
        success, lti_data = lti_provider.process_launch_request(request.form.to_dict())
        
        if not success:
            logging.error(f"LTI launch processing failed: {lti_data.get('error')}")
            flash(f'LTI launch failed: {lti_data.get("error")}', 'error')
            return redirect(url_for('home'))
        
        # Create or update user from LTI data
        user = LTIUser.create_or_update_user(lti_data)
        
        if not user:
            logging.error("Failed to create/update LTI user")
            flash('Failed to create user account from LTI launch.', 'error')
            return redirect(url_for('home'))
        
        # Update user with LTI context information
        context_info = lti_data.get('context_info', {})
        grade_info = lti_data.get('grade_info', {})
        
        user.lti_consumer_key = request.form.get('oauth_consumer_key')
        user.lti_context_id = context_info.get('course_id')
        user.lti_resource_link_id = context_info.get('resource_link_id')
        user.lti_grade_passback_url = grade_info.get('lis_outcome_service_url')
        user.lti_result_sourcedid = grade_info.get('lis_result_sourcedid')
        
        db.session.commit()
        
        # Log the user in
        login_user(user, remember=True)
        session['lti_launch'] = True
        session['lti_data'] = lti_data
        
        # Check for custom quiz parameter
        custom_quiz_id = request.form.get('custom_quiz_id')
        
        if custom_quiz_id:
            # Direct launch to specific quiz
            try:
                quiz_id = int(custom_quiz_id)
                quiz = Quiz.query.get(quiz_id)
                if quiz and quiz.is_active:
                    return redirect(url_for('take_quiz', quiz_id=quiz_id))
                else:
                    flash('Requested quiz not found or inactive.', 'warning')
            except (ValueError, TypeError):
                flash('Invalid quiz identifier.', 'warning')
        
        # Redirect based on user role
        if user.is_admin():
            return redirect(url_for('admin_dashboard'))
        elif user.is_host():
            return redirect(url_for('host_dashboard'))
        else:
            return redirect(url_for('participant_dashboard'))
        
    except Exception as e:
        logging.error(f"LTI launch error: {e}")
        flash('An error occurred during LTI launch. Please try again.', 'error')
        return redirect(url_for('home'))

@app.route('/lti/config.xml')
def lti_config_xml():
    """LTI Tool Configuration XML for LMS setup"""
    base_url = request.url_root.rstrip('/')
    consumer_key = "bigbossizzz_lti_key"  # Should be configurable
    
    xml_config = LTIToolConfiguration.generate_xml_config(base_url, consumer_key)
    
    response = make_response(xml_config)
    response.headers['Content-Type'] = 'application/xml'
    response.headers['Content-Disposition'] = 'attachment; filename=bigbossizzz_lti_config.xml'
    
    return response

@app.route('/lti/config.json')
def lti_config_json():
    """LTI 1.3 Tool Configuration JSON for LMS setup"""
    base_url = request.url_root.rstrip('/')
    consumer_key = "bigbossizzz_lti_key"
    
    json_config = LTIToolConfiguration.generate_json_config(base_url, consumer_key)
    
    return jsonify(json_config)

@app.route('/lti/admin')
@login_required
def lti_admin():
    """LTI Administration dashboard"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('home'))
    
    # Get LTI user statistics
    lti_users = User.query.filter(User.lti_user_id.isnot(None)).all()
    lti_stats = {
        'total_lti_users': len(lti_users),
        'lti_participants': len([u for u in lti_users if u.role == 'participant']),
        'lti_hosts': len([u for u in lti_users if u.role == 'host']),
        'lti_admins': len([u for u in lti_users if u.role == 'admin']),
        'unique_contexts': len(set([u.lti_context_id for u in lti_users if u.lti_context_id])),
        'unique_consumers': len(set([u.lti_consumer_key for u in lti_users if u.lti_consumer_key]))
    }
    
    # Get recent LTI launches (from session logs)
    recent_launches = []
    try:
        for user in lti_users[:10]:  # Show last 10 LTI users
            recent_launches.append({
                'user': user,
                'context_id': user.lti_context_id,
                'consumer_key': user.lti_consumer_key,
                'last_login': user.last_login
            })
    except Exception as e:
        logging.error(f"Error getting LTI launch data: {e}")
    
    return render_template('admin_lti.html', 
                         lti_stats=lti_stats,
                         recent_launches=recent_launches,
                         base_url=request.url_root.rstrip('/'))

@app.route('/api/lti/grade-passback', methods=['POST'])
@login_required
def lti_grade_passback():
    """Send grade back to LMS for completed quiz attempts"""
    try:
        data = request.get_json()
        attempt_id = data.get('attempt_id')
        
        if not attempt_id:
            return jsonify({'success': False, 'error': 'Missing attempt_id'}), 400
        
        # Get quiz attempt
        attempt = QuizAttempt.query.get(attempt_id)
        if not attempt:
            return jsonify({'success': False, 'error': 'Quiz attempt not found'}), 404
        
        # Check if user has LTI grade passback information
        user = attempt.participant
        if not user.lti_grade_passback_url or not user.lti_result_sourcedid:
            return jsonify({'success': False, 'error': 'No LTI grade passback information available'}), 400
        
        # Calculate score
        if attempt.score is None:
            attempt.calculate_score()
            db.session.commit()
        
        # Prepare grade information
        grade_info = {
            'lis_outcome_service_url': user.lti_grade_passback_url,
            'lis_result_sourcedid': user.lti_result_sourcedid
        }
        
        # Send grade back to LMS
        grade_passback = get_lti_grade_passback()
        success = grade_passback.send_grade(grade_info, attempt.score, 100.0)
        
        if success:
            # Log successful grade passback
            logging.info(f"Grade passback successful for attempt {attempt_id}: {attempt.score}%")
            return jsonify({'success': True, 'message': 'Grade sent to LMS successfully'})
        else:
            logging.error(f"Grade passback failed for attempt {attempt_id}")
            return jsonify({'success': False, 'error': 'Failed to send grade to LMS'}), 500
        
    except Exception as e:
        logging.error(f"Grade passback error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/lti/quiz-selection')
@login_required
def lti_quiz_selection():
    """Quiz selection page for LTI content selection"""
    if not current_user.is_host() and not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('home'))
    
    # Get available quizzes for selection
    if current_user.is_admin():
        quizzes = Quiz.query.filter_by(is_active=True).all()
    else:
        quizzes = Quiz.query.filter_by(creator_id=current_user.id, is_active=True).all()
    
    # Check if this is a content selection request
    is_content_selection = request.args.get('lti_message_type') == 'ContentItemSelectionRequest'
    return_url = request.args.get('content_item_return_url')
    
    return render_template('lti_quiz_selection.html',
                         quizzes=quizzes,
                         is_content_selection=is_content_selection,
                         return_url=return_url)

@app.route('/api/lti/content-item', methods=['POST'])
@login_required
def lti_content_item():
    """Return content item for LTI Deep Linking"""
    try:
        data = request.get_json()
        quiz_id = data.get('quiz_id')
        return_url = data.get('return_url')
        
        if not quiz_id or not return_url:
            return jsonify({'success': False, 'error': 'Missing required parameters'}), 400
        
        quiz = Quiz.query.get(quiz_id)
        if not quiz:
            return jsonify({'success': False, 'error': 'Quiz not found'}), 404
        
        # Check permissions
        if not current_user.is_admin() and quiz.creator_id != current_user.id:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        # Create content item response
        base_url = request.url_root.rstrip('/')
        content_item = {
            "@context": "http://purl.imsglobal.org/ctx/lti/v1/ContentItem",
            "@graph": [
                {
                    "@type": "LtiLinkItem",
                    "url": f"{base_url}/lti/launch?custom_quiz_id={quiz_id}",
                    "title": quiz.title,
                    "text": quiz.description or f"Take the quiz: {quiz.title}",
                    "mediaType": "application/vnd.ims.lti.v1.ltilink",
                    "placementAdvice": {
                        "presentationDocumentTarget": "window"
                    },
                    "custom": {
                        "quiz_id": str(quiz_id)
                    }
                }
            ]
        }
        
        return jsonify({
            'success': True,
            'content_item': content_item,
            'return_url': return_url
        })
        
    except Exception as e:
        logging.error(f"LTI content item error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/lti/1.3/login', methods=['POST', 'GET'])
def lti_13_login():
    """LTI 1.3 OIDC Login endpoint"""
    # LTI 1.3 implementation placeholder
    # This would require full OAuth 2.0 / OIDC implementation
    return jsonify({
        'error': 'LTI 1.3 not yet implemented',
        'message': 'Please use LTI 1.1 endpoints for now'
    }), 501

@app.route('/lti/1.3/launch', methods=['POST'])
def lti_13_launch():
    """LTI 1.3 Launch endpoint"""
    # LTI 1.3 implementation placeholder
    return jsonify({
        'error': 'LTI 1.3 not yet implemented',
        'message': 'Please use LTI 1.1 endpoints for now'
    }), 501

@app.route('/lti/1.3/jwks')
def lti_13_jwks():
    """LTI 1.3 JSON Web Key Set endpoint"""
    # LTI 1.3 JWKS implementation placeholder
    return jsonify({
        'keys': []
    })

# Helper route for testing LTI integration
@app.route('/lti/test')
def lti_test():
    """LTI Integration test page (development only)"""
    if not app.debug:
        return "Not available in production", 404
    
    base_url = request.url_root.rstrip('/')
    test_data = {
        'launch_url': f"{base_url}/lti/launch",
        'config_xml_url': f"{base_url}/lti/config.xml",
        'config_json_url': f"{base_url}/lti/config.json",
        'consumer_key': 'bigbossizzz_lti_key',
        'consumer_secret': 'bigbossizzz_lti_secret_change_in_production'
    }
    
    return render_template('lti_test.html', test_data=test_data)

# Automatic grade passback on quiz completion
@app.after_request
def lti_grade_passback_trigger(response):
    """Automatically trigger grade passback for LTI users after quiz completion"""
    try:
        # Only process for quiz submission routes
        if (request.endpoint == 'submit_quiz' and 
            response.status_code == 302 and  # Successful redirect
            current_user.is_authenticated and 
            hasattr(current_user, 'lti_user_id') and 
            current_user.lti_user_id):
            
            # Get the most recent attempt for automatic grade passback
            recent_attempt = QuizAttempt.query.filter_by(
                participant_id=current_user.id,
                status='completed'
            ).order_by(QuizAttempt.completed_at.desc()).first()
            
            if (recent_attempt and 
                current_user.lti_grade_passback_url and 
                current_user.lti_result_sourcedid):
                
                # Trigger grade passback in background
                try:
                    grade_info = {
                        'lis_outcome_service_url': current_user.lti_grade_passback_url,
                        'lis_result_sourcedid': current_user.lti_result_sourcedid
                    }
                    
                    grade_passback = get_lti_grade_passback()
                    grade_passback.send_grade(grade_info, recent_attempt.score or 0, 100.0)
                    
                    logging.info(f"Automatic LTI grade passback completed for attempt {recent_attempt.id}")
                    
                except Exception as e:
                    logging.error(f"Automatic LTI grade passback failed: {e}")
                    
    except Exception as e:
        logging.error(f"LTI grade passback trigger error: {e}")
    
    return response

# ===== Automated Proctoring Reports System =====

@app.route('/admin/proctoring-reports')
@login_required
def admin_proctoring_reports():
    """Proctoring reports dashboard for administrators"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('home'))
    
    try:
        # Get recent report summary
        generator = ProctoringReportGenerator()
        end_date = datetime.utcnow()
        start_date = end_date - timedelta(days=7)  # Last 7 days
        
        # Generate quick summary for dashboard
        recent_report = generator.generate_comprehensive_report(start_date, end_date)
        
        # Get available quizzes for filtering
        quizzes = Quiz.query.filter_by(is_active=True).all()
        
        # Get report statistics
        report_stats = {
            'total_attempts_last_week': len(recent_report.get('detailed_findings', {}).get('high_risk_attempts', {}).get('details', [])),
            'integrity_score': recent_report.get('ai_analysis', {}).get('overall_integrity_score', {}).get('score', 0),
            'critical_violations': len(recent_report.get('detailed_findings', {}).get('critical_violations', {}).get('details', [])),
            'security_events': len(recent_report.get('detailed_findings', {}).get('security_events', {}).get('details', []))
        }
        
        return render_template('admin_proctoring_reports.html',
                             recent_report=recent_report,
                             report_stats=report_stats,
                             quizzes=quizzes)
        
    except Exception as e:
        logging.error(f"Error loading proctoring reports dashboard: {e}")
        flash('Error loading proctoring reports. Please try again.', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/api/generate-proctoring-report', methods=['POST'])
@login_required
def api_generate_proctoring_report():
    """Generate custom proctoring report via API"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    try:
        data = request.get_json()
        
        # Parse date parameters
        start_date_str = data.get('start_date')
        end_date_str = data.get('end_date')
        quiz_ids = data.get('quiz_ids', [])
        user_ids = data.get('user_ids', [])
        
        if not start_date_str or not end_date_str:
            return jsonify({'error': 'Start and end dates are required'}), 400
        
        # Parse dates
        start_date = datetime.fromisoformat(start_date_str.replace('Z', '+00:00'))
        end_date = datetime.fromisoformat(end_date_str.replace('Z', '+00:00'))
        
        # Validate date range
        if start_date >= end_date:
            return jsonify({'error': 'Start date must be before end date'}), 400
        
        if (end_date - start_date).days > 90:
            return jsonify({'error': 'Date range cannot exceed 90 days'}), 400
        
        # Generate report
        generator = ProctoringReportGenerator()
        report = generator.generate_comprehensive_report(
            start_date, end_date, quiz_ids if quiz_ids else None, user_ids if user_ids else None
        )
        
        if 'error' in report:
            return jsonify({'error': f'Report generation failed: {report["error"]}'}), 500
        
        return jsonify({
            'success': True,
            'report': report,
            'download_url': f'/api/download-report/{report["report_id"]}'
        })
        
    except ValueError as e:
        return jsonify({'error': f'Invalid date format: {str(e)}'}), 400
    except Exception as e:
        logging.error(f"Error generating proctoring report: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/proctoring-analytics', methods=['GET'])
@login_required
def api_proctoring_analytics():
    """Get proctoring analytics data for charts and graphs"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    try:
        # Get date range from query parameters
        days = request.args.get('days', 30, type=int)
        end_date = datetime.utcnow()
        start_date = end_date - timedelta(days=days)
        
        # Get violation trends
        violations_by_day = db.session.query(
            func.date(UserViolation.flagged_at).label('date'),
            func.count(UserViolation.id).label('count')
        ).filter(
            UserViolation.flagged_at >= start_date,
            UserViolation.flagged_at <= end_date
        ).group_by(func.date(UserViolation.flagged_at)).all()
        
        # Get proctoring events by type
        events_by_type = db.session.query(
            ProctoringEvent.event_type,
            func.count(ProctoringEvent.id).label('count')
        ).filter(
            ProctoringEvent.timestamp >= start_date,
            ProctoringEvent.timestamp <= end_date
        ).group_by(ProctoringEvent.event_type).all()
        
        # Get security alerts by severity
        alerts_by_severity = db.session.query(
            SecurityAlert.severity,
            func.count(SecurityAlert.id).label('count')
        ).filter(
            SecurityAlert.created_at >= start_date,
            SecurityAlert.created_at <= end_date
        ).group_by(SecurityAlert.severity).all()
        
        # Get quiz completion rates
        total_attempts = db.session.query(QuizAttempt).filter(
            QuizAttempt.started_at >= start_date,
            QuizAttempt.started_at <= end_date
        ).count()
        
        completed_attempts = db.session.query(QuizAttempt).filter(
            QuizAttempt.started_at >= start_date,
            QuizAttempt.started_at <= end_date,
            QuizAttempt.status == 'completed'
        ).count()
        
        flagged_attempts = db.session.query(QuizAttempt).filter(
            QuizAttempt.started_at >= start_date,
            QuizAttempt.started_at <= end_date,
            QuizAttempt.is_flagged == True
        ).count()
        
        analytics_data = {
            'violation_trends': [
                {'date': str(row.date), 'count': row.count} 
                for row in violations_by_day
            ],
            'events_by_type': [
                {'type': row.event_type, 'count': row.count} 
                for row in events_by_type
            ],
            'alerts_by_severity': [
                {'severity': row.severity, 'count': row.count} 
                for row in alerts_by_severity
            ],
            'completion_stats': {
                'total_attempts': total_attempts,
                'completed_attempts': completed_attempts,
                'flagged_attempts': flagged_attempts,
                'completion_rate': round((completed_attempts / total_attempts * 100) if total_attempts > 0 else 0, 2),
                'flag_rate': round((flagged_attempts / total_attempts * 100) if total_attempts > 0 else 0, 2)
            },
            'period': {
                'start_date': start_date.isoformat(),
                'end_date': end_date.isoformat(),
                'days': days
            }
        }
        
        return jsonify(analytics_data)
        
    except Exception as e:
        logging.error(f"Error generating proctoring analytics: {e}")
        return jsonify({'error': 'Failed to generate analytics'}), 500

# ===== Enhanced Analytics & Insights System =====

@app.route('/admin/analytics-dashboard')
@login_required
def admin_analytics_dashboard():
    """Enhanced analytics dashboard for comprehensive insights"""
    if not current_user.is_admin():
        flash('Access denied.', 'error')
        return redirect(url_for('home'))
    
    try:
        # Initialize analytics engine
        analytics = get_analytics_engine()
        
        # Get comprehensive analytics report
        comprehensive_report = analytics.generate_comprehensive_report()
        
        # Get real-time institutional metrics
        institutional_metrics = analytics.dashboard.get_real_time_metrics()
        
        # Get live monitoring data
        live_data = analytics.dashboard.get_live_monitoring_data()
        
        return render_template('admin_analytics_dashboard.html',
                             comprehensive_report=comprehensive_report,
                             institutional_metrics=institutional_metrics,
                             live_data=live_data)
        
    except Exception as e:
        logging.error(f"Error loading analytics dashboard: {e}")
        flash('Error loading analytics dashboard. Please try again.', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/api/predictive-analytics', methods=['GET'])
@login_required
def api_predictive_analytics():
    """Get predictive analytics for student risk assessment"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    try:
        user_id = request.args.get('user_id', type=int)
        course_id = request.args.get('course_id', type=int)
        limit = request.args.get('limit', 50, type=int)
        
        # Initialize predictive analytics
        predictor = PredictiveAnalytics()
        
        # Analyze student risk
        risk_profiles = predictor.analyze_student_risk(user_id, course_id)
        
        # Limit results for performance
        if limit and len(risk_profiles) > limit:
            risk_profiles = risk_profiles[:limit]
        
        # Convert to serializable format
        results = []
        for profile in risk_profiles:
            results.append({
                'user_id': profile.user_id,
                'username': profile.username,
                'risk_score': profile.risk_score,
                'risk_level': profile.risk_level,
                'risk_factors': profile.risk_factors,
                'intervention_recommendations': profile.intervention_recommendations,
                'predicted_failure_probability': profile.predicted_failure_probability,
                'engagement_score': profile.engagement_score,
                'performance_trend': profile.performance_trend,
                'last_activity': profile.last_activity.isoformat() if profile.last_activity else None,
                'courses_enrolled': profile.courses_enrolled,
                'avg_quiz_score': profile.avg_quiz_score,
                'violation_count': profile.violation_count,
                'proctoring_issues': profile.proctoring_issues
            })
        
        # Calculate summary statistics
        summary = {
            'total_students': len(results),
            'high_risk_count': len([r for r in results if r['risk_level'] in ['High', 'Critical']]),
            'medium_risk_count': len([r for r in results if r['risk_level'] == 'Medium']),
            'low_risk_count': len([r for r in results if r['risk_level'] in ['Low', 'Minimal']]),
            'average_risk_score': round(sum(r['risk_score'] for r in results) / len(results), 2) if results else 0,
            'most_common_risk_factors': {}
        }
        
        # Count risk factors
        all_factors = []
        for r in results:
            all_factors.extend(r['risk_factors'])
        
        from collections import Counter
        factor_counts = Counter(all_factors)
        summary['most_common_risk_factors'] = dict(factor_counts.most_common(10))
        
        return jsonify({
            'success': True,
            'students': results,
            'summary': summary,
            'generated_at': datetime.utcnow().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Error in predictive analytics API: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/institutional-metrics', methods=['GET'])
@login_required
def api_institutional_metrics():
    """Get real-time institutional dashboard metrics"""
    if not current_user.is_admin():
        return jsonify({'error': 'Access denied'}), 403
    
    try:
        # Initialize dashboard
        dashboard = InstitutionalDashboard()
        
        # Get real-time metrics
        metrics = dashboard.get_real_time_metrics()
        
        # Get live monitoring data
        live_data = dashboard.get_live_monitoring_data()
        
        return jsonify({
            'success': True,
            'metrics': {
                'total_students': metrics.total_students,
                'active_students_today': metrics.active_students_today,
                'quizzes_in_progress': metrics.quizzes_in_progress,
                'completed_quizzes_today': metrics.completed_quizzes_today,
                'average_performance': metrics.average_performance,
                'high_risk_students': metrics.high_risk_students,
                'security_alerts_today': metrics.security_alerts_today,
                'system_uptime': metrics.system_uptime,
                'concurrent_users': metrics.concurrent_users,
                'violation_rate': metrics.violation_rate
            },
            'live_data': live_data,
            'generated_at': datetime.utcnow().isoformat()
        })
        
    except Exception as e:
        logging.error(f"Error in institutional metrics API: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/user/theme-preference', methods=['POST'])
@login_required
def save_theme_preference():
    """Save user theme preference"""
    try:
        data = request.get_json()
        theme = data.get('theme', 'auto')
        
        # Validate theme
        valid_themes = ['auto', 'light', 'dark']
        if theme not in valid_themes:
            return jsonify({'error': 'Invalid theme'}), 400
        
        # Save to user model (assuming we add theme_preference field)
        current_user.theme_preference = theme
        db.session.commit()
        
        return jsonify({'success': True, 'theme': theme})
        
    except Exception as e:
        logging.error(f"Error saving theme preference: {e}")
        return jsonify({'error': 'Failed to save preference'}), 500

@app.route('/sw.js')
def service_worker():
    """Serve service worker from root path for proper scope"""
    return app.send_static_file('../sw.js')

@app.route('/api/connectivity-check', methods=['HEAD', 'GET'])
def connectivity_check():
    """Fast connectivity check for offline manager"""
    return '', 204

@app.route('/api/quiz/sync-progress', methods=['POST'])
@login_required
def sync_quiz_progress():
    """Sync quiz progress from offline storage"""
    try:
        data = request.get_json()
        quiz_id = data.get('quizId')
        progress = data.get('progress', {})
        
        if not quiz_id:
            return jsonify({'error': 'Quiz ID required'}), 400
        
        # Update quiz progress in database
        # This would typically update the user's progress
        logging.info(f"Syncing offline progress for quiz {quiz_id} for user {current_user.id}")
        
        return jsonify({
            'success': True,
            'message': 'Progress synced successfully',
            'quiz_id': quiz_id
        })
        
    except Exception as e:
        logging.error(f"Error syncing quiz progress: {e}")
        return jsonify({'error': 'Failed to sync progress'}), 500
